"""
HydroMind Pro - Upgraded Edition
================================

This script extends the original HydroMind Pro platform by adding several new
analytics and decision‑support features. Enhancements include:

1. **Water Quality Index (WQI)** calculation using the weighted arithmetic
   method. Users can specify permissible limits for key parameters or rely on
   reasonable defaults (e.g. WHO/BIS standards). Each sample receives an
   overall WQI score and a corresponding classification (Excellent, Good,
   Poor, etc.) following the ranges outlined by Brown et al. (1972)【378712564134928†L400-L417】.

2. **Irrigation Water Classification** that determines salinity and sodium
   hazard classes (C1–C4 and S1–S4). The classification scheme follows
   guidance from the North Dakota State University Extension Service, where
   low‑salinity water (C1) is suitable for most crops and soils, whereas
   very high salinity water (C4) is only suitable under special circumstances;
   similarly, low sodium (S1) poses little hazard while very high sodium (S4)
   is generally unsuitable for irrigation【669674659774536†L584-L612】.

3. **Statistical Summary** including descriptive statistics, Pearson
   correlation matrix and an interactive heatmap. A scatter matrix is also
   provided for quick visual inspection of relationships between parameters.

4. **Unsupervised Clustering** of sampling locations based on their
   hydrochemical signatures using k‑means. Users can choose the number of
   clusters and visualise results on a scatter plot.

The rest of the application retains the original capabilities:
chemical diagrams (Piper, Schoeller, Gibbs, Stiff), GIS mapping with
interpolation, SAR and Wilcox plots, lithology parsing, OCR, predictive
analysis, Cooper‑Jacob pumping test interpretation, AI‑generated reports and
project archiving. Authentication and user management remain unchanged.

To run this app, install the required packages (streamlit, pandas, folium,
plotly, scikit‑learn, etc.) and launch via `streamlit run hydromind_pro_upgrade.py`.

"""

import streamlit as st
import pandas as pd
import folium
from streamlit_folium import st_folium
import branca.colormap as cm
import io
from google import genai
from google.genai import types
from datetime import datetime
import sqlite3
import hashlib
import matplotlib.pyplot as plt
from matplotlib.patches import Polygon
import numpy as np
from scipy.interpolate import griddata, Rbf
import geojsoncontour
import json
from docx import Document
import plotly.graph_objects as go
import plotly.express as px
import base64

# Additional imports for new features
from sklearn.cluster import KMeans
from sklearn.preprocessing import StandardScaler
from sklearn.decomposition import PCA
import seaborn as sns
from matplotlib.backends.backend_pdf import PdfPages
from pykrige.ok import OrdinaryKriging
from shapely.geometry import MultiPoint, Point, Polygon as ShapelyPoly
from shapely.ops import unary_union
from sklearn.model_selection import LeaveOneOut
from sklearn.metrics import mean_squared_error, mean_absolute_error, r2_score


# --- HELPER: MAP & GEOSTATS ---
def calculate_idw(x, y, z, xi, yi, power=2):
    """Simple Inverse Distance Weighting"""
    dist = np.sqrt((x[:, None, None] - xi)**2 + (y[:, None, None] - yi)**2)
    # Avoid div by zero
    dist = np.where(dist < 1e-10, 1e-10, dist)
    weights = 1.0 / (dist ** power)
    zi = np.sum(weights * z[:, None, None], axis=0) / np.sum(weights, axis=0)
    return zi

def clip_grid_to_convex_hull(xi, yi, zi, x_pts, y_pts, boundary_poly=None):
    """Clip interpolation grid to convex hull or custom boundary"""
    if boundary_poly:
        # Use provided polygon
        clip_poly = boundary_poly
    else:
        # Calc Convex Hull
        points = np.column_stack((x_pts, y_pts))
        if len(points) < 3: return zi
        hull = MultiPoint(points).convex_hull
        clip_poly = hull

    # Create mask
    # This acts on the grid points
    mask = np.zeros_like(zi, dtype=bool)
    rows, cols = xi.shape
    
    # Check each point (this can be slow, optimize if needed)
    # For speed, using matplotlib Path
    from matplotlib.path import Path
    
    if isinstance(clip_poly, ShapelyPoly):
        ex, ey = clip_poly.exterior.coords.xy
        poly_verts = list(zip(ex, ey))
        path = Path(poly_verts)
        points_grid = np.column_stack((xi.flatten(), yi.flatten()))
        mask_flat = path.contains_points(points_grid)
        mask = mask_flat.reshape(rows, cols)
        
        zi_clipped = np.where(mask, zi, np.nan)
        return zi_clipped
    else:
        # fallback
        return zi

def calculate_flow_direction(zi, xi, yi):
    """Calculate gradient (flow direction) from grid"""
    # np.gradient returns (dy, dx)
    dy, dx = np.gradient(zi) 
    # Current grid spacing (assuming uniform for visualization or use actual)
    # We just need direction
    return dx, dy

def generate_contours_advanced(df, param, levels, colormap, technique, resolution_factor=100, variogram_model='linear', range_val=None, sill_val=None, nugget_val=None, clip_hull=True, boundary_json=None):
    if len(df) < 4: return None, None, None, None, None, None
    
    x = df['Lon'].values
    y = df['Lat'].values
    z = df[param].values
    
    # Grid definition
    xi = np.linspace(min(x), max(x), resolution_factor)
    yi = np.linspace(min(y), max(y), resolution_factor)
    xi, yi = np.meshgrid(xi, yi)
    
    zi = None
    sigma_sq = None # Uncertainty
    
    # 1. INTERPOLATION
    if technique == "IDW":
        zi = calculate_idw(x, y, z, xi, yi)
        
    elif technique == "Kriging (Ordinary)":
        try:
            # Basic Ordinary Kriging
            # We can expose variogram params later if needed
            ok = OrdinaryKriging(
                x, y, z, 
                variogram_model=variogram_model,
                verbose=False, enable_plotting=False,
                nlags=6,
                variogram_parameters={'sill': sill_val, 'range': range_val, 'nugget': nugget_val} if sill_val else None
            )
            zi, ss = ok.execute('grid', np.linspace(min(x), max(x), resolution_factor), np.linspace(min(y), max(y), resolution_factor))
            sigma_sq = ss # Variance
        except Exception as e:
            st.error(t("kriging_error").format(error=e))
            return None, None, None, None, None, None

    elif technique == "Spline / Cubic":
        zi = griddata((x, y), z, (xi, yi), method='cubic')
        
    elif technique == "Natural Neighbor / Linear":
        zi = griddata((x, y), z, (xi, yi), method='linear')
    
    # 2. CLIPPING
    boundary_poly = None
    if boundary_json:
        try:
            # Parse GeoJSON to Shapely
            from shapely.geometry import shape
            # Assuming FeatureCollection or Feature
            if boundary_json.get('type') == 'FeatureCollection':
                polys = [shape(f['geometry']) for f in boundary_json['features']]
                boundary_poly = unary_union(polys)
            elif boundary_json.get('type') == 'Feature':
                boundary_poly = shape(boundary_json['geometry'])
        except:
            pass

    if clip_hull or boundary_poly:
        zi = clip_grid_to_convex_hull(xi, yi, zi, x, y, boundary_poly)
        if sigma_sq is not None:
             sigma_sq = clip_grid_to_convex_hull(xi, yi, sigma_sq, x, y, boundary_poly)

    # 3. CONTOURING
    # Handle NaNs
    if zi is None or np.isnan(zi).all(): return None, None, None, None, None, None

    contour = plt.figure()
    ax = contour.add_subplot(111)
    
    # Filled contours
    try:
        cntr = ax.contourf(xi, yi, zi, levels=levels, cmap=colormap, alpha=0.7)
        cbar = plt.colorbar(cntr, ax=ax) # Dummy colorbar to get colors
    except:
        plt.close(contour)
        return None, None, None, None, None, None
        
    geojson = geojsoncontour.contourf_to_geojson(
        contourf=cntr,
        min_angle_deg=3.0,
        ndigits=3,
        stroke_width=2,
        fill_opacity=0.5
    )
    plt.close(contour)
    
    return None, geojson, cm.LinearColormap(colors=['blue', 'green', 'yellow', 'red'], vmin=min(z), vmax=max(z)), zi, sigma_sq, {'xi': xi, 'yi': yi}

def perform_cv_analysis(df, param, method, variogram_model='linear'):
    """Quick LOOCV"""
    x = df['Lon'].values
    y = df['Lat'].values
    z = df[param].values
    ids = df.get("ID", df.index.astype(str)).values
    
    predicted = []
    actual = []
    residuals = []
    p_ids = []
    
    loo = LeaveOneOut()
    
    # Limit for performance on big datasets
    if len(df) > 100:
        indices = np.random.choice(len(df), 100, replace=False)
    else:
        indices = np.arange(len(df))

    # This loop is simplified/naive for demo. 
    # For Kriging proper CV exists in pykrige.
    
    for train_ix, test_ix in loo.split(indices):
        idx_test = indices[test_ix[0]]
        train_idxs = np.delete(indices, test_ix[0])
        
        x_train, y_train = x[train_idxs], y[train_idxs]
        z_train = z[train_idxs]
        
        x_test, y_test = x[idx_test], y[idx_test]
        z_actual = z[idx_test]
        
        z_pred = np.nan
        
        try:
            if method == "IDW":
                # Simple IDW
                dist = np.sqrt((x_train - x_test)**2 + (y_train - y_test)**2)
                w = 1.0 / (dist**2 + 1e-9)
                z_pred = np.sum(w * z_train) / np.sum(w)
                
            elif method == "Kriging (Ordinary)":
                 ok = OrdinaryKriging(x_train, y_train, z_train, variogram_model=variogram_model, verbose=False)
                 val, ss = ok.execute("points", [x_test], [y_test])
                 z_pred = val[0]
                 
            else: # Linear/Cubic
                 # Griddata requires at least 3 points and convex hull issue
                 # We simply use nearest for stability in this quick Hack or fallback
                 # Actually griddata on single point
                 z_pred = griddata((x_train, y_train), z_train, (x_test, y_test), method='nearest')
                 
        except:
            z_pred = np.nan
            
        if not np.nan in [z_pred]:
            predicted.append(z_pred)
            actual.append(z_actual)
            residuals.append(z_actual - z_pred)
            p_ids.append(ids[idx_test])
            
    if not predicted: return None
    
    results = pd.DataFrame({
        "ID": p_ids,
        "Actual": actual,
        "Predicted": predicted,
        "Residual": residuals,
        "AbsResidual": np.abs(residuals)
    })
    
    metrics = {
        "RMSE": np.sqrt(mean_squared_error(actual, predicted)),
        "MAE": mean_absolute_error(actual, predicted),
        "R2": r2_score(actual, predicted)
    }
    
    return metrics, results

# --- 1. SETTINGS ---
# Configuration helper for Google Gemini
def get_genai_client():
    # Require that the user has accepted the AI usage terms; if not, refuse
    if not st.session_state.get("ai_accept", False):
        return None
        
    # Only use custom key logic for now, built-in key removed as per request
    key = st.session_state.get("user_api_key", "")
    if not key:
        return None
        
    try:
        return genai.Client(api_key=key)
    except Exception:
        return None

import time

def call_gemini_with_retry(client, model, contents, retries=3, delay=1):
    """
    Calls Gemini API with retry logic for 503/Overloaded errors.
    """
    last_exception = None
    for attempt in range(retries):
        try:
            return client.models.generate_content(model=model, contents=contents)
        except Exception as e:
            last_exception = e
            error_str = str(e)
            # Check for 503 or overload related messages
            if "503" in error_str or "overloaded" in error_str.lower() or "resource exhausted" in error_str.lower():
                if attempt < retries - 1:
                    time.sleep(delay * (2 ** attempt)) # Exponential backoff: 1s, 2s, 4s
                    continue
            # If it's not a temporary error, raise immediately
            raise e
    raise last_exception

# --- LANGUAGE ---
LANGS = {
    "en": {
        "language": "Language",
        "home": "🏠 Home",
        "new_analysis": "🧪 New Analysis",
        "archive": "🗄️ Archive",
        "logout": "Log out",
        "home_sub": "AI‑Assisted Hydrogeological Analysis Platform",
        "chemistry": "🧪 Chemistry",
        "gis": "🗺️ GIS",
        "agriculture": "🚜 Agriculture",
        "ai": "🤖 AI",
        "what_can": "### 🚀 What can you do?",
        "start": "👉 Start New Analysis",
        "analysis_panel": "🛠️ Analysis Panel",
        "data_input": "Data Input:",
        "upload_excel": "📁 Upload Excel",
        "manual_entry": "✍️ Manual Entry",
        "excel_file": "Excel File (.xlsx)",
        "ai_fill": "🤖 AI Data Fill",
        "ai_caption": "Paste raw values (rows, CSV, or free text). AI will build the table.",
        "ai_input": "AI input",
        "random_rows": "Random rows",
        "random_fill": "Random Fill",
        "ai_fill_btn": "AI Fill Table",
        "active_cols_caption": "Choose active columns for Statistics/Clustering (hidden columns are not auto‑detected).",
        "active_cols": "Active columns",
        "tabs_agri": "🌾 Agriculture",
        "tabs_pump": "📉 Pumping",
        "tabs_lith": "🪨 Lithology",
        "tabs_darcy": "🧮 Darcy",
        "tabs_forecast": "🔮 Forecast",
        "tabs_report": "🤖 Report",
        "tabs_save": "💾 Save",
        "tabs_stats": "📊 Statistics",
        "tabs_cluster": "🔗 Clustering",
        "tabs_level": "💧 Water Level",
        "no_data": "No data",
        "na_required": "Na, Ca, Mg are required.",
        "discharge": "Discharge (Q)",
        "ai_input_text": "##### 🤖 AI Input (Text)",
        "ai_parse": "AI Parse",
        "table": "##### 📝 Table",
        "image": "Image",
        "analyze": "Analyze",
        "generate": "Generate Report",
        "generating": "Generating...",
        "download_word": "📄 Download Word",
        "parameter": "Parameter",
        "method": "Method",
        "name": "Name",
        "save": "Save",
        "saved": "Saved",
        "ai_settings": "AI Service Settings",
        "use_builtin_service": "Use built‑in AI service",
        "use_custom_key": "Provide your own Gemini API key",
        "api_key": "Gemini API key",
        "accept_terms": "I accept terms and understand charges may apply.",
        "ai_not_configured": "AI service is not configured. Please configure AI settings in the sidebar.",
        
        # Darcy
        "darcy_title": "💧 Darcy Law Calculator",
        "darcy_caption": "Calculate groundwater travel time and specific discharge.",
        "input_params": "Input Parameters",
        "darcy_k": "Hydraulic Conductivity (K)",
        "darcy_i": "Hydraulic Gradient (i)",
        "darcy_n": "Effective Porosity (ne)",
        "formulas": "Formulas",
        "darcy_v": "Specific Discharge (q)",
        "darcy_vs": "Seepage Velocity (v)",
        "darcy_annual": "Annual Movement",
        "travel_time_title": "Travel Time",
        "distance": "Distance (m)",
        "travel_time_res": "Water will travel this distance in",
        "days": "days",
        "years": "years",
        
        # Lithology
        "manager": "Manager",
        "lith_input": "Data Input",
        "well_id": "Well ID:",
        "manual_entry": "Manual Entry",
        "import": "Import",
        "from_m": "From (m)",
        "to_m": "To (m)",
        "lithology": "Lithology",
        "desc": "Description",
        "add_layer": "Add Layer",
        "paste_log_hint": "Paste your drilling log text below.",
        "log_text": "Log Text",
        "parse_ai": "🔮 Parse with AI",
        "analyzing": "Analyzing...",
        "preview": "Preview:",
        "apply_log": "✅ Apply to Log",
        "discard": "❌ Discard",
        "upload_csv": "Upload CSV",
        "load_csv": "Load CSV",
        "qc_edit": "QC & Edit",
        "issues_found": "Issues Detected",
        "auto_fix": "Auto-Fix",
        "log_ok": "Log Integrity: OK ✅",
        "strip_log_vis": "Strip Log Visual",
        "success": "Success!",
        
        # OCR
        "ocr_title": "📷 AI OCR & Parser",
        "ocr_info": "Analyzes ArcGIS tables, Lab reports, etc.",
        "upload_file": "Upload File",
        "ai_warning": "⚠️ Please enter API Key in sidebar.",
        "ocr_how": "How it works?",
        "ocr_steps": "1. Upload\n2. Wait\n3. Check results",
        "ocr_waiting": "Analyzing document...",
        "ocr_result": "Result Table",
        "ocr_ai_report": "AI Analysis Report",
        "report_writing": "Writing report...",
        "ocr_error_hint": "Try a clearer image.",
        "success_rows": "Success! {count} rows found.",
        
        # Forecast
        "forecast_title": "🔮 Future Trend Forecast",
        "forecast_desc": "Polynomial regression (2nd degree).",
        "forecast_caption": "Polynomial regression prediction.",
        "history": "History",
        "history_data": "Historical Data",
        "years_forward": "Years forward?",
        "result": "Result",
        "results": "Results",
        "year_forecast": "Year Forecast",
        "forecast_val": "Prediction",
        "trend_direction": "Trend Direction",
        "trend_dir": "Direction",
        "increasing": "Increasing",
        "decreasing": "Decreasing",
        "trend_up": "Increasing 📈",
        "trend_down": "Decreasing 📉",
        
        # Report
        "ai_report_caption": "Professional hydraulic review by Gemini.",
        "ai_report_info": "Analyzes all project data to write a final text.",
        "generate_report": "🚀 Generate Report",
        "ai_writing": "AI is reading and writing...",
        "report_structure": "##### Report Content:\n1. Intro\n2. Chemistry\n3. WQI\n4. Conclusion",
        
        # Maps
        "maps_title": "Geospatial Analysis (GIS)",
        "layers": "Layers",
        "kriging_info": "Kriging fills gaps between points.",
        "show_points": "Show Points",
        "opacity": "Opacity",
        "no_map_data": "No data for selected parameter.",
        "lat_lon_required": "Lat/Lon columns required.",
        "gis_tab_interpolation": "🗺️ Interpolation Map",
        "gis_tab_cv": "📉 Cross-Validation",
        "gis_tab_flow": "🌊 Flow Analysis",
        "gis_data_layer": "##### 1. Data Layer",
        "gis_method": "##### 2. Method",
        "gis_interpolation_method": "Interpolation Method",
        "gis_formatting": "##### 3. Formatting",
        "gis_colormap": "Colormap",
        "gis_contour_levels": "Contour Levels",
        "gis_constraints": "##### 4. Constraints",
        "gis_clip_hull": "Clip to Convex Hull",
        "gis_clip_hull_help": "Masks area outside data points.",
        "gis_resolution": "Resolution (Grid)",
        "gis_resolution_help": "Higher = smoother but slower.",
        "gis_variogram_settings": "Variogram Settings (Kriging)",
        "gis_variogram_model": "Variogram Model",
        "gis_manual_params": "Manual Parameters?",
        "gis_sill": "Sill",
        "gis_range": "Range",
        "gis_nugget": "Nugget",
        "gis_not_enough_points": "Not enough points for interpolation. Need at least 4. (N={n})",
        "gis_computing": "Computing {method}...",
        "gis_layer_view": "Layer View",
        "gis_layer_interpolated": "Interpolated Value",
        "gis_layer_uncertainty": "Uncertainty (Variance)",
        "gis_layer_name": "Layer",
        "gis_cv_title": "#### 📉 Leave-One-Out Cross Validation (LOOCV)",
        "gis_cv_caption": "Validates the model by removing one point at a time and predicting it.",
        "gis_cv_param": "CV Parameter",
        "gis_cv_method": "CV Method",
        "gis_run_validation": "Run Validation",
        "gis_cv_running": "Running LOOCV...",
        "gis_rmse_help": "Root Mean Square Error (Lower is better)",
        "gis_r2_help": "Coefficient of Determination (Closer to 1 is better)",
        "gis_cv_plot_title": "Predicted vs Observed ({param})",
        "gis_detailed_residuals": "Detailed Residuals",
        "gis_select_params_prompt": "Select parameters and click 'Run Validation'",
        "gis_flow_title": "#### 🌊 Groundwater Flow Direction",
        "gis_flow_info": "Estimated from Water Level gradient (requires 'Water Level' column).",
        "gis_flow_missing": "Dataset must have 'Water Level' column.",
        "gis_flow_grid_resolution": "Grid Resolution",
        "gis_not_enough_data": "Not enough data.",
        "gis_ideal_line": "Ideal (1:1)",
        "gis_intersection": "Intersection",
        "gis_intersection_name": "Intersection (Bias=0)",
        "kriging_error": "Kriging error: {error}",
        "uncertainty_plot_error": "Uncertainty plot error: {error}",
        "method_idw": "IDW",
        "method_kriging": "Kriging (Ordinary)",
        "method_spline": "Spline / Cubic",
        "method_natural": "Natural Neighbor / Linear",
        
        # WQI
        "wqi_title": "Water Quality Index (WQI)",
        "wqi_desc": "Weighted Arithmetic Method.",
        "wqi_params_ok": "All parameters available.",
        "avg_wqi": "Average WQI",
        "class_counts": "Count by Class:",
        "wqi_chart_title": "WQI per Well",
        "detailed_table": "Detailed Table",
        
        # Stats
        "stats_caption": "Distribution and relationships.",
        "what_is_this": "What is this?",
        "stats_explainer": "Descriptive stats and correlations.",
        "select_param_warning": "Select at least 1 parameter.",
        "general_stats": "General Statistics",
        "correlation": "Correlation",
        "scatter_matrix": "Scatter Matrix",
        "scatter_warning": "Scatter Matrix might be slow.",
        "corr_req": "Select at least 2 parameters.",
        
        # Cluster
        "cluster_caption": "ML Grouping.",
        "how_it_works": "How it works?",
        "cluster_explainer": "K-Means algorithm groups similar waters.",
        "parameters": "Parameters",
        "cluster_req": "Need 2+ parameters.",
        "cluster_count": "Cluster Count (K)",
        "selected_params": "Selected params",
        "cluster_details": "Details per Group",
        
        # Levels
        "water_level_req": "Water Level column required.",
        "classification_method": "Classification Method",
        "level_class_explainer": "Low/Normal/High based on quartiles.",
        "min_level": "Min Level",
        "avg_level": "Average",
        "max_level": "Max Level",
        "level_chart": "Chart",
        "static_levels": "Static Levels",
        "histogram": "Histogram",
        "freq_chart": "Frequency",
        "level_map": "Level Map",
        "interpolation": "Interpolation",
        "isoline_count": "Isoline Count",
        "assumptions": "Assumptions",
        "level_assumptions_text": "data implies depth to water.",
        "save_info": "All analysis saved to DB.",
        "project_name_lbl": "Project Name",
        "save_project_btn": "Save Project",
        "project_saved": "Project saved as",

        # Chemistry Tabs (0-3)
        "piper_info": "### ℹ️ Piper Diagram",
        "piper_desc": "Used to determine water type (facies).",
        "piper_zones": """**Zones:**
* **Ca-HCO3**: Fresh water, shallow groundwater.
* **Na-Cl**: Saline water, mixture with seawater.
* **Mixed**: Transition zone.""",
        "schoeller_caption": "Logarithmic comparison of chemical composition across wells.",
        "gibbs_title": "### 🧬 Gibbs Plot",
        "gibbs_desc": "Determines the source of dissolved ions.",
        "gibbs_zones": """* **Precipitation:** Low TDS, high Na/(Na+Ca).
* **Rock-Weathering:** Medium TDS, low Na.
* **Evaporation:** High TDS, high Na.""",
        "stiff_title": "### 🔷 Stiff Diagrams",
        "stiff_caption": "Visual 'signature' for each sample.",
        
        # Agri Tab (4)
        "agri_analysis": "Analysi",
        "c1_class": "C1 (Excellent)",
        "c2_class": "C2 (Good)",
        "c3_class": "C3 (Risky)",
        "c4_class": "C4 (Unsuitable)",
        "s1_class": "S1 (Excellent)",
        "s2_class": "S2 (Medium)",
        "s3_class": "S3 (Risky)",
        "s4_class": "S4 (Unsuitable)",
        "wilcox_title": "Wilcox Diagram (Irrigation Suitability)",
        "classification_table": "#### Classification Table",
        
        # Pumping Tab (5)
        "cooper_title": "📉 Cooper-Jacob Analysis",
        "cooper_desc": "Determination of Transmissivity (T) from Pumping Test data.",
        "time_drawdown": "##### 2. Time-Drawdown",
        "results_header": "##### 3. Results",
        "transmissivity": "Transmissivity (T)",
        "slope": "Slope",
        "error_trend": "Error: Could not calculate trend.",
        "time_drawdown_plot": "Time-Drawdown Semilog Plot",
        "excellent_zone": "Excellent",
        "piper_chart_title": "Piper Diagram",
        "schoeller_chart_title": "Schoeller Diagram",
        "gibbs_chart_title": "Gibbs Diagram",
        "stiff_chart_title": "Stiff Diagram",
        
        # Validation
        "ai_parse_warning": "Please enter data to parse.",
        "ai_error": "AI error:",
        "lith_invalid_depth": "Error: Layer {0} has invalid depth range {1}-{2}",
        "lith_overlap": "Overlap detected at {0}m",
        "lith_gap": "Gap detected between {0}m and {1}m",
        
        # Well Design & QC
        "design_title": "🏗️ Well Design & QC",
        "clay_fraction": "Clay Fraction",
        "sand_fraction": "Sand Fraction",
        "screen_rec": "Screen Recommendations",
        "blind_rec": "Blind Pipe",
        "filter_pack": "Filter Pack",
        "slot_size": "Slot Size",
        "risk_flags": "Risk Flags",
        "risk_thin": "Thin layer (<1m) at {0}m",
        "risk_unstable": "Potential unstable zone at {0}m",
        "design_report_btn": "📄 Generate Design Report",
        "design_header": "### Well Design Proposal",
        "rec_interval": "Recommended Interval",
        "grain_warn": "⚠️ Note: Use grain size analysis for precise sizing.",

        # Maps - NEW
        "maps_sidebar_title": "🗺️ Map Settings",
        "interpolation_method": "Interpolation Method",
        "method_idw": "IDW (Inverse Distance Weighting)",
        "method_kriging": "Kriging (Ordinary)",
        "method_cubic": "Spline / Cubic",
        "method_linear": "Natural Neighbor / Linear",
        "desc_idw": "Weight by distance (Fast, simple)",
        "desc_kriging": "Variogram-based (Geostatistical, Best variance)",
        "desc_cubic": "Smooth surface (May overshoot)",
        "desc_linear": "Triangulation based (Robust)",
        "params_hydro": "Hydrochemistry",
        "params_wl": "Hydraulics (Water Level)",
        "params_irr": "Irrigation",
        "params_derived": "Derived Ratios",
        "kriging_settings": "Kriging Settings (Pro)",
        "variogram_model": "Variogram Model",
        "uncertainty_toggle": "Show Uncertainty (Variance)",
        "cv_panel": "📉 Cross-Validation",
        "perform_cv": "Run Cross-Validation",
        "qc_map": "Data QC & Boundary",
        "clip_hull": "Clip to Convex Hull",
        "show_outliers": "Highlight Outliers",
        "export_map": "Export Options",
        "layer_ctrl": "Layer Control",
        "flow_dir": "Flow Direction (Gradient)",
        "contours": "Contours",
        "upload_boundary": "Upload Boundary (GeoJSON)",
        "grid_res": "Grid Resolution",
        "res_low": "Fast",
        "res_med": "Balanced",
        "res_high": "High Quality",
        "cv_results": "CV Results:",
        "worst_points": "Worst Fitting Points:",

        # New Translations
        "info_ocr": "**Photo Analysis (OCR):** Capture lab sheets and let AI convert them to tables.",
        "info_forecast": "**Forecasting:** Predict future water levels from historical data.",
        "info_lithology": "**Lithology Log:** Enter drilling text and auto-generate the chart.",
        "info_report": "**Word Report:** Download a professional report with one click.",
        "system_stats": "System Statistics",
        "month_jan": "Jan", "month_feb": "Feb", "month_mar": "Mar", "month_apr": "Apr", "month_may": "May",
        "template_btn": "📥 Template",
        "ai_data_fill_header": "##### AI & Data Fill",
        "fill_mode": "Fill Mode",
        "fill_mode_manual": "Manual / Paste",
        "fill_mode_synthetic": "Generate Synthetic (Demo)",
        "gibbs_engineer_help": "### What should this page provide? (Minimum Output)",
        "dominant_process": "#### 🔹 Dominant process (dataset summary)",
        "rock_weathering": "Rock weathering",
        "evaporation": "Evaporation",
        "precipitation": "Precipitation",
        "auto_classification": "#### 🔹 Auto-classification per sample",
        "qc_warnings": "#### ⚠️ QC Warnings",
        "qc_tds_issue": "TDS unit compatible? ❌ (Issues: {0})",
        "qc_tds_ok": "- TDS unit compatible? ✅",
        "qc_ratio_issue": "Ratio between 0–1? ❌ (Outside: {0})",
        "qc_ratio_ok": "- Ratio between 0–1? ✅",
        "qc_outlier_issue": "Outliers present? ⚠️ (Found: {0})",
        "qc_outlier_ok": "- Outliers present? ❌ (None)",
        "controls": "**Controls**",
        "select_sample_focus": "Select Sample (Click-to-focus)",
        "export": "**Export**",
        "download_png_grid": "Download PNG Grid",
        "download_pdf_report": "Download PDF (Report)",
        "map_integration": "🗺️ Map Integration",
        "focused_view": "#### Focused View: {0}",
        "stiff_map_preview": "##### 🗺️ Stiff Map Preview",
        "all_diagrams": "#### All Diagrams",
        "summary_overview": "#### 📊 Summary Overview",
        "dominant_classification": "**Dominant Classification:** {0} / {1}",
        "dominant_desc": "Most samples fall into this category. Check detailed table for outliers.",
        "recommended_actions": "#### 🛠️ Recommended Actions",
        "drainage_improvement": "Drainage improvement required.",
        "salt_tolerant": "Select salt-tolerant crops.",
        "gypsum_app": "Gypsum (CaSO4) application recommended.",
        "monitor_irrigation": "Standard irrigation monitoring.",
        "assumptions_standards": "ℹ️ Assumptions & Standards Used",
        
        # Cooper-Jacob Specific
        "select_linear_section": "Select the straight line (Cooper-Jacob) section on the plot:",
        "analysis_range": "Analysis Range (Time Range)",
        "min_3_points": "At least 3 points must be selected for the trend.",
        "day_unit": "m²/day",
        "measured_all": "Measured (All)",
        "fitted_selected": "Fitted (Selected)",
        "time_log_axis": "Time (min) [Log]",
        "drawdown_axis": "Drawdown (m)",
        "calc_error": "Calculation Error: {0}",
        
        # Batch 3
        "ocr_mode_hydro": "Hydrochemical / Lab Analysis",
        "ocr_mode_geo": "Borehole Log / Geotechnical (Beta)",
        "forecast_suffix": "(Data Trend Analysis)",
        "forecast_extended_caption": "Statistical Trend Analysis & Future Forecasting with Confidence Intervals",
    },
    "az": {
        "language": "Dil",
        "home": "🏠 Əsas Səhifə",
        "new_analysis": "🧪 Yeni Analiz",
        "archive": "🗄️ Arxiv",
        "logout": "Çıxış",
        "home_sub": "Süni İntellekt Dəstəkli Hidrogeoloji Analiz Platforması",
        "chemistry": "🧪 Kimya",
        "gis": "🗺️ GIS",
        "agriculture": "🚜 Aqrar",
        "ai": "🤖 AI",
        "what_can": "### 🚀 Nələr edə bilərsiniz?",
        "start": "👉 Yeni Analizə Başla",
        "analysis_panel": "🛠️ Analiz Paneli",
        "data_input": "Məlumat Girişi:",
        "upload_excel": "📁 Excel Yüklə",
        "manual_entry": "✍️ Manual Giriş",
        "excel_file": "Excel Faylı (.xlsx)",
        "ai_fill": "🤖 AI ilə Doldur",
        "ai_caption": "Sətirlər, CSV və ya sərbəst mətni yapışdırın. AI cədvəl quracaq.",
        "ai_input": "AI girişi",
        "random_rows": "Təsadüfi sətir",
        "random_fill": "Təsadüfi Doldur",
        "ai_fill_btn": "AI ilə Doldur",
        "active_cols_caption": "Statistika/Klaster üçün aktiv sütunları seçin (gizlədilən sütunlar avtomatik götürülmür).",
        "active_cols": "Aktiv sütunlar",
        "tabs_agri": "🌾 Aqrar",
        "tabs_pump": "📉 Nasos",
        "tabs_lith": "🪨 Litologiya",
        "tabs_darcy": "🧮 Darsi",
        "tabs_forecast": "🔮 Proqnoz",
        "tabs_report": "🤖 Hesabat",
        "tabs_save": "💾 Yaddaş",
        "tabs_stats": "📊 Statistik",
        "tabs_cluster": "🔗 Klaster",
        "tabs_level": "💧 Səviyyə",
        "no_data": "Data yoxdur",
        "na_required": "Na, Ca, Mg lazımdır.",
        "discharge": "Debit (Q)",
        "ai_input_text": "##### 🤖 AI Giriş (Text)",
        "ai_parse": "AI Çevir",
        "table": "##### 📝 Cədvəl",
        "image": "Şəkil",
        "analyze": "Analiz",
        "generate": "Analiz Yaz",
        "generating": "Yazılır...",
        "download_word": "📄 Word Yüklə",
        "parameter": "Parametr",
        "method": "Metod",
        "name": "Ad",
        "save": "Yadda Saxla",
        "saved": "Yadda saxlandı",
        "ai_settings": "AI Xidməti Ayarları",
        "use_builtin_service": "Daxili AI xidmətindən istifadə et",
        "use_custom_key": "Öz Gemini API açarınızı daxil edin",
        "api_key": "Gemini API açarı",
        "accept_terms": "Şərtləri qəbul edirəm və ödəniş ola biləcəyini anlayıram.",
        "ai_not_configured": "AI xidməti qurulmayıb. Zəhmət olmasa sağ paneldən AI ayarlarını tənzimləyin.",
        
        # Darcy
        "darcy_title": "💧 Darcy Qanunu Kalkulyatoru",
        "darcy_caption": "Yeraltı suların süzülmə sürətini hesablayın.",
        "input_params": "Giriş Parametrləri",
        "darcy_k": "Süzülmə Əmsalı (K)",
        "darcy_i": "Hidravlik Meyllik (i)",
        "darcy_n": "Effektiv Məsaməlilik (ne)",
        "formulas": "Düsturlar",
        "darcy_v": "Xüsusi Sərf (q)",
        "darcy_vs": "Süzülmə Sürəti (v)",
        "darcy_annual": "İllik Hərəkət",
        "travel_time_title": "Yayılma Müddəti",
        "distance": "Məsafə (m)",
        "travel_time_res": "Su bu məsafəni qət edəcək:",
        "days": "gün",
        "years": "il",
        
        # Lithology
        "manager": "Menecer",
        "lith_input": "Data Girişi",
        "well_id": "Quyu ID:",
        "manual_entry": "Manual",
        "import": "İmport",
        "from_m": "Hardan (m)",
        "to_m": "Hara (m)",
        "lithology": "Litologiya",
        "desc": "Təsviri",
        "add_layer": "Lay Əlavə Et",
        "paste_log_hint": "Qazma məlumatını bura yapışdırın.",
        "log_text": "Log Mətni",
        "parse_ai": "🔮 AI ilə Oxu",
        "analyzing": "Analiz edilir...",
        "preview": "Önizləmə:",
        "apply_log": "✅ Loya tətbiq et",
        "discard": "❌ Ləğv et",
        "upload_csv": "CSV Yüklə",
        "load_csv": "CSV Oxu",
        "qc_edit": "QC & Redaktə",
        "issues_found": "xəta tapıldı",
        "auto_fix": "Avto-Düzəliş",
        "log_ok": "Log Tamlığı: OK ✅",
        "strip_log_vis": "Sütun Vizualizasiyası",
        "success": "Uğurlu!",
        
        # OCR
        "ocr_title": "📷 AI OCR & Parser",
        "ocr_info": "ArcGIS cədvəllərini, Lab hesabatlarını analiz edir.",
        "upload_file": "Fayl Yüklə",
        "ai_warning": "⚠️ API açarını daxil edin.",
        "ocr_how": "Necə işləyir?",
        "ocr_steps": "1. Yüklə\n2. Gözlə\n3. Yoxla",
        "ocr_waiting": "Sənəd oxunur...",
        "ocr_result": "Nəticə Cədvəli",
        "ocr_ai_report": "AI Analiz Rəyi",
        "report_writing": "Rəy yazılır...",
        "ocr_error_hint": "Şəkil keyfiyyətini yoxlayın.",
        "success_rows": "Uğurlu! {count} sətir.",
        
        # Forecast
        "forecast_title": "🔮 Gələcək Zaman Proqnozu",
        "forecast_desc": "Polinomial reqressiya (2-ci dərəcəli).",
        "forecast_caption": "Trend proqnozu.",
        "history": "Tarixçə",
        "history_data": "Tarixi Məlumat",
        "years_forward": "Neçə il irəli?",
        "result": "Nəticə",
        "results": "Nəticələr",
        "year_forecast": "İl Proqnozu",
        "forecast_val": "Proqnoz",
        "trend_direction": "Trend İstiqaməti",
        "trend_dir": "İstiqamət",
        "increasing": "Artan",
        "decreasing": "Azalan",
        "trend_up": "Artan 📈",
        "trend_down": "Azalan 📉",
        
        # Report
        "ai_report_caption": "Google Gemini tərəfindən hidravlik rəy.",
        "ai_report_info": "Bütün layihəni analiz edib rəy yazır.",
        "generate_report": "🚀 Rəy Yaz",
        "ai_writing": "AI düşünür və yazır...",
        "report_structure": "##### Məzmun:\n1. Giriş\n2. Kimya\n3. WQI\n4. Nəticə",
        
        # Maps
        "maps_title": "Geospatial Analiz (GIS)",
        "layers": "Laylar (Layers)",
        "kriging_info": "Kriging boşluqları doldurur.",
        "show_points": "Nöqtələri Göstər",
        "opacity": "Şəffaflıq",
        "no_map_data": "Seçilmiş parametrdə məlumat yoxdur.",
        "lat_lon_required": "Lat/Lon tələb olunur.",
        "gis_tab_interpolation": "🗺️ İnterpolasiya Xəritəsi",
        "gis_tab_cv": "📉 Kross-Validasiya",
        "gis_tab_flow": "🌊 Axın Analizi",
        "gis_data_layer": "##### 1. Məlumat Qatı",
        "gis_method": "##### 2. Metod",
        "gis_interpolation_method": "İnterpolasiya Metodu",
        "gis_formatting": "##### 3. Formatlama",
        "gis_colormap": "Rəng Xəritəsi",
        "gis_contour_levels": "Kontur Səviyyələri",
        "gis_constraints": "##### 4. Məhdudiyyətlər",
        "gis_clip_hull": "Konveks Qabıqla Kəsmə",
        "gis_clip_hull_help": "Nöqtələrdən kənar sahəni maskalayır.",
        "gis_resolution": "Rezolyusiya (Şəbəkə)",
        "gis_resolution_help": "Yüksək dəyər = daha hamar, amma daha yavaş.",
        "gis_variogram_settings": "Variogram Ayarları (Kriging)",
        "gis_variogram_model": "Variogram Modeli",
        "gis_manual_params": "Əl ilə Parametrlər?",
        "gis_sill": "Sill",
        "gis_range": "Range",
        "gis_nugget": "Nugget",
        "gis_not_enough_points": "İnterpolasiya üçün kifayət qədər nöqtə yoxdur. Ən az 4 lazımdır. (N={n})",
        "gis_computing": "Hesablanır: {method}...",
        "gis_layer_view": "Lay Görünüşü",
        "gis_layer_interpolated": "İnterpolasiya Dəyəri",
        "gis_layer_uncertainty": "Qeyri-müəyyənlik (Varians)",
        "gis_layer_name": "Lay",
        "gis_cv_title": "#### 📉 Leave-One-Out Çarpaz Yoxlama (LOOCV)",
        "gis_cv_caption": "Modeli bir-bir nöqtəni çıxarıb proqnozlaşdıraraq yoxlayır.",
        "gis_cv_param": "CV Parametri",
        "gis_cv_method": "CV Metodu",
        "gis_run_validation": "Yoxlamanı İşə Sal",
        "gis_cv_running": "LOOCV işləyir...",
        "gis_rmse_help": "Root Mean Square Error (Aşağı olması yaxşıdır)",
        "gis_r2_help": "Determinasiya əmsalı (1-ə yaxın daha yaxşıdır)",
        "gis_cv_plot_title": "Proqnoz vs Müşahidə ({param})",
        "gis_detailed_residuals": "Ətraflı Qalıqlar",
        "gis_select_params_prompt": "Parametrləri seçin və 'Yoxlamanı İşə Sal' düyməsinə basın",
        "gis_flow_title": "#### 🌊 Yeraltı Su Axın İstiqaməti",
        "gis_flow_info": "Su səviyyəsi qradienti əsasında hesablanır ('Water Level' sütunu tələb olunur).",
        "gis_flow_missing": "Datasetdə 'Water Level' sütunu olmalıdır.",
        "gis_flow_grid_resolution": "Şəbəkə Rezolyusiyası",
        "gis_not_enough_data": "Məlumat kifayət deyil.",
        "gis_ideal_line": "İdeal (1:1)",
        "gis_intersection": "Kəsişmə",
        "gis_intersection_name": "Kəsişmə (Bias=0)",
        "kriging_error": "Kriqinq xətası: {error}",
        "uncertainty_plot_error": "Qeyri-müəyyənlik qrafiki xətası: {error}",
        "method_idw": "IDW",
        "method_kriging": "Kriqinqlə (Ordinary)",
        "method_spline": "Splin / Kubik",
        "method_natural": "Təbii Qonşu / Xətti",
        
        # WQI
        "wqi_title": "Su Keyfiyyəti İndeksi (WQI)",
        "wqi_desc": "Weighted Arithmetic Metodu.",
        "wqi_params_ok": "Bütün parametrlər var.",
        "avg_wqi": "Ortalama WQI",
        "class_counts": "Siniflər üzrə say:",
        "wqi_chart_title": "Quyu üzrə WQI",
        "detailed_table": "Ətraflı Cədvəl",
        
        # Stats
        "stats_caption": "Paylanma və əlaqələr.",
        "what_is_this": "Bu nədir?",
        "stats_explainer": "Təsviri statistika və korrelyasiya.",
        "select_param_warning": "Ən azı 1 parametr seçin.",
        "general_stats": "Ümumi Statistika",
        "correlation": "Korrelyasiya",
        "scatter_matrix": "Səpələnmə Matriksi",
        "scatter_warning": "Scatter Matrix ağır işləyə bilər.",
        "corr_req": "Ən azı 2 parametr seçin.",
        
        # Cluster
        "cluster_caption": "ML Qruplaşdırma.",
        "how_it_works": "Necə işləyir?",
        "cluster_explainer": "K-Means oxşar suları qruplaşdırır.",
        "parameters": "Parametrlər",
        "cluster_req": "2+ parametr lazımdır.",
        "cluster_count": "Qrup sayı (K)",
        "selected_params": "Seçildi",
        "cluster_details": "Qrup detalları",
        
        # Levels
        "water_level_req": "Water Level sütunu lazımdır.",
        "classification_method": "Təsnifat Metodu",
        "level_class_explainer": "Kvartil əsaslı High/Low bölgüsü.",
        "min_level": "Min Səviyyə",
        "avg_level": "Ortalama",
        "max_level": "Max Səviyyə",
        "level_chart": "Qrafik",
        "static_levels": "Statik Səviyyələr",
        "histogram": "Histofram",
        "freq_chart": "Tezlik",
        "level_map": "Səviyyə Xəritəsi",
        "interpolation": "İnterpolyasiya",
        "isoline_count": "İzohips Sayı",
        "assumptions": "Qeydlər",
        "level_assumptions_text": "Məlumat yer səthinə nəzərən dərinliyi bildirir.",
        "save_info": "Məlumat bazaya yazılır.",
        "project_name_lbl": "Layihə Adı",
        "save_project_btn": "Yadda Saxla",
        "project_saved": "Uğurla yazıldı:",

        # Chemistry Tabs (0-3)
        "piper_info": "### ℹ️ Piper Diaqramı",
        "piper_desc": "Su tipini (fasiyasını) təyin etmək üçün istifadə olunur.",
        "piper_zones": """**Zonalar:**
* **Ca-HCO3**: Şirin su, dayaz yeraltı sular.
* **Na-Cl**: Duzlu su, dəniz suyu qarışığı.
* **Qarışıq**: Keçid zonası.""",
        "schoeller_caption": "Müxtəlif quyuların kimyəvi tərkibinin loqarifmik müqayisəsi.",
        "gibbs_title": "### 🧬 Gibbs Plot",
        "gibbs_desc": "Suyun minerallaşma mənbəyini göstərir.",
        "gibbs_zones": """* **Yağıntı (Precipitation):** Aşağı TDS, yüksək Na/(Na+Ca).
* **Süxur (Rock-Weathering):** Orta TDS, aşağı Na rəhbərliyi.
* **Buxarlanma (Evaporation):** Yüksək TDS, yüksək Na.""",
        "stiff_title": "### 🔷 Stiff Diaqramları",
        "stiff_caption": "Hər bir nümunə üçün vizual 'imza'.",
        
        # Agri Tab (4)
        "agri_analysis": "Analizi",
        "c1_class": "C1 (Əla)",
        "c2_class": "C2 (Yaxşı)",
        "c3_class": "C3 (Riskli)",
        "c4_class": "C4 (Yararsız)",
        "s1_class": "S1 (Əla)",
        "s2_class": "S2 (Orta)",
        "s3_class": "S3 (Riskli)",
        "s4_class": "S4 (Yararsız)",
        "wilcox_title": "Wilcox Diagramı (Suvarma Yararlılığı)",
        "classification_table": "#### Təsnifat Cədvəli",
        
        # Pumping Tab (5)
        "cooper_title": "📉 Cooper-Jacob Analizi",
        "cooper_desc": "Nasos təcrübəsi (Pumping Test) məlumatlarından Transmissivliyin (T) təyini.",
        "time_drawdown": "##### 2. Zaman-Səviyyə",
        "results_header": "##### 3. Nəticələr",
        "transmissivity": "Transmissivlik (T)",
        "slope": "Meyllik (Slope)",
        "error_trend": "Xəta: Trend hesablana bilmədi.",
        "time_drawdown_plot": "Zaman-Səviyyə (Semilog)",
        "excellent_zone": "Əla",
        "piper_chart_title": "Piper Diaqramı",
        "schoeller_chart_title": "Schoeller Diaqramı",
        "gibbs_chart_title": "Gibbs Diaqramı",
        "stiff_chart_title": "Stiff Diaqramı",
        
        # Validation
        "ai_parse_warning": "Zəhmət olmasa təhlil üçün data daxil edin.",
        "ai_error": "AI xətası:",
        "lith_invalid_depth": "Dərinlik xətası: Layer {0}, {1}-{2} aralığı yanlışdır",
        "lith_overlap": "Çarpazlaşma aşkar edildi: {0}m",
        "lith_gap": "Boşluq aşkar edildi: {0}m və {1}m arasında",
        
         # Well Design & QC
        "design_title": "🏗️ Quyu Dizaynı və QC",
        "clay_fraction": "Gil Faizi",
        "sand_fraction": "Qum/Suxur Faizi",
        "screen_rec": "Filtr Tövsiyələri",
        "blind_rec": "Karoy/Kor Boru",
        "filter_pack": "Çınqıl Paket",
        "slot_size": "Darıq Gözü (Slot)",
        "risk_flags": "Risk Faktorları",
        "risk_thin": "Nazik lay (<1m): {0}m",
        "risk_unstable": "Potensial dayanıqsız zona: {0}m",
        "design_report_btn": "📄 Dizayn Rəyi Yarat",
        "design_header": "### Quyu Konstruksiya Təklifi",
        "rec_interval": "Tövsiyə Olunan İnterval",
        "grain_warn": "⚠️ Qeyd: Dəqiq ölçülər üçün qranulometrik analiz lazımdır.",

        # Maps - NEW
        "maps_sidebar_title": "🗺️ Xəritə Ayarları",
        "interpolation_method": "İnterpolyasiya Metodu",
        "method_idw": "IDW (Tərs Məsafə Çəkili)",
        "method_kriging": "Kriging (Ordinary)",
        "method_cubic": "Spline / Cubic",
        "method_linear": "Natural Neighbor / Linear",
        "desc_idw": "Məsafəyə görə çəki (Sürətli, sadə)",
        "desc_kriging": "Varioqrama əsaslanan (Geostatistik, Ən yaxşı varians)",
        "desc_cubic": "Hamar səth (Overshoot ola bilər)",
        "desc_linear": "Triangulyasiyaya əsaslanan (Stabil)",
        "params_hydro": "Hidrokimya",
        "params_wl": "Hidravlika (Su Səviyyəsi)",
        "params_irr": "Suvarma",
        "params_derived": "Törəmə Nisbətlər",
        "kriging_settings": "Kriging Ayarları (Pro)",
        "variogram_model": "Varioqram Modeli",
        "uncertainty_toggle": "Qeyri-müəyyənliyi göstər (Variance)",
        "cv_panel": "📉 Cross-Validation (Yoxlama)",
        "perform_cv": "Yoxlamanı Başlat",
        "qc_map": "Data QC & Sərhədlər",
        "clip_hull": "Convex Hull daxilində kəs",
        "show_outliers": "Outlier-ləri vurğula",
        "export_map": "Eksport Seçimləri",
        "layer_ctrl": "Lay İdarəetməsi (Layer Control)",
        "flow_dir": "Axın İstiqaməti (Qradient)",
        "contours": "İzoxətlər (Contours)",
        "upload_boundary": "Sərhəd yüklə (GeoJSON)",
        "grid_res": "Grid Keyfiyyəti",
        "res_low": "Sürətli",
        "res_med": "Balanslı",
        "res_high": "Yüksək",
        "cv_results": "CV Nəticələri:",
        "worst_points": "Ən pis uyğunlaşan nöqtələr:",

        "info_ocr": "**Foto Analizi (OCR):** Laboratoriya vərəqlərini çəkin və AI onları cədvələ çevirsin.",
        "info_forecast": "**Proqnozlaşdırma:** Tarixi məlumatlardan gələcək su səviyyələrini proqnozlaşdırın.",
        "info_lithology": "**Litologiya Logu:** Qazma mətnini daxil edin və avtomatik qrafik yaradın.",
        "info_report": "**Word Hesabatı:** Bir kliklə professional hesabat yükləyin.",
        "system_stats": "Sistem Statistikası",
        "month_jan": "Yan", "month_feb": "Fev", "month_mar": "Mar", "month_apr": "Apr", "month_may": "May",
        "template_btn": "📥 Şablon",
        "ai_data_fill_header": "##### AI və Məlumat Doldurma",
        "fill_mode": "Doldurma Rejimi",
        "fill_mode_manual": "Manual / Yapışdır",
        "fill_mode_synthetic": "Sintetik Yarad (Demo)",
        "gibbs_engineer_help": "### Mühəndis üçün bu səhifə nə verməlidir? (Minimum Output)",
        "dominant_process": "#### 🔹 Dominant proses (dataset summary)",
        "rock_weathering": "Süxur aşınması (Rock weathering)",
        "evaporation": "Buxarlanma (Evaporation)",
        "precipitation": "Yağıntı (Precipitation)",
        "auto_classification": "#### 🔹 Hər nümunə üçün avtomatik klasifikasiya",
        "qc_warnings": "#### ⚠️ QC xəbərdarlıqları",
        "qc_tds_issue": "TDS vahidi uyğundur? ❌ (Problemlər: {0})",
        "qc_tds_ok": "- TDS vahidi uyğundur? ✅",
        "qc_ratio_issue": "Ratio 0–1 arasıdır? ❌ (Xaricdə: {0})",
        "qc_ratio_ok": "- Ratio 0–1 arasıdır? ✅",
        "qc_outlier_issue": "Outlier var? ⚠️ (Tapıldı: {0})",
        "qc_outlier_ok": "- Outlier var? ❌ (Yoxdur)",
        "controls": "**İdarəetmə**",
        "select_sample_focus": "Nümunə Seç (Fokuslanmaq üçün klikləyin)",
        "export": "**Eksport**",
        "download_png_grid": "PNG Grid Yüklə",
        "download_pdf_report": "PDF Yüklə (Hesabat)",
        "map_integration": "🗺️ Xəritə İnteqrasiyası",
        "focused_view": "#### Fokuslanmış Görünüş: {0}",
        "stiff_map_preview": "##### 🗺️ Stiff Xəritə Önizləməsi",
        "all_diagrams": "#### Bütün Diaqramlar",
        "summary_overview": "#### 📊 Yekun İcmal",
        "dominant_classification": "**Dominant Təsnifat:** {0} / {1}",
        "dominant_desc": "Əksər nümunələr bu kateqoriyaya düşür. Outlierlər üçün ətraflı cədvələ baxın.",
        "recommended_actions": "#### 🛠️ Tövsiyə Olunan Tədbirlər",
        "drainage_improvement": "Drenajın yaxşılaşdırılması tələb olunur.",
        "salt_tolerant": "Duzluluğa davamlı bitkilər seçin.",
        "gypsum_app": "Gips (CaSO4) tətbiqi məsləhətdir.",
        "monitor_irrigation": "Standart suvarma monitorinqi.",
        "assumptions_standards": "ℹ️ Fərziyyələr və İstinadlar",
        
        # Cooper-Jacob Specific
        "select_linear_section": "📉 Qrafikdə düz xətt (Cooper-Jacob) hissəsini seçin:",
        "analysis_range": "Analiz Aralığı (Time Range)",
        "min_3_points": "Trend üçün ən azı 3 nöqtə seçilməlidir.",
        "day_unit": "m²/gün",
        "measured_all": "Ölçülən (All)",
        "fitted_selected": "Fit üçün (Selected)",
        "time_log_axis": "Zaman (dəq) [Log]",
        "drawdown_axis": "Səviyyə düşməsi (m)",
        "calc_error": "Hesablama xətası: {0}",

        # Batch 3
        "ocr_mode_hydro": "Hidrokimyəvi / Lab Analizi",
        "ocr_mode_geo": "Qazma Logu / Geotexniki (Beta)",
        "forecast_suffix": "(Məlumat Trend Analizi)",
        "forecast_extended_caption": "Statistik Trend Analizi və İnam İntervalları ilə Gələcək Proqnozu",
    }
}

def t(key):
    lang = st.session_state.get("lang", "en")
    return LANGS.get(lang, LANGS["en"]).get(key, key)

# Build reverse lookup for auto-translation of literal UI strings
EN_VALUE_TO_KEY = {v: k for k, v in LANGS.get("en", {}).items() if isinstance(v, str)}

def tr_text(text: str) -> str:
    """Translate literal UI strings if they match an EN value in LANGS."""
    lang = st.session_state.get("lang", "en")
    if lang == "en":
        return text
    key = EN_VALUE_TO_KEY.get(text)
    if key:
        return LANGS.get(lang, LANGS["en"]).get(key, text)
    return text

def _localize_arg0(func):
    def wrapper(*args, **kwargs):
        if args and isinstance(args[0], str):
            args = (tr_text(args[0]),) + args[1:]
        return func(*args, **kwargs)
    return wrapper

def apply_streamlit_localization():
    """Auto-translate first string argument for common Streamlit UI calls."""
    for name in [
        "title", "header", "subheader", "markdown", "caption", "write",
        "info", "warning", "error", "success", "text", "checkbox", "radio",
        "selectbox", "multiselect", "slider", "button", "expander",
        "download_button", "text_input", "number_input", "metric"
    ]:
        if hasattr(st, name):
            setattr(st, name, _localize_arg0(getattr(st, name)))

apply_streamlit_localization()

st.set_page_config(page_title="HydroMind Pro", page_icon="💧", layout="wide")

# --- CSS (UI FIX & DESIGN) ---
st.markdown("""
<style>
    :root {
        color-scheme: light dark;
        --card-bg: #ffffff;
        --card-text: #111111;
        --card-subtext: #444444;
    }
    .block-container {
        padding-top: 4rem !important; 
        padding-bottom: 2rem;
    }
    .stTabs [data-baseweb="tab-list"] {
        gap: 6px;
        background-color: transparent;
        padding-bottom: 10px;
        flex-wrap: wrap;
    }
    .stTabs [data-baseweb="tab"] {
        height: 40px;
        white-space: nowrap !important;
        background-color: #f0f2f6;
        border-radius: 20px;
        color: #444;
        font-weight: 600;
        font-size: 13px !important;
        border: 1px solid #ddd;
        padding: 0 15px !important;
        flex-grow: 1;
        text-align: center;
        display: flex;
        align-items: center;
        justify-content: center;
    }
    .stTabs [aria-selected="true"] {
        background-color: #0079c1 !important;
        color: white !important;
        border: none;
        box-shadow: 0 3px 6px rgba(0, 121, 193, 0.3);
    }
    .stTabs [data-baseweb="tab"]:hover {
        background-color: #e3e6ea;
        color: #0079c1;
        transform: translateY(-2px);
    }
    .stTabs [data-baseweb="tab-border"] { display: none; }
    .feature-card {
        background-color: var(--card-bg);
        color: var(--card-text);
        padding: 20px;
        border-radius: 10px;
        border: 1px solid #e0e0e0;
        box-shadow: 0 2px 5px rgba(0,0,0,0.05);
        text-align: center;
        height: 100%;
    }
    .feature-card h3 { color: var(--card-text); }
    .feature-card p { color: var(--card-subtext); }
    .big-header { font-size: 3rem; font-weight: 800; color: #0079c1; text-align: center; margin-bottom: 0;}
    .sub-header { font-size: 1.2rem; color: #666; text-align: center; margin-bottom: 2rem;}

    /* Dark mode fixes for readability */
    @media (prefers-color-scheme: dark) {
        :root {
            --card-bg: #1b1f27;
            --card-text: #e6e6e6;
            --card-subtext: #bfc7d5;
        }
        .block-container {
            color: #e6e6e6;
        }
        .stTabs [data-baseweb="tab"] {
            background-color: #1f232b;
            color: #e6e6e6;
            border: 1px solid #2f3440;
        }
        .stTabs [data-baseweb="tab"]:hover {
            background-color: #2a2f3a;
            color: #7cc4ff;
        }
        .stTabs [aria-selected="true"] {
            background-color: #0079c1 !important;
            color: #ffffff !important;
            border: none;
            box-shadow: 0 3px 6px rgba(0, 121, 193, 0.45);
        }
        .feature-card {
            border: 1px solid #2f3440;
            box-shadow: 0 2px 8px rgba(0,0,0,0.35);
        }
        .sub-header {
            color: #bfc7d5;
        }
        .big-header {
            color: #7cc4ff;
        }
        .stMarkdown, .stMarkdown p, .stMarkdown li, .stText, .stTextArea, .stCaption {
            color: #e6e6e6;
        }
    }

    /* Streamlit dark theme override (when prefers-color-scheme is not respected) */
    .stApp[data-theme="dark"] {
        --card-bg: #1b1f27;
        --card-text: #e6e6e6;
        --card-subtext: #bfc7d5;
    }
    .stApp[data-theme="dark"] .block-container {
        color: #e6e6e6;
    }
    .stApp[data-theme="dark"] .stTabs [data-baseweb="tab"] {
        background-color: #1f232b;
        color: #e6e6e6;
        border: 1px solid #2f3440;
    }
    .stApp[data-theme="dark"] .stTabs [data-baseweb="tab"]:hover {
        background-color: #2a2f3a;
        color: #7cc4ff;
    }
    .stApp[data-theme="dark"] .stTabs [aria-selected="true"] {
        background-color: #0079c1 !important;
        color: #ffffff !important;
        border: none;
        box-shadow: 0 3px 6px rgba(0, 121, 193, 0.45);
    }
    .stApp[data-theme="dark"] .feature-card {
        border: 1px solid #2f3440 !important;
        box-shadow: 0 2px 8px rgba(0,0,0,0.35) !important;
    }
    .stApp[data-theme="dark"] .sub-header {
        color: #bfc7d5;
    }
    .stApp[data-theme="dark"] .big-header {
        color: #7cc4ff;
    }
    .stApp[data-theme="dark"] .stMarkdown,
    .stApp[data-theme="dark"] .stMarkdown p,
    .stApp[data-theme="dark"] .stMarkdown li,
    .stApp[data-theme="dark"] .stText,
    .stApp[data-theme="dark"] .stTextArea,
    .stApp[data-theme="dark"] .stCaption {
        color: #e6e6e6;
    }
</style>
""", unsafe_allow_html=True)

# --- 2. DATABASE ---
if 'logged_in' not in st.session_state: st.session_state['logged_in'] = False
if 'username' not in st.session_state: st.session_state['username'] = "Guest"

conn = sqlite3.connect('hydromind_db.db', check_same_thread=False)
c = conn.cursor()
c.execute('CREATE TABLE IF NOT EXISTS userstable(username TEXT PRIMARY KEY, password TEXT)')
c.execute('CREATE TABLE IF NOT EXISTS projects(id INTEGER PRIMARY KEY AUTOINCREMENT, username TEXT, date TEXT, filename TEXT, summary TEXT, data BLOB)')
conn.commit()

def make_hashes(password): return hashlib.sha256(str.encode(password)).hexdigest()
def login_user(username, password):
    c.execute('SELECT * FROM userstable WHERE username =? AND password = ?', (username, password))
    return c.fetchall()
def add_user(username, password): 
    c.execute('INSERT INTO userstable(username,password) VALUES (?,?)', (username, password))
    conn.commit()
def save_project(username, filename, summary, df):
    buffer = io.BytesIO(); df.to_excel(buffer, index=False); blob = buffer.getvalue()
    date = datetime.now().strftime("%Y-%m-%d %H:%M")
    c.execute('INSERT INTO projects(username, date, filename, summary, data) VALUES (?,?,?,?,?)', (username, date, filename, summary, blob)); conn.commit()
def get_user_projects(username):
    c.execute('SELECT date, filename, summary, data FROM projects WHERE username =? ORDER BY id DESC', (username,)); return c.fetchall()

# --- 3. HELPER FUNCTIONS ---
def force_rerun(): st.rerun()

def gps_duzelt(deyer, tip="lat"):
    """Normalize latitude/longitude values by ensuring they fall within ±90/±180.
    Strings containing commas are converted to floats. Values that exceed the
    allowable range are iteratively divided by 10 until they fit."""
    try: 
        if isinstance(deyer, str): deyer = deyer.replace(',', '.')
        val = float(deyer)
        limit = 90 if tip == "lat" else 180
    except: return None 
    while abs(val) > limit: val = val / 10
    return val

def clean_chemistry_data(df, cols):
    """Convert selected columns to numeric and replace non‑numeric values with zero."""
    for col in cols:
        if col in df.columns:
            if df[col].dtype == 'object': df[col] = df[col].astype(str).str.replace(',', '.')
            df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0)
    return df

# --- NEW FEATURE: WATER QUALITY INDEX ---
def calculate_wqi(df, standards=None):
    """Compute the weighted arithmetic Water Quality Index (WQI).

    Parameters
    ----------
    df : pandas.DataFrame
        Data frame containing water quality measurements. Columns should include
        the parameters defined in `standards`.
    standards : dict, optional
        A dictionary mapping parameter names to their permissible limit (Si).
        If not provided, default values based on common drinking water
        guidelines are used.

    Returns
    -------
    wqi_df : pandas.DataFrame
        A new dataframe with the computed sub‑indices, weights, WQI and
        classification for each sample.
    weights : dict
        The relative unit weights used for each parameter.
    used_standards : dict
        The actual standards user for calculation.
    """
    if standards is None:
        standards = {
            # permissible values (Si) approximate WHO/BIS guidelines
            'Ca': 75,      # mg/L (Calcium)
            'Mg': 50,      # mg/L (Magnesium)
            'Na': 200,     # mg/L (Sodium)
            'HCO3': 500,   # mg/L (Bicarbonate)
            'Cl': 250,     # mg/L (Chloride)
            'SO4': 200,    # mg/L (Sulfate)
            'TDS': 500,    # mg/L (Total dissolved solids)
            'NO3': 45,     # mg/L (Nitrate)
            'pH': 8.5      # Standard pH upper limit
        }
    
    # Filter for columns present in the dataframe
    used_standards = {p: v for p, v in standards.items() if p in df.columns}
    
    if not used_standards:
        # Return empty with correct columns to avoid errors
        return pd.DataFrame({'WQI': [], 'Class': []}), {}, {}

    # constant k for weight calculation
    inv_sum = sum(1.0/s for s in used_standards.values() if s > 0)
    k = 1.0 / inv_sum if inv_sum > 0 else 1.0
    weights = {p: k / Si for p, Si in used_standards.items() if Si > 0}

    result_records = []
    for _, row in df.iterrows():
        sub_indices = {}
        weighted_sub_indices = {}
        
        # compute quality rating qi for each parameter (measured Ci vs standard Si)
        for p, Si in used_standards.items():
            Ci = row.get(p, 0)
            if pd.isna(Ci): Ci = 0
            
            # For pH, ideally we want deviation from 7.0, but keeping standard formulation for now
            # wqi = Σ(Wi * qi) / ΣWi
            # qi = (Ci / Si) * 100
            
            if p == 'pH':
                 # Standard WQI often uses ideal=7.0 for pH
                 # qi = |Ci - 7.0| / (8.5 - 7.0) * 100
                 qi = (abs(Ci - 7.0) / (8.5 - 7.0)) * 100
            else:
                 qi = (Ci / Si) * 100 if Si > 0 else 0
                 
            sub_indices[p] = qi
            weighted_sub_indices[p] = weights[p] * qi
            
        # compute weighted sum
        numerator = sum(weighted_sub_indices.values())
        denominator = sum(weights.values())
        wqi_val = numerator / denominator if denominator != 0 else 0
        
        # classify according to Brown et al. (1972) ranges
        if wqi_val <= 25:
            cls = "Excellent"
        elif wqi_val <= 50:
            cls = "Good"
        elif wqi_val <= 70:
            cls = "Poor"
        elif wqi_val <= 90:
            cls = "Very Poor"
        else:
            cls = "Unsuitable" # >90 is typically Unsuitable/Very Poor
        
        record = {'WQI': wqi_val, 'Class': cls}
        for p in used_standards:
            record[f'{p}_qi'] = sub_indices[p]
            # Contribution % = (Wi * qi) / (Sum Wi * qi)  <-- Contribution to the Numerator Sum
            # Or contribution to the final Index?
            # Metric: "Driver" implies what pushed the WQI up.
            # (Wi * qi) is the term in the numerator.
            if numerator > 0:
                record[f'{p}_contribution_pct'] = (weighted_sub_indices[p] / numerator) * 100
            else:
                record[f'{p}_contribution_pct'] = 0.0

        result_records.append(record)
        
    wqi_df = pd.DataFrame(result_records)
    return wqi_df, weights, used_standards

# --- 4. GIS ENGINE ---
def generate_contours(df, parameter, levels, colormap_name, method='Cubic'):
    try:
        df_map = df.dropna(subset=['Lat', 'Lon', parameter])
        if len(df_map) < 3: return None, None, None
        x = df_map['Lon'].values; y = df_map['Lat'].values; z = df_map[parameter].values
        xi = np.linspace(x.min(), x.max(), 100); yi = np.linspace(y.min(), y.max(), 100); xi, yi = np.meshgrid(xi, yi)
        if method == 'Kriging (RBF)':
            try:
                rbf = Rbf(x, y, z, function='linear') 
                zi = rbf(xi, yi)
                zi = np.clip(zi, z.min(), z.max())
            except: zi = griddata((x, y), z, (xi, yi), method='cubic')
        else:
            zi = griddata((x, y), z, (xi, yi), method='cubic')
            mask = np.isnan(zi)
            if mask.any(): zi[mask] = griddata((x, y), z, (xi[mask], yi[mask]), method='nearest')
        fig, ax = plt.subplots()
        contour = ax.contour(xi, yi, zi, levels=levels, linewidths=1, colors='black')
        contour_filled = ax.contourf(xi, yi, zi, levels=levels, cmap=colormap_name, alpha=0.6)
        geojson_lines = json.loads(geojsoncontour.contour_to_geojson(contour=contour, ndigits=5))
        geojson_filled = json.loads(geojsoncontour.contourf_to_geojson(contourf=contour_filled, ndigits=5))
        plt.close(fig)
        try:
            cmap = plt.get_cmap(colormap_name)
            colors = [cmap(i) for i in np.linspace(0, 1, levels)]
            hex_colors = ['#{:02x}{:02x}{:02x}'.format(int(c[0]*255), int(c[1]*255), int(c[2]*255)) for c in colors]
            colormap = cm.LinearColormap(colors=hex_colors, vmin=z.min(), vmax=z.max()).to_step(n=levels)
            colormap.caption = f"{parameter} (mg/l) - {method}"
        except: colormap = None
        return geojson_lines, geojson_filled, colormap
    except: return None, None, None

def get_marker_color(val, min_v, max_v, cmap_name='jet'):
    try:
        norm_val = (val - min_v) / (max_v - min_v) if max_v > min_v else 0.5
        cmap = plt.get_cmap(cmap_name); rgba = cmap(norm_val)
        return '#{:02x}{:02x}{:02x}'.format(int(rgba[0]*255), int(rgba[1]*255), int(rgba[2]*255))
    except: return '#000000'

# --- 6. AI ANALYSIS ---

def prepare_report_context(df, project_name="General Project", client_name="Valued Client"):
    """
    Generates a structured context string for the AI report.
    Includes Metadata, QC checks, Exceedance, and Water Facies hints.
    """
    try:
        # 1. Metadata
        sample_count = len(df)
        params = [c for c in df.columns if c not in ['geometry', 'Lat', 'Lon', 'WKT']]
        
        # 2. QC Summary (CBE)
        # We try to calculate CBE if ions are present
        qc_summary = "QC Analysis skipped (missing ions)."
        cbe_reliability = "Unknown"
        
        req_ions = ['Ca', 'Mg', 'Na', 'HCO3', 'Cl', 'SO4']
        if all(col in df.columns for col in req_ions):
            try:
                # Use local calculation to ensure availability or use simple logic here
                # Re-implementing simplified CBE here to avoid dependency ordering issues if calculate_cbe is below
                def quick_cbe(row):
                    cat = (row['Ca']/20.04) + (row['Mg']/12.15) + (row['Na']/22.99) + (row.get('K',0)/39.10)
                    an_ = (row['HCO3']/61.02) + (row['Cl']/35.45) + (row['SO4']/48.06) + (row.get('NO3',0)/62.0)
                    return ((cat - an_) / (cat + an_) * 100) if (cat+an_) > 0 else 0
                
                cbe_vals = df.apply(quick_cbe, axis=1).abs()
                avg_cbe = cbe_vals.mean()
                outliers = (cbe_vals > 10).sum()
                cbe_reliability = "High" if avg_cbe < 5 else "Moderate" if avg_cbe < 10 else "Low"
                qc_summary = f"Average CBE: {avg_cbe:.2f}%. Outliers (>10%): {outliers} samples. Data Reliability: {cbe_reliability}."
            except Exception as e:
                qc_summary = f"QC Error: {str(e)}"

        # 3. Exceedance Check (WHO/General Standards)
        limits = {'TDS': 1000, 'NO3': 50, 'Cl': 250, 'SO4': 250, 'Na': 200, 'F': 1.5, 'Fe': 0.3}
        exceed_text = []
        for p, lim in limits.items():
            if p in df.columns:
                cnt = df[df[p] > lim].shape[0]
                if cnt > 0:
                    exceed_text.append(f"{p} > {lim} ({cnt} samples)")
        
        exceed_str = ", ".join(exceed_text) if exceed_text else "None detected."

        # 4. Water Facies (Simplified)
        # Check dominant Cation and Anion based on means
        facies = "Mixed"
        if all(col in df.columns for col in ['Na', 'Ca', 'Mg', 'Cl', 'HCO3', 'SO4']):
            means = df.mean()
            dom_cat = max(['Na', 'Ca', 'Mg'], key=lambda x: means.get(x,0))
            dom_an = max(['Cl', 'HCO3', 'SO4'], key=lambda x: means.get(x,0))
            facies = f"{dom_cat}-{dom_an} Type (based on avg concentrations)"

        context = f"""
        PROJECT HEADER:
        - Project: {project_name}
        - Client: {client_name}
        - Samples: {sample_count}
        - Date: {datetime.now().strftime('%Y-%m-%d')}
        
        QUALITY CONTROL:
        - {qc_summary}
        
        HYDROCHEMISTRY HIGHLIGHTS:
        - Dominant Facies: {facies}
        - Exceedances (WHO Standards): {exceed_str}
        
        STATISTICAL SUMMARY (Averages & Ranges):
        {df.describe().to_string()}
        """
        return context
    except Exception as e:
        return f"Error preparing context: {e}\nOriginal Stats: {df.describe().to_string()}"

def ai_analiz_et(stats_context):
    try:
        # Configure Gemini
        client = get_genai_client()
        if not client:
            st.warning(t("ai_not_configured"))
            return t("ai_not_configured")
        
        # Detect language
        lang_code = st.session_state.get("lang", "en")
        lang_name = "Azerbaijani" if lang_code == "az" else "English"
        
        prompt = f"""
        Act as a Senior Hydrogeologist and Consultant. 
        Write a professional, "Audit-Ready" Hydrogeological Report in {lang_name} based on the data below.

        INPUT DATA & CONTEXT:
        {stats_context}
        
        REQUIRED REPORT STRUCTURE:
        1. **Executive Summary**: Brief overview of the project and key findings.
        2. **Introduction & Methodology**: 
           - Mention sample count and parameters analyzed.
           - State that standard methods were assumed.
           - Mention limitations (e.g., analysis based on provided major ions).
        3. **Data Quality Assurance (QA/QC)**:
           - Discuss the Charge Balance Error (CBE) results provided in the context.
           - Statement on data reliability.
        4. **Hydrochemical Characterization**:
           - Interpret the water type/facies.
           - Discuss the dominant ions.
        5. **Regulatory Compliance & Water Quality**:
           - Compare results against WHO standards (TDS, NO3, etc.).
           - Highlight any exceedances mentioned in the context.
        6. **Suitability Assessment**:
           - Suitability for Irrigation (general comment based on Salinity/Sodium if evident).
           - Suitability for Drinking/Industrial use.
        7. **Conclusion & Recommendations**:
           - Clear, actionable advice for the client.
           
        TONE: Professional, Technical, Objective.
        FORMAT: Markdown. Use tables if necessary to summarize key stats.
        Style: Formal, Scientific, yet accessible. Use proper formatting (headers, bolds).
        """
        
        response = call_gemini_with_retry(
            client,
            model="gemini-3-flash-preview",
            contents=prompt
        )
        return response.text
    except Exception as e: return f"AI Error: {str(e)}"

def ai_parse_geotech_log(file_bytes, mime_type):
    try:
        client = get_genai_client()
        if not client: return {"error": "AI not configured"}
        
        prompt = """
        Analyze this image as a Geotechnical Borehole Log / Well Log.
        Act as a Senior Geotechnical Engineer.
        
        Your task is to extract structured engineering data with HIGH PRECISION.
        
        **EXTRACT THE FOLLOWING SECTIONS:**
        
        1. **Metadata (Header)**:
           - Project Name, Well ID / Borehole No.
           - Drill Rig, Drilling Diameter (mm).
           - Location (Coordinates/Name).
           - Groundwater Level (GWL), Date, Elevation.
           - *Confidence*: Give a score (0-100) for header OCR reliability.
           
        2. **Stratigraphy (Soil Layers)**:
           - Iterate through every depth interval.
           - format: { "from": float, "to": float, "desc": "...", "lithology": "...", "color": "...", "consistency": "...", "uscs": "..." }
           - "lithology": Standardize to [Clay, Silt, Sand, Gravel, Cobbles, Limestone, Bedrock, Fill, Topsoil].
           - "uscs": Estimate USCS symbol (e.g., CL, CH, SM, GW) based on description.
           
        3. **In-situ Tests (SPT & Others)**:
           - Extract Standard Penetration Test (SPT) data.
           - format: { "depth": float, "N1": int, "N2": int, "N3": int, "N_total": int, "type": "SPT" }
           - If Pressuremeter or other tests exist, extract them too.
           
        **OUTPUT FORMAT:**
        Return ONLY a valid JSON object. Do not include markdown formatting like ```json.
        Structure:
        {
          "metadata": { "project": "...", "well_id": "...", "date": "...", "rig": "...", "diameter": "...", "gw_level": "...", "elevation": "...", "confidence": 85 },
          "stratigraphy": [{ "from": 0.0, "to": 1.5, "desc": "...", "lithology": "Clay", "uscs": "CL", "color": "Brown", "consistency": "Stiff" }],
          "tests": [{ "depth": 3.0, "N_total": 28, "type": "SPT" }],
          "exports": {
             "lithology_csv_rows": [{ "From": 0.0, "To": 1.5, "Rock": "Clay" }],
             "spt_csv_rows": [{ "From": 3.0, "To": 3.45, "N": 28 }]
          },
          "qc": { "issues": ["List any unreadable parts"], "ocr_quality": "High/Medium/Low" }
        }
        """
        
        response = call_gemini_with_retry(
            client,
            model="gemini-3-flash-preview",
            contents=[prompt, types.Part.from_bytes(data=file_bytes, mime_type=mime_type)]
        )
        
        txt = response.text
        # Clean JSON
        if "```json" in txt:
             txt = txt.split("```json")[1]
        if "```" in txt:
             txt = txt.split("```")[0]
             
        return json.loads(txt.strip())
    
    except Exception as e:
        return {"error": str(e)}

def render_geotech_visual_log(layers_df, tests_df, meta):
    """
    Draws a professional Borehole Log strip using Plotly (similar to Lithology tab).
    Includes Stratigraphy column and SPT N-value plot.
    """
    from plotly.subplots import make_subplots
    
    # Calculate ranges
    max_d = 20
    if not layers_df.empty and 'to' in layers_df.columns:
        max_d = layers_df['to'].max()
    if not tests_df.empty and 'depth' in tests_df.columns:
        valid_depths = pd.to_numeric(tests_df['depth'], errors='coerce')
        if not valid_depths.empty:
             val = valid_depths.max()
             if pd.notna(val) and val > max_d: max_d = val + 2

    fig = make_subplots(rows=1, cols=2, shared_yaxes=True, column_widths=[0.3, 0.7], 
                        subplot_titles=["Lithology", "SPT (N-Value)"])
    
    # 1. Stratigraphy Column
    if not layers_df.empty:
        for _, row in layers_df.iterrows():
            lith = row.get('lithology', 'Unknown')
            # Fallback color map
            c_map = {
                'Clay': '#8B4513', 'Silt': '#A0522D', 'Sand': '#F4A460', 'Gravel': '#808080', 
                'Rock': '#696969', 'Fill': '#D3D3D3', 'Topsoil': '#228B22', 'Limestone': '#C0C0C0'
            }
            color = c_map.get(lith, '#D2B48C')
            
            try:
                base = float(row['from'])
                th = float(row['to']) - base
                
                fig.add_trace(go.Bar(
                    x=[1], y=[th], base=[base],
                    orientation='v', name=lith,
                    marker=dict(color=color, line=dict(width=1, color='black')),
                    text=f"{lith}<br>{row.get('uscs','-')}", hoverinfo="text",
                    hovertext=f"<b>{lith}</b><br>{row.get('desc','')}<br>USCS: {row.get('uscs','')}"
                ), row=1, col=1)
                
                # Text inside bar
                fig.add_annotation(
                    x=0.5, y=base + th/2,
                    text=str(row.get('uscs', lith[:2])),
                    showarrow=False, font=dict(size=10, color='white'),
                    xref="x1", yref="y1"
                )
            except: pass

    # 2. SPT Plot
    if not tests_df.empty and 'N_total' in tests_df.columns:
        # Filter valid N
        t_df = tests_df.copy()
        t_df['N_total'] = pd.to_numeric(t_df['N_total'], errors='coerce')
        t_df['depth'] = pd.to_numeric(t_df['depth'], errors='coerce')
        t_df = t_df.dropna(subset=['N_total', 'depth'])
        
        if not t_df.empty:
            t_df = t_df.sort_values('depth')
            fig.add_trace(go.Scatter(
                x=t_df['N_total'], y=t_df['depth'], mode='lines+markers',
                name='SPT N', line=dict(color='blue', width=2),
                marker=dict(size=8, color='red')
            ), row=1, col=2)
            
            # Add Consistency Lines approx (Terzaghi & Peck)
            fig.add_vline(x=4, line_width=1, line_dash="dash", line_color="green", opacity=0.5, row=1, col=2, annotation_text="Loose")
            fig.add_vline(x=10, line_width=1, line_dash="dash", line_color="orange", opacity=0.5, row=1, col=2, annotation_text="Medium")
            fig.add_vline(x=30, line_width=1, line_dash="dash", line_color="red", opacity=0.5, row=1, col=2, annotation_text="Dense")
            fig.add_vline(x=50, line_width=1, line_dash="dash", line_color="purple", opacity=0.5, row=1, col=2, annotation_text="V.Dense")

    # Layout
    title_txt = f"Borehole Log: {meta.get('well_id','?')} | Proj: {meta.get('project','?')}"
    fig.update_layout(
        title=title_txt,
        yaxis=dict(autorange="reversed", title="Depth (m)", range=[0, max_d]),
        xaxis1=dict(showticklabels=False, title="Stratigraphy"),
        xaxis2=dict(title="SPT N-Value", range=[0, 60]),
        height=700,
        showlegend=False,
        plot_bgcolor='white',
        template='plotly_white',
        margin=dict(l=20, r=20, t=50, b=20)
    )
    return fig

def create_geotech_report_word(meta, layers_df, tests_df, interpretation_text):
    doc = Document()
    doc.add_heading('Geotechnical Borehole Log Analysis Report', 0)
    
    # Meta
    p = doc.add_paragraph()
    p.add_run(f"Project: {meta.get('project','-')}\n").bold = True
    p.add_run(f"Well ID: {meta.get('well_id','-')}\n")
    p.add_run(f"Date: {meta.get('date','-')}\n")
    p.add_run(f"Equipment: {meta.get('rig','-')}")

    doc.add_heading('1. Observed Stratigraphy', level=1)
    if not layers_df.empty:
        t = doc.add_table(rows=1, cols=4)
        t.style = 'Table Grid'
        cols = ['from', 'to', 'lithology', 'desc']
        hdr = t.rows[0].cells
        hdr[0].text = "From (m)"; hdr[1].text = "To (m)"; hdr[2].text = "Lithology"; hdr[3].text = "Description"
        
        for _, row in layers_df.iterrows():
            cells = t.add_row().cells
            cells[0].text = str(row.get('from',''))
            cells[1].text = str(row.get('to',''))
            cells[2].text = str(row.get('lithology',''))
            cells[3].text = str(row.get('desc',''))

    doc.add_heading('2. In-situ Tests (SPT)', level=1)
    if not tests_df.empty:
        t2 = doc.add_table(rows=1, cols=3)
        t2.style = 'Table Grid'
        cols2 = ['depth', 'N_total', 'type']
        hdr2 = t2.rows[0].cells
        hdr2[0].text = "Depth (m)"; hdr2[1].text = "N Value"; hdr2[2].text = "Type"
        for _, row in tests_df.iterrows():
            cells = t2.add_row().cells
            cells[0].text = str(row.get('depth',''))
            cells[1].text = str(row.get('N_total',''))
            cells[2].text = str(row.get('type',''))

    doc.add_heading('3. Engineering Interpretation', level=1)
    doc.add_paragraph(interpretation_text)
    
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def ai_generate_ocr_detailed_report(file_bytes, mime_type, data_summary, lang_code="en"):
    """
    Generates a detailed report by analyzing the image/PDF visually AND the extracted data.
    """
    try:
        client = get_genai_client()
        if not client: return t("ai_not_configured")

        lang_name = "Azerbaijani" if lang_code == "az" else "English"

        prompt = f"""
        Act as a Senior Hydrogeological Consultant.
        Perform a **Multi-Modal Analysis** of the provided document (Image/PDF) and the extracted data summary.
        
        **CONTEXT:**
        The user has uploaded a document (likely a Lab Report, Borehole Log, or GIS Attribute Table).
        We have extracted some data into a table format (summary below), but the document might contain more context (headers, footnotes, handwritten remarks, map context, etc.).
        
        **EXTRACTED DATA SUMMARY (for reference):**
        {{data_summary}}
        
        **TASK:**
        Write a detailed **Assessment Report** in {{lang_name}}.
        
        **REQUIRED CONTENT:**
        1. **Document Identification**:
           - What type of document is this? (e.g., "Standard Lab Analysis Sheet", "Handwritten Field Note", "ArcGIS Export").
           - Extract any visible metadata: Project Name, Well IDs, Dates, Location names (if visible but not in extracted table).
           
        2. **Visual QA/QC & Integrity Check**:
           - Does the extracted data summary match the visual document? 
           - Are there any flagged values, stamps (e.g., "Preliminary", "Void"), or signatures visible?
           - Note any unreadable or ambiguous sections.
           
        3. **Detailed Hydro-Chemical / Technical Interpretation**:
           - Analyze the values (pH, TDS, Ions). 
           - Determine the **Water Facies/Type** (e.g., Na-Cl, Ca-HCO3) if major ions are present.
           - Assessment of **Suitability** (Drinking, Irrigation, Concrete Aggressiveness) based on generic standards (WHO, FAO) if applicable.
           
        4. **Geospatial & Contextual Insights** (if map/coords visible):
           - If a map or coordinates are visible, comment on the location context (e.g., "Located near coast", "Inland").
           
        5. **Conclusion & Recommendations**:
           - Summarize the key findings.
           - Recommend further actions (e.g., "Resample well X", "Complete full ionic balance").
           
        **STYLE:**
        - Professional, Engineering standard.
        - Use Bullet points and Tables for readability.
        - Return strictly Markdown format.
        """
        
        response = call_gemini_with_retry(
            client,
            model="gemini-3-flash-preview",
            contents=[
                prompt,
                types.Part.from_bytes(data=file_bytes, mime_type=mime_type)
            ]
        )
        return response.text
    except Exception as e:
        return f"Error generating detailed report: {{str(e)}}"

# --- REPORT V2 PROMPTS ---
REP_SYS_PROMPT = """You are HydroMind Pro Report Engine. Produce a technical hydrogeology report from the provided dataset ONLY.
Hard rules:
1) Do NOT invent values, limits, standards, or parameters not present in the input. If missing, write “Not provided”.
2) Always state data limitations (sample size, missing parameters, units).
3) Always run QC-first: units check + missing values + basic plausibility + (if major ions provided) charge balance error (CBE%). If you cannot compute CBE due to missing ions/units, state it.
4) Use consistent units. If conversion is applied (mg/L→meq/L), show conversion basis (equivalent weights) in an appendix table or note.
5) All interpretations must be traceable to computed results (tables/figures). No generic filler.
6) Output must be in Azerbaijani, professional tone, concise, engineering style.

Input:
- Project meta: {{project_meta}}
- Coordinates: {{coordinates}}
- Dataset table: {{dataset_table}}
- Units per column: {{units_table}}
- Available analyses/charts already computed by the app: {{available_outputs}}
- Standards/threshold set selected by user (if any): {{selected_standard_set}}
"""

REP_TEMPLATES = {
    "A": """TASK: Generate a Hydrochem & Water Quality report.

Scope (include ONLY these):
1) Executive Summary (5–8 lines): water type(s), mineralization (TDS/EC), key risks, confidence level.
2) Data & Methods: sample count (n), date range (if provided), lab parameters present, units, QC workflow.
3) Quality Control:
   - Missing values summary
   - Unit consistency check
   - Charge Balance Error (CBE%) per sample IF major ions available: Ca, Mg, Na, K, HCO3/CO3, Cl, SO4 (mg/L). Convert to meq/L and compute CBE%.
   - Flag samples with |CBE| > 10% as “analytical concern”.
4) Hydrochemical Characterization:
   - Convert major ions to meq/L and %meq (if possible).
   - Identify dominant cations/anions and facies.
   - Use Piper (mandatory if available). If Piper output not available, infer facies from %meq and state “inferred”.
   - Use Schoeller OR Stiff (choose whichever is available; if both, use both only if n<=10).
   - Use Gibbs only if the app provides the plotted ratios; otherwise do not mention it.
5) Water Quality (intended use: drinking suitability as “screening” only):
   - Evaluate only parameters present.
   - If microbial/trace metals are missing, explicitly state potability cannot be confirmed; only salinity/major-ion screening performed.
   - If a standards set is provided ({{selected_standard_set}}), compare to those limits; otherwise do not cite numeric limits.
6) Spatial Notes (only if coordinates exist):
   - Map summary: any spatial gradient in TDS/Cl/Na (qualitative + computed min/max by area if possible).
7) Conclusions & Recommendations:
   - Monitoring items
   - Additional analyses recommended (microbiology, nitrate, trace metals) if missing
   - If salinization indicators (Na/Cl, TDS/EC) are high, note possible drivers (marine influence, evaporation, anthropogenic) as hypotheses, not facts.

Required Figures (embed references only; do not generate images):
- Fig 1: Piper diagram (if available)
- Fig 2: Schoeller or Stiff
- Fig 3: Gibbs plot (only if available)
- Fig 4: Map of wells colored by TDS/EC (only if available)

Required Tables:
- Table 1: Summary stats (min/median/max) for present parameters
- Table 2: Major ions in mg/L and meq/L (if possible)
- Table 3: CBE% per sample (if possible)

Formatting:
- Use numbered headings and bullet recommendations.
- Any claim must reference a table/figure output name from {{available_outputs}}.
""",

    "B": """TASK: Generate an Irrigation Suitability report.

Scope (include ONLY these):
1) Executive Summary (5–8 lines): irrigation class summary, main hazards (salinity vs sodium), key management advice.
2) Data & QC:
   - n, missing values, units
   - If EC missing but TDS present, convert ONLY if the app has a defined conversion; otherwise state EC not provided.
3) Core Irrigation Indices (compute/quote only if inputs exist):
   - SAR (mandatory if Ca, Mg, Na available in meq/L)
   - Na% (optional if cations available)
   - USSL class (C1–C4, S1–S4) if EC (or TDS proxy) and SAR exist
   - Wilcox (EC vs Na%) if both exist
   - RSC (Residual Sodium Carbonate) if HCO3/CO3 and Ca/Mg exist
4) Results:
   - Provide per-sample table: EC/TDS, SAR, Na%, class labels
   - Provide dataset summary: counts by class, worst-case wells
5) Interpretation (engineering style):
   - Separate “salinity hazard” and “sodium hazard”
   - Give actionable guidance: crop tolerance, leaching requirement, drainage, blending, monitoring soil EC
   - Do NOT claim “safe/unsafe” without thresholds; if thresholds are absent, use qualitative ranking.
6) Limitations:
   - Mention missing boron, nitrate, heavy metals, soil data; irrigation suitability depends on soil texture/drainage and crop type.

Required Figures (references only):
- Fig 1: USSL diagram (EC vs SAR) if available
- Fig 2: Wilcox diagram if available
- Fig 3: Map risk layer (optional)

Required Tables:
- Table 1: Irrigation indices per sample
- Table 2: Classification counts
- Table 3: Input ions in meq/L used for SAR/Na%

Output language: Azerbaijani.
Tone: concise, management-oriented.
""",

    "C": """TASK: Generate a Pumping Test & Groundwater Dynamics report.

Scope (include ONLY these):
1) Executive Summary (5–8 lines): test type, transmissivity/storage results, reliability, operational implications.
2) Test Description:
   - Pumping well ID, observation well ID(s) if provided
   - Pumping rate Q and units
   - Test duration, time units
   - Distance r between wells (if provided)
   - Any step test phases (if provided)
3) Data Quality & Assumptions:
   - Confirm time units consistency with T units
   - Identify early-time points vs late-time straight line
   - State assumptions for Cooper–Jacob (confined, radial flow, infinite aquifer, late-time approximation). If conditions likely violated, flag.
4) Analysis Results (use only available outputs):
   - Cooper–Jacob: slope (Δs per log cycle), T computed.
   - Compute S only if t0 and r are available; otherwise state “S not computed (missing r or t0)”.
   - If recovery data exists, include recovery analysis summary.
5) Interpretation:
   - What T implies for aquifer productivity (qualitative)
   - Expected drawdown behavior; boundary/leakage hints if the semilog deviates
6) Operational Recommendations:
   - Sustainable pumping suggestions (qualitative unless you have long-term data)
   - Monitoring plan: water level, EC/Cl (saline intrusion risk if coastal)
7) Limitations:
   - n points, fit window choice, missing r, missing well construction, partial penetration, well losses.

Required Figures (references only):
- Fig 1: Time–drawdown semilog plot with fitted window
- Fig 2: Recovery plot (if available)
- Fig 3: Water-level time series (if available)

Required Tables:
- Table 1: Time–drawdown data
- Table 2: Computed parameters (Q, slope, T, S if computed) + units
- Table 3: Fit diagnostics (R², selected fitting interval) if available

Output: Azerbaijani, engineer tone, no generic claims.
"""
}

REP_AUTO_PROMPT = """TASK: Choose the best report type (A/B/C) based on available data columns and outputs.
Rules:
- If pumping test columns (time, drawdown, Q) exist → choose C.
- Else if irrigation indices can be computed (EC/TDS + Na/Ca/Mg) → choose B.
- Else choose A.
Then generate the report using the corresponding template strictly.
Also output a one-line justification: “Selected report type: X because …”
"""

def ai_generate_report_v2(df, report_type, meta_info, available_outputs, standard_set="None"):
    client = get_genai_client()
    if not client: return "AI not configured."
    
    # Data Prep
    ds_table = df.head(50).to_markdown(index=False) 
    units = {c: "mg/L (assumed)" if c in ['Ca','Mg','Na','K','HCO3','Cl','SO4','NO3'] else "-" for c in df.columns}
    coords = "Lat/Lon present" if 'Lat' in df.columns and 'Lon' in df.columns else "No coordinates"
    
    # Fill System Prompt placeholders
    sys_filled = REP_SYS_PROMPT.replace("{{project_meta}}", str(meta_info))
    sys_filled = sys_filled.replace("{{coordinates}}", coords)
    sys_filled = sys_filled.replace("{{dataset_table}}", ds_table)
    sys_filled = sys_filled.replace("{{units_table}}", str(units))
    sys_filled = sys_filled.replace("{{available_outputs}}", str(available_outputs))
    sys_filled = sys_filled.replace("{{selected_standard_set}}", standard_set)
    
    # Select Template
    if report_type == "Auto":
        user_task = REP_AUTO_PROMPT + "\\n\\nAvailable Templates:\\n" + \
                    "A:\\n" + REP_TEMPLATES["A"] + "\\n\\n" + \
                    "B:\\n" + REP_TEMPLATES["B"] + "\\n\\n" + \
                    "C:\\n" + REP_TEMPLATES["C"]
    else:
        user_task = REP_TEMPLATES.get(report_type, REP_TEMPLATES["A"])
        
    final_prompt = sys_filled + "\\n\\n" + user_task
    
    try:
        response = call_gemini_with_retry(
            client,
            model="gemini-3-flash-preview",
            contents=final_prompt
        )
        return response.text
    except Exception as e: return f"AI Error: {e}"

def ai_read_image(file_bytes, mime_type="image/jpeg"):
    try:
        client = get_genai_client()
        if not client:
            st.warning(t("ai_not_configured"))
            return None
            
        prompt = (
            "Analyze this image/document (potentially an ArcGIS export, Attribute Table, or Map). "
            "Goal: Extract water quality sampling data and coordinates. "
            "1. Look for tables, lists, or labeled markers with hydrochemical data. "
            "2. Extract fields: ID (Sample Name), Lat (Latitude), Lon (Longitude), Ca, Mg, Na, Cl, SO4, HCO3, TDS, pH, WL (Water Level). "
            "3. If Lat/Lon are separate columns (X/Y), convert or keep them. "
            "4. Output a STRICT JSON array of objects. Example: [{\"ID\":\"1\",\"Lat\":40.5,\"Ca\":120}]. "
            "5. If no specific data is visible, return []. "
            "Return ONLY the JSON string, no markdown."
        )
        
        st.info(f"🔍 Analyzing document with Gemini 3 Flash Preview...")
        response = call_gemini_with_retry(
            client,
            model="gemini-3-flash-preview",
            contents=[
                prompt,
                types.Part.from_bytes(data=file_bytes, mime_type=mime_type)
            ]
        )
        content = response.text
        
        # Debug: check raw content if needed
        # st.text(content)

        # Clean up JSON
        json_str = content.replace("```json", "").replace("```", "").strip()
        if not json_str.startswith("["):
            start = json_str.find("[")
            end = json_str.rfind("]")
            if start != -1 and end != -1 and end > start:
                json_str = json_str[start:end+1]
            else:
                # If no array found, try to find a single object and wrap it? Or just fail gracefully.
                return pd.DataFrame() # Return empty DF if parse fails

        data = json.loads(json_str)
        if not data:
            return pd.DataFrame()

        return pd.DataFrame(data)
    except Exception as e:
        st.error(f"OCR Error: {e}")
        return None

def ai_parse_lithology(text):
    try:
        if not text or not text.strip():
            st.warning("Please enter a description first.")
            return None
        client = get_genai_client()
        if not client:
            st.warning(t("ai_not_configured"))
            return None
            
        prompt = (
            "Extract lithology layers from the text. "
            "Classify 'Lithology' strictly into one of: Clay, Sand, Silt, Gravel, Limestone, Sandstone, Shale, Coal, Other. "
            "Return ONLY a valid JSON array with keys: "
            "From (number), To (number), Lithology (string), Description (string), Confidence (number 0.0-1.0). "
            "Example: [{\"From\":0,\"To\":5,\"Lithology\":\"Clay\",\"Description\":\"Brown sticky clay\",\"Confidence\":0.95}]\n\n"
            f"Text: {text}"
        )
        response = call_gemini_with_retry(
            client,
            model="gemini-3-flash-preview",
            contents=prompt
        )
        content = response.text.strip()
        
        # Remove code fences if present
        if content.startswith("```"):
            content = content.replace("```json", "").replace("```", "").strip()
        # Try to extract JSON array if extra text exists
        if not content.startswith("["):
            start = content.find("[")
            end = content.rfind("]")
            if start != -1 and end != -1 and end > start:
                content = content[start:end+1]
        data = json.loads(content)
        df = pd.DataFrame(data)
        # Normalize possible column names
        if 'Süxur' in df.columns and 'Rock' not in df.columns:
            df = df.rename(columns={'Süxur': 'Rock'})
            
        # Ensure standard columns exist
        if 'Lithology' not in df.columns and 'Rock' in df.columns:
             df['Lithology'] = df['Rock']
        
        # Add metadata if missing
        if 'Source' not in df.columns: df['Source'] = 'AI'
        if 'ID' not in df.columns: df['ID'] = [f'L{i+1}' for i in range(len(df))]
        
        return df
    except Exception as e:
        st.error(f"Lithology AI error: {e}")
        return None

# --- LITHOLOGY HELPERS ---

LITHOLOGY_STYLES = {
    "Clay": {"color": "#8D6E63", "pattern": "."},  # Brownish
    "Sand": {"color": "#FFF59D", "pattern": "."},  # Yellow
    "Silt": {"color": "#E0E0E0", "pattern": ""},   # Grey
    "Gravel": {"color": "#9E9E9E", "pattern": "x"}, # Pebbles (using 'x' as 'o' is not supported)
    "Limestone": {"color": "#4FC3F7", "pattern": "+"}, # Blueish blocky
    "Sandstone": {"color": "#FFCC80", "pattern": "."}, # Orange sand
    "Shale": {"color": "#3E2723", "pattern": "-"}, # Dark lines
    "Coal": {"color": "#212121", "pattern": ""},   # Black
    "Other": {"color": "#CFD8DC", "pattern": "/"},
    "Unknown": {"color": "#FFFFFF", "pattern": "/"}
}

def get_lith_style(lith_name):
    # Normalize
    clean_name = str(lith_name).strip().capitalize()
    
    # Direct match
    if clean_name in LITHOLOGY_STYLES:
        return LITHOLOGY_STYLES[clean_name]

    # Partial match
    for k in LITHOLOGY_STYLES.keys():
        if k.lower() in clean_name.lower():
            return LITHOLOGY_STYLES[k]
            
    return LITHOLOGY_STYLES["Other"]

def validate_lithology_log(df):
    issues = []
    if df is None or df.empty: return issues
    
    # Sort by depth
    df_sorted = df.sort_values("From")
    
    # Check From < To
    invalid_depth = df_sorted[df_sorted['From'] >= df_sorted['To']]
    for _, row in invalid_depth.iterrows():
        issues.append(f"⚠️ {t('lith_invalid_depth').format(row.get('Lithology','?'), row['From'], row['To'])}")
        
    # Check Overlaps/Gaps
    prev_to = 0
    if not df_sorted.empty:
        prev_to = df_sorted.iloc[0]['From']
        
    for i, row in df_sorted.iterrows():
        curr_from = row['From']
        curr_to = row['To']
        
        if curr_from < prev_to:
            issues.append(f"⚠️ {t('lith_overlap').format(curr_from)}")
        elif curr_from > prev_to:
            issues.append(f"⚠️ {t('lith_gap').format(prev_to, curr_from)}")
            
        prev_to = curr_to
        
    return issues

def analyze_well_design(df):
    """Generates design recommendations and QC metrics."""
    if df.empty: return {}, [], []
    
    # ensure thickness
    if 'Thickness' not in df.columns:
        df['Thickness'] = df['To'] - df['From']
        
    # 1. Metrics & Fraction
    total_depth = df['To'].max()
    clay_types = ['Clay', 'Silt', 'Shale', 'Loam', 'Soil']
    aquifer_types = ['Sand', 'Gravel', 'Sandstone', 'Limestone', 'Fractured Rock']
    
    clay_th = df[df['Lithology'].isin(clay_types)]['Thickness'].sum()
    sand_th = df[df['Lithology'].isin(aquifer_types)]['Thickness'].sum()
    
    metrics = {
        "clay_pct": (clay_th / total_depth * 100) if total_depth else 0,
        "sand_pct": (sand_th / total_depth * 100) if total_depth else 0,
        "total_depth": total_depth
    }
    
    # 2. Risk Flags
    risks = []
    # Thin layers (<0.5m) inside aquifer might be issues
    thin = df[df['Thickness'] < 1.0]
    for _, r in thin.iterrows():
        risks.append(t("risk_thin").format(r['From']))
        
    # 3. Design Suggestions
    design_recs = []
    
    for _, r in df.sort_values("From").iterrows():
        lith = r['Lithology']
        # normalize
        lith_clean = str(lith).strip().capitalize()
        style = get_lith_style(lith_clean)
        
        item = {"from": r['From'], "to": r['To'], "lith": lith}
        
        # Heuristic based on LITHOLOGY_STYLES keys or partial match
        if "Sand" in lith_clean and "Stone" not in lith_clean: # Sand
            item['type'] = "Screen"
            item['slot'] = "0.5 - 0.75 mm"
            item['pack'] = "1.0 - 2.0 mm (Quartz)"
        elif "Gravel" in lith_clean:
            item['type'] = "Screen"
            item['slot'] = "1.5 - 2.5 mm"
            item['pack'] = "3.0 - 5.0 mm"
        elif "Sandstone" in lith_clean or "Limestone" in lith_clean:
             item['type'] = "Screen/Open"
             item['slot'] = "1.0 mm"
             item['pack'] = "Stabilizer"
        else:
            item['type'] = "Blind Pipe"
            item['slot'] = "-"
            item['pack'] = "-"
            
        design_recs.append(item)
        
    return metrics, risks, design_recs


def ai_parse_table(text):
    try:
        if not text or not text.strip():
            st.warning(t("ai_parse_warning"))
            return None
        cmd = text.strip().lower()
        if cmd.startswith("random"):
            return generate_random_table(5)
        
        client = get_genai_client()
        if not client:
            st.warning(t("ai_not_configured"))
            return None
            
        prompt = (
            "Convert the following raw water-sample data into a JSON array. "
            "Return ONLY valid JSON. Use columns if present; otherwise infer: "
            "ID, Lat, Lon, Water Level, Ca, Mg, Na, HCO3, Cl, SO4. "
            "Example: [{\"ID\":\"Q1\",\"Lat\":40.4,\"Lon\":49.86,\"Water Level\":10.2,"
            "\"Ca\":50,\"Mg\":20,\"Na\":200,\"HCO3\":150,\"Cl\":200,\"SO4\":50}]\n\n"
            f"Data:\n{text}"
        )
        response = call_gemini_with_retry(
            client,
            model="gemini-3-flash-preview",
            contents=prompt
        )
        content = response.text.strip()
        
        if content.startswith("```"):
            content = content.replace("```json", "").replace("```", "").strip()
        if not content.startswith("["):
            start = content.find("[")
            end = content.rfind("]")
            if start != -1 and end != -1 and end > start:
                content = content[start:end+1]
        data = json.loads(content)
        return pd.DataFrame(data)
    except Exception as e:
        st.error(f"AI table parse error: {e}")
        return None

def generate_random_table(n=5):
    try:
        n = int(n)
        rng = np.random.default_rng()
        df = pd.DataFrame({
            'ID': [f'Q{i+1}' for i in range(n)],
            'Lat': np.round(rng.uniform(40.35, 40.45, n), 4),
            'Lon': np.round(rng.uniform(49.80, 49.90, n), 4),
            'Water Level': np.round(rng.uniform(8.0, 13.0, n), 2),
            'Ca': np.round(rng.uniform(20, 80, n), 0),
            'Mg': np.round(rng.uniform(5, 35, n), 0),
            'Na': np.round(rng.uniform(80, 260, n), 0),
            'HCO3': np.round(rng.uniform(80, 220, n), 0),
            'Cl': np.round(rng.uniform(80, 240, n), 0),
            'SO4': np.round(rng.uniform(20, 90, n), 0),
        })
        return df
    except Exception as e:
        st.error(f"Random table error: {e}")
        return None

def predict_future(df, years_to_predict=5, model_type='Linear'):
    try:
        # Check for English or Azerbaijani column names
        col_year = 'Year' if 'Year' in df.columns else 'İl'
        col_value = 'Value' if 'Value' in df.columns else 'Dəyər'

        if col_year not in df.columns or col_value not in df.columns:
            return None
            
        df = df.sort_values(col_year)
        x = df[col_year].values
        y = df[col_value].values
        n = len(x)
        
        if n < 2: return None

        # Degree selection
        degree = 2 if model_type == 'Polynomial' and n >= 3 else 1
        
        # Fit
        z = np.polyfit(x, y, degree)
        p = np.poly1d(z)
        
        # Historical prediction for metrics
        y_pred = p(x)
        
        # Metrics
        if n > 1:
            sss_tot = np.sum((y - np.mean(y))**2)
            sss_res = np.sum((y - y_pred)**2)
            r2 = 1 - (sss_res / sss_tot) if sss_tot > 0 else 0
            mae = np.mean(np.abs(y - y_pred))
        else:
            r2, mae = 0, 0
            
        slope = z[0] if degree == 1 else None

        # Future
        last_year = int(x.max())
        future_years = np.arange(last_year + 1, last_year + years_to_predict + 1)
        future_values = p(future_years)
        
        # Confidence Intervals (95%) for Linear
        ci_lower, ci_upper = None, None
        if degree == 1 and n > 2:
            # t-critical for 95% (approx 2.0 for simplicity)
            t_val = 2.0 
            
            # Standard Error of Estimate
            residuals = y - y_pred
            sse = np.sum(residuals**2)
            se = np.sqrt(sse / (n - 2))
            
            # CI Calculation
            x_bar = np.mean(x)
            s_xx = np.sum((x - x_bar)**2)
            
            # Forecast Error
            se_forecast = se * np.sqrt(1 + (1/n) + ((future_years - x_bar)**2) / s_xx)
            
            ci_lower = future_values - t_val * se_forecast
            ci_upper = future_values + t_val * se_forecast

        return {
            "years": future_years,
            "values": future_values,
            "ci_lower": ci_lower,
            "ci_upper": ci_upper,
            "metrics": {"R2": r2, "MAE": mae, "Slope": slope},
            "model": "Linear" if degree == 1 else "Polynomial",
            "history_x": x,
            "history_y": y,
            "history_pred": y_pred
        }
    except Exception as e: 
        return None

def create_word_report(text, stats):
    doc = Document(); doc.add_paragraph(text); bio = io.BytesIO(); doc.save(bio); bio.seek(0); return bio

# --- 7. CHARTS (existing diagram functions) ---
def draw_piper_diagram(df):
    try:
        plot_df = df.copy()
        req = ['Ca', 'Mg', 'Na', 'HCO3', 'Cl', 'SO4', 'TDS']
        for col in req:
            if col in plot_df.columns: plot_df[col] = pd.to_numeric(plot_df[col], errors='coerce').fillna(0)
        cat_sum = (plot_df['Ca']/20.04) + (plot_df['Mg']/12.15) + (plot_df['Na']/23.0)
        an_sum = (plot_df['HCO3']/61.0) + (plot_df['Cl']/35.45) + (plot_df['SO4']/48.0)
        x = ((plot_df['Ca']/20.04) / cat_sum.replace(0,1)) * 100
        y = ((plot_df['Cl']/35.45) / an_sum.replace(0,1)) * 100
        fig = go.Figure(go.Scatter(x=x, y=y, mode='markers', marker=dict(size=14, color=plot_df['TDS'], colorscale='Viridis', showscale=True)))
        fig.update_layout(title="Piper", xaxis_title="Ca%", yaxis_title="Cl%")
        return fig
    except: return None

def draw_schoeller_diagram(df):
    try:
        ions = ['Ca', 'Mg', 'Na', 'Cl', 'SO4', 'HCO3']
        fig = go.Figure()
        for i, row in df.iterrows():
            vals = [row.get(ion, 0) if row.get(ion, 0)>0 else 0.01 for ion in ions]
            fig.add_trace(go.Scatter(x=ions, y=vals, mode='lines+markers', name=str(row.get('ID',i))))
        fig.update_layout(title="Schoeller", yaxis_type="log")
        return fig
    except: return None

def draw_stiff_diagram(df):
    charts = []
    try:
        for i, row in df.iterrows():
            na = row.get('Na',0)/23.0; ca = row.get('Ca',0)/20.04; mg = row.get('Mg',0)/12.15
            cl = row.get('Cl',0)/35.45; hco3 = row.get('HCO3',0)/61.0; so4 = row.get('SO4',0)/48.0
            fig, ax = plt.subplots(figsize=(3, 3))
            verts = [(-na, 3), (cl, 3), (hco3, 2), (so4, 1), (-mg, 1), (-ca, 2)]
            poly = Polygon(verts, closed=True, facecolor='#0079c1', edgecolor='black', alpha=0.7, linewidth=1)
            ax.add_patch(poly)
            max_val = max(na, ca, mg, cl, hco3, so4); limit = max_val * 1.2 if max_val > 0 else 1
            ax.set_xlim(-limit, limit); ax.set_ylim(0.5, 3.5); ax.axvline(0, color='black', linewidth=1)
            ax.grid(True, linestyle=':', alpha=0.5); ax.set_yticks([1, 2, 3])
            ax.set_yticklabels(['Mg|SO4', 'Ca|HCO3', 'Na|Cl'], fontsize=8, fontweight='bold')
            ax.set_title(str(row.get('ID', f'Q-{i}')), fontsize=10)
            plt.tight_layout(); charts.append(fig)
    except: return []
    return charts

# --- STIFF HELPERS ---
def create_stiff_on_ax(ax, row, sample_label, fontsize=10):
    try:
        na = pd.to_numeric(row.get('Na',0), errors='coerce')/23.0
        ca = pd.to_numeric(row.get('Ca',0), errors='coerce')/20.04
        mg = pd.to_numeric(row.get('Mg',0), errors='coerce')/12.15
        cl = pd.to_numeric(row.get('Cl',0), errors='coerce')/35.45
        hco3 = pd.to_numeric(row.get('HCO3',0), errors='coerce')/61.0
        so4 = pd.to_numeric(row.get('SO4',0), errors='coerce')/48.0
        
        na = 0 if pd.isna(na) else na; ca = 0 if pd.isna(ca) else ca; mg = 0 if pd.isna(mg) else mg
        cl = 0 if pd.isna(cl) else cl; hco3 = 0 if pd.isna(hco3) else hco3; so4 = 0 if pd.isna(so4) else so4

        verts = [(-na, 3), (cl, 3), (hco3, 2), (so4, 1), (-mg, 1), (-ca, 2)]
        poly = Polygon(verts, closed=True, facecolor='#0079c1', edgecolor='black', alpha=0.7, linewidth=1)
        ax.add_patch(poly)
        max_val = max(na, ca, mg, cl, hco3, so4)
        limit = max_val * 1.2 if max_val > 0 else 1
        ax.set_xlim(-limit, limit); ax.set_ylim(0.5, 3.5); ax.axvline(0, color='black', linewidth=1)
        ax.grid(True, linestyle=':', alpha=0.5); ax.set_yticks([1, 2, 3])
        ax.set_yticklabels(['Mg|SO4', 'Ca|HCO3', 'Na|Cl'], fontsize=8, fontweight='bold')
        ax.set_title(str(sample_label), fontsize=fontsize)
    except: pass

def create_single_stiff_fig(row, sample_label, figsize=(3, 3), fontsize=10):
    fig, ax = plt.subplots(figsize=figsize)
    create_stiff_on_ax(ax, row, sample_label, fontsize)
    plt.tight_layout()
    return fig

def get_stiff_base64(row, sample_label):
    fig = create_single_stiff_fig(row, sample_label, figsize=(2, 2), fontsize=8)
    if fig:
        buf = io.BytesIO()
        fig.savefig(buf, format='png', bbox_inches='tight', transparent=True)
        plt.close(fig)
        buf.seek(0)
        return base64.b64encode(buf.read()).decode('utf-8')
    return None

def perform_gibbs_analysis(df):
    """
    Analyzes Gibbs diagram data to provide:
    1. Dataset summary (Dominant Process)
    2. Per-sample classification
    3. QC Warnings
    """
    analysis = {
        "summary": {"Rock weathering": 0, "Evaporation": 0, "Precipitation": 0},
        "samples": [],
        "qc": {
            "tds_unit_issues": [],
            "ratio_issues": [],
            "outliers": []
        }
    }
    
    # Check for required columns
    required = ['TDS', 'Na', 'Ca']
    if not all(col in df.columns for col in required):
        return analysis

    # Calculate global stats for outlier detection
    mean_tds = df['TDS'].mean()
    std_tds = df['TDS'].std()
    
    total_samples = len(df)
    
    for i, row in df.iterrows():
        sample_id = row.get('ID', f"Sample-{i}")
        tds = pd.to_numeric(row.get('TDS', 0), errors='coerce') or 0
        na = pd.to_numeric(row.get('Na', 0), errors='coerce') or 0
        ca = pd.to_numeric(row.get('Ca', 0), errors='coerce') or 0
        
        # Avoid division by zero
        denom = na + ca
        if denom == 0:
            ratio = 0 
        else:
            ratio = na / denom
            
        # Classification Logic
        # Precipitation: Low TDS (< 100 approx)
        # Evaporation: High TDS (> 1000 approx)
        # Rock Weathering: Intermediate TDS
        
        zone = "Rock weathering"
        if tds < 100:
            zone = "Precipitation"
        elif tds > 1000:
             zone = "Evaporation"
        
        analysis["summary"][zone] += 1
        
        # QC Checks
        # 1. Ratio 0-1
        if not (0 <= ratio <= 1.05): # slight tolerance
             analysis["qc"]["ratio_issues"].append(sample_id)
        
        # 2. TDS Unit (Heuristic)
        if tds <= 0:
             analysis["qc"]["tds_unit_issues"].append(f"{sample_id} (<=0)")

        # 3. Outlier ( > 3 std dev)
        if total_samples >= 3 and std_tds > 0:
             z_score = abs(tds - mean_tds) / std_tds
             if z_score > 3:
                 analysis["qc"]["outliers"].append(f"{sample_id} (Z={z_score:.1f})")

        # Comment
        comment = f"TDS={tds:.0f}, Ratio={ratio:.2f}"
        
        analysis["samples"].append({
            "Sample ID": sample_id,
            "Zone": zone,
            "Qısa şərh": comment
        })

    return analysis

def draw_gibbs_diagram(df):
    try:
        ratio = df['Na'] / (df['Na'] + df['Ca']).replace(0,1)
        fig = go.Figure(go.Scatter(x=ratio, y=df['TDS'], mode='markers'))
        fig.update_yaxes(type="log", range=[1, 5]); fig.update_layout(title="Gibbs")
        return fig
    except: return None

def cooper_jacob_analysis(t, s, Q):
    try:
        log_t = np.log10(t)
        slope, intercept = np.polyfit(log_t, s, 1)
        T = (0.183 * Q) / abs(slope)
        return T, slope, intercept
    except: return None, None, None

# --- NEW FEATURE: STATISTICS AND CLUSTERING ---
def compute_statistics(df, cols):
    """Return descriptive statistics and correlation matrix for selected columns."""
    stats = df[cols].describe().transpose()
    corr = df[cols].corr(method='pearson')
    return stats, corr

def perform_clustering(df, cols, n_clusters=3):
    """Apply k-means clustering to the specified columns and return labels."""
    data = df[cols].fillna(0).values
    model = KMeans(n_clusters=n_clusters, n_init=10, random_state=42)
    labels = model.fit_predict(data)
    return labels

# --- NEW HELPERS FOR STATISTICS TAB ---
def calculate_hardness(row):
    try:
        ca = float(row.get('Ca', 0)) if pd.notnull(row.get('Ca')) else 0
        mg = float(row.get('Mg', 0)) if pd.notnull(row.get('Mg')) else 0
        return 2.497 * ca + 4.118 * mg
    except:
        return np.nan

def calculate_cbe(row):
    """
    Calculate Charge Balance Error (CBE).
    Formula: (Sum Cations - Sum Anions) / (Sum Cations + Sum Anions) * 100
    Units: meq/L.
    Assumes input columns (Ca, Mg, Na, K, HCO3, Cl, SO4) are in mg/L.
    """
    try:
        # Cations (meq/L)
        # Ca: 40.08 (z=2) -> /20.04
        # Mg: 24.31 (z=2) -> /12.15
        # Na: 22.99 (z=1) -> /22.99
        # K: 39.10 (z=1) -> /39.10
        ca = float(row.get('Ca', 0)) if pd.notnull(row.get('Ca')) else 0
        mg = float(row.get('Mg', 0)) if pd.notnull(row.get('Mg')) else 0
        na = float(row.get('Na', 0)) if pd.notnull(row.get('Na')) else 0
        k = float(row.get('K', 0)) if pd.notnull(row.get('K')) else 0
        
        sum_cations = (ca/20.04 + mg/12.15 + na/22.99 + k/39.10)
        
        # Anions (meq/L)
        # HCO3: 61.02 (z=1) -> /61.02
        # Cl: 35.45 (z=1) -> /35.45
        # SO4: 96.06 (z=2) -> /48.03
        hco3 = float(row.get('HCO3', 0)) if pd.notnull(row.get('HCO3')) else 0
        cl = float(row.get('Cl', 0)) if pd.notnull(row.get('Cl')) else 0
        so4 = float(row.get('SO4', 0)) if pd.notnull(row.get('SO4')) else 0

        sum_anions = (hco3/61.02 + cl/35.45 + so4/48.03)
        
        # Add optional if exist
        if 'CO3' in row and pd.notnull(row['CO3']): 
            sum_anions += float(row['CO3'])/30.00
        if 'NO3' in row and pd.notnull(row['NO3']): 
            sum_anions += float(row['NO3'])/62.00
            
        if (sum_cations + sum_anions) == 0:
            return 0.0
            
        return (sum_cations - sum_anions) / (sum_cations + sum_anions) * 100.0
    except:
        return np.nan

def generate_insights(df, selected_cols):
    insights = []
    
    # 1. Correlation Insights
    if len(selected_cols) >= 2:
        # Check numeric only just in case
        numeric_df = df[selected_cols].select_dtypes(include=[np.number])
        if not numeric_df.empty:
            corr_matrix = numeric_df.corr().abs()
            # Get upper triangle
            upper = corr_matrix.where(np.triu(np.ones(corr_matrix.shape), k=1).astype(bool))
            strong_pairs = upper.stack()
            strong_pairs = strong_pairs[strong_pairs > 0.7] # Threshold
            
            # Sort by strength
            strong_pairs = strong_pairs.sort_values(ascending=False).head(5)
            
            for idx, val in strong_pairs.items():
                col1, col2 = idx
                insights.append(f"• Strong correlation detected between **{col1}** and **{col2}** (r={val:.2f}).")
            
    # 2. TDS Variability
    if 'TDS' in df.columns and pd.api.types.is_numeric_dtype(df['TDS']):
        mean_tds = df['TDS'].mean()
        std_tds = df['TDS'].std()
        if mean_tds > 0:
            cv = (std_tds / mean_tds) * 100
            variability = "High" if cv > 50 else "Low"
            insights.append(f"• TDS variability is **{variability}** (CV={cv:.1f}%).")
            
    # 3. High Salinity
    if 'TDS' in df.columns and pd.api.types.is_numeric_dtype(df['TDS']):
        high_salinity = df[df['TDS'] > 1000].shape[0]
        if high_salinity > 0:
            insights.append(f"• **{high_salinity}** samples have high Salinity (>1000 mg/L).")
            
    if not insights:
        insights.append("No specialized insights detected.")
        
    return insights

# --- HELPER FUNCTIONS FOR ANALYSIS PANEL ---
def render_active_columns_selection(df):
    """
    Refactored Active Columns Selection with 3 modes.
    Returns the list of selected columns.
    """
    st.markdown("### 🛠 Column Configuration")
    
    all_cols = list(df.columns)
    
    # Standardize column names for checking but keep original for selection
    common_ions = ['Ca', 'Mg', 'Na', 'K', 'HCO3', 'Cl', 'SO4', 'EC', 'TDS', 'pH', 'WL']
    hydro_preset = ['Ca', 'Mg', 'Na', 'K', 'HCO3', 'Cl', 'SO4', 'TDS']
    irrigation_preset = ['Na', 'Ca', 'Mg', 'HCO3', 'Cl']
    hydraulics_preset = ['WL', 'K', 'T']
    
    mode = st.radio(
        "Selection Mode", 
        ["Auto-Detect", "Presets", "Advanced (Manual)"], 
        horizontal=True,
        key="col_select_mode"
    )
    
    selected_cols = []
    
    if mode == "Auto-Detect":
        selected_cols = [c for c in common_ions if c in df.columns]
        if not selected_cols:
            st.warning("No standard hydrochemistry columns detected.")
            if len(all_cols) > 0:
                selected_cols = all_cols[:min(5, len(all_cols))]
        else:
            st.info(f"Auto-detected: {', '.join(selected_cols)}")
            
    elif mode == "Presets":
        preset_choice = st.selectbox("Choose Preset", ["Hydrochemistry", "Irrigation", "Hydraulics"], key="preset_choice")
        target_list = []
        if preset_choice == "Hydrochemistry":
            target_list = hydro_preset
        elif preset_choice == "Irrigation":
            target_list = irrigation_preset
        elif preset_choice == "Hydraulics":
            target_list = hydraulics_preset
            
        selected_cols = [c for c in target_list if c in df.columns]
        
        missing = set(target_list) - set(df.columns)
        if missing:
             pass # Just filter silent
        st.write(f"Active: {', '.join(selected_cols)}")
        
    else: # Advanced (Manual)
        default_sel = [c for c in common_ions if c in df.columns]
        if not default_sel: default_sel = all_cols[:5] if len(all_cols) >=5 else all_cols
        selected_cols = st.multiselect("Select Columns (Manual Override)", all_cols, default=default_sel, key="manual_col_select")
        
    return selected_cols

def render_data_health_bar(df):
    """
    Displays a Data Health (QC) bar with metrics.
    """
    with st.container():
        st.markdown("### 🏥 Data Health")
        cols = st.columns(4)
        
        # 1. Missing Values
        missing_count = df.isnull().sum().sum()
        cols[0].metric("Missing Values", f"{missing_count}", delta="High" if missing_count > 0 else None, delta_color="inverse")
        
        # 2. Dimensions
        cols[1].metric("Dimensions", f"{df.shape[0]} Rows, {df.shape[1]} Cols")
        
        # 3. CBE Warning (Quick Estimate)
        factors = {
             'Ca': 1/20.04, 'Mg': 1/12.15, 'Na': 1/22.99, 'K': 1/39.09,
             'HCO3': 1/61.01, 'Cl': 1/35.45, 'SO4': 1/48.03
        }
        
        cations = 0; anions = 0
        has_chem = False
        for c in ['Ca', 'Mg', 'Na', 'K']:
            if c in df.columns:
                cations += pd.to_numeric(df[c], errors='coerce').fillna(0) * factors.get(c, 0)
                has_chem = True
        
        for a in ['HCO3', 'Cl', 'SO4']:
            if a in df.columns:
                anions += pd.to_numeric(df[a], errors='coerce').fillna(0) * factors.get(a, 0)
                has_chem = True

        if has_chem:
            denom = cations + anions
            non_zero = denom > 0
            if non_zero.any(): 
                # If scalar
                if isinstance(denom, (int, float)):
                   cbe = ((cations - anions) / denom) * 100
                   avg_cbe = abs(cbe)
                else: 
                   cbe = ((cations - anions).abs() / denom) * 100
                   cbe = cbe[non_zero]
                   avg_cbe = cbe.mean()
                
                delta_str = "OK"
                delta_col = "normal"
                if avg_cbe > 10:
                    delta_str = "High Error"
                    delta_col = "inverse"
                
                cols[2].metric("Avg CBE", f"{avg_cbe:.1f}%", delta=delta_str, delta_color=delta_col)
            else:
                cols[2].metric("Avg CBE", "N/A")
        else:
            cols[2].metric("Avg CBE", "No Ions")

        # 4. Units Detected
        header_str = " ".join(df.columns).lower()
        if any(x in header_str for x in ['mg/l', 'ppm', 'meq/l']):
            cols[3].metric("MetaData", "Units Found")
        else:
            cols[3].metric("MetaData", "No Units", delta="Check units", delta_color="off")
        st.markdown("---") 

def get_template_df():
    tpl = {
        'ID': ['Q1', 'Q2', 'Q3', 'Q4', 'Q5'],
        'Lat': [40.40, 40.41, 40.39, 40.42, 40.38],
        'Lon': [49.86, 49.87, 49.85, 49.88, 49.89],
        'Water Level': [10.2, 9.8, 11.1, 10.6, 12.0],
        'Ca': [50, 60, 40, 55, 30],
        'Mg': [20, 25, 15, 22, 10],
        'Na': [200, 220, 180, 210, 150],
        'K': [5, 6, 4, 5, 3],
        'HCO3': [150, 160, 140, 155, 130],
        'Cl': [200, 210, 190, 205, 120],
        'SO4': [50, 55, 45, 52, 40],
        'pH': [7.2, 7.5, 7.1, 7.4, 7.3],
        'TDS': [500, 600, 480, 550, 400]
    }
    return pd.DataFrame(tpl)

# --- 9. MAIN APP LOGIC ---
def main_app():
    st.sidebar.image("https://cdn-icons-png.flaticon.com/512/2823/2823620.png", width=50)
    st.sidebar.title("HydroMind Pro")
    if 'lang' not in st.session_state:
        st.session_state['lang'] = "en"
    st.sidebar.selectbox(t("language"), options=["English", "Azərbaycan"], index=0 if st.session_state['lang']=="en" else 1,
                         key="lang_select")
    st.session_state['lang'] = "en" if st.session_state['lang_select'] == "English" else "az"
    st.sidebar.write(f"👤 **{st.session_state.get('username', 'User')}**")

    # ---------------------------------------------------------------------
    # AI Service Configuration UI
    #
    # Allow the user to provide their Gemini API key.
    # An acceptance checkbox is also provided.
    
    if "user_api_key" not in st.session_state:
        st.session_state["user_api_key"] = ""
    if "ai_accept" not in st.session_state:
        st.session_state["ai_accept"] = False

    with st.sidebar.expander(t("ai_settings"), expanded=False):
        st.session_state["user_api_key"] = st.text_input(
            t("api_key"),
            value=st.session_state.get("user_api_key", ""),
            type="password",
            help="Get your key from Google AI Studio"
        )
        st.session_state["ai_accept"] = st.checkbox(
            t("accept_terms"),
            value=st.session_state.get("ai_accept", False)
        )
        st.info("Powered by Google Gemini 3 Flash Preview")

    if 'page_key' not in st.session_state:
        st.session_state['page_key'] = "home"
    menu_keys = ["home", "new_analysis", "archive"]
    menu_labels = [t("home"), t("new_analysis"), t("archive")]
    current_index = menu_keys.index(st.session_state['page_key']) if st.session_state['page_key'] in menu_keys else 0
    choice_label = st.sidebar.radio("Menyu", menu_labels, index=current_index)
    choice_key = menu_keys[menu_labels.index(choice_label)]
    if choice_key != st.session_state['page_key']:
        st.session_state['page_key'] = choice_key
        force_rerun()
    if st.sidebar.button(t("logout")):
        st.session_state['logged_in'] = False
        st.session_state['username'] = "Guest"
        force_rerun()
    # --- HOME PAGE ---
    if st.session_state['page_key'] == "home":
        st.markdown("<h1 class='big-header'>💧 HydroMind Pro</h1>", unsafe_allow_html=True)
        st.markdown(f"<p class='sub-header'>{t('home_sub')}</p>", unsafe_allow_html=True)
        c1, c2, c3, c4 = st.columns(4)
        with c1: st.markdown(f"""<div class='feature-card'><h3>{t('chemistry')}</h3><p>Piper, Stiff, Schoeller</p></div>""", unsafe_allow_html=True)
        with c2: st.markdown(f"""<div class='feature-card'><h3>{t('gis')}</h3><p>Kriging & Contours</p></div>""", unsafe_allow_html=True)
        with c3: st.markdown(f"""<div class='feature-card'><h3>{t('agriculture')}</h3><p>SAR & Irrigation</p></div>""", unsafe_allow_html=True)
        with c4: st.markdown(f"""<div class='feature-card'><h3>{t('ai')}</h3><p>OCR & Reports</p></div>""", unsafe_allow_html=True)
        st.markdown("---")
        col_text, col_img = st.columns([2, 1])
        with col_text:
            st.markdown(t("what_can"))
            st.info(f"""
            * {t('info_ocr')}
            * {t('info_forecast')}
            * {t('info_lithology')}
            * {t('info_report')}
            """)
            if st.button(t("start"), type="primary"):
                st.session_state['page_key'] = "new_analysis"
                force_rerun()
        with col_img:
            # Localized mock data
            m_months = [t('month_jan'), t('month_feb'), t('month_mar'), t('month_apr'), t('month_may')]
            mock_data = pd.DataFrame({'Date': m_months, t('discharge'): [10, 15, 13, 18, 22]})
            fig = px.bar(mock_data, x='Date', y=t('discharge'), title=t('system_stats'), color=t('discharge'))
            st.plotly_chart(fig, use_container_width=True)
    # --- ANALIZ MODULU ---
    elif st.session_state['page_key'] == "new_analysis":
        st.subheader(t("analysis_panel"))
        im = st.radio(t("data_input"), [t("upload_excel"), t("manual_entry")], horizontal=True)
        df = None
        if im == t("upload_excel"):
            col_up1, col_up2 = st.columns([3, 1])
            with col_up1:
                 f = st.file_uploader(t("excel_file"), ["xlsx"])
            with col_up2:
                 st.write("") 
                 st.write("")
                 # Download Template button
                 if st.button(t("template_btn")):
                     tpl_df = get_template_df()
                     csv = tpl_df.to_csv(index=False).encode('utf-8')
                     st.download_button(
                         label="Download CSV",
                         data=csv,
                         file_name="hydro_template.csv",
                         mime="text/csv",
                     )
            if f:
                df = pd.read_excel(f); df.columns = df.columns.str.strip()
        else:
            # Manual Table Entry
            with st.expander(t("ai_fill"), expanded=True):
                st.markdown(t("ai_data_fill_header"))
                fill_mode = st.radio(t("fill_mode"), [t("fill_mode_manual"), t("fill_mode_synthetic")], horizontal=True)
                
                if fill_mode == t("fill_mode_synthetic"):
                     n_rows = st.number_input(t("random_rows"), min_value=2, max_value=50, value=5, step=1)
                     if st.button(t("random_fill")):
                        rand_df = generate_random_table(n_rows)
                        st.session_state['ai_table_df'] = rand_df
                else:
                    st.caption(t("ai_caption"))
                    ai_text = st.text_area(t("ai_input"), height=100, key="ai_table_input", placeholder="Paste table data or description here...")
                    if st.button(t("ai_fill_btn")):
                        parsed = ai_parse_table(ai_text)
                        if parsed is not None and not parsed.empty:
                            st.session_state['ai_table_df'] = parsed

            # Initialize editor with template or session state
            if 'ai_table_df' not in st.session_state:
                st.session_state['ai_table_df'] = get_template_df()
                
            df = st.data_editor(st.session_state['ai_table_df'], num_rows="dynamic", key="main_data_editor")
                        
        if df is not None and not df.empty:
            # Basic Cleaning
            req = ['Ca','Mg','Na','Cl','SO4','HCO3']
            cols_to_clean = [c for c in req if c in df.columns]
            if cols_to_clean:
                df = clean_chemistry_data(df, cols_to_clean)
            
            if 'Water Level' in df.columns:
                 df = clean_chemistry_data(df, ['Water Level'])
                 
            # Calc TDS if missing
            if 'TDS' not in df.columns and set(req).issubset(df.columns):
                df['TDS'] = df[req].sum(axis=1)
            
            # Calc Water Type if TDS available
            if 'TDS' in df.columns:
                df['Water Type'] = df.apply(lambda row: "Fresh" if row['TDS'] < 1000 else "Saline", axis=1)
            
            if 'Lat' in df.columns and 'Lon' in df.columns:
                 df['Lat'] = df['Lat'].apply(lambda x: gps_duzelt(x,'lat'))
                 df['Lon'] = df['Lon'].apply(lambda x: gps_duzelt(x,'lon'))

            # --- 1. Data Health ---
            render_data_health_bar(df)
            
            # --- 2. Active Columns ---
            active_cols = render_active_columns_selection(df)

            # --- 3. Tabs ---
            # Define the visual order
            tab_labels = [
                "📊 Piper", "📈 Schoeller", "🧬 Gibbs", "🔷 Stiff", "💧 WQI",      # Hydrochem
                t("tabs_agri"),                                                 # Irrigation (Agri)
                t("tabs_pump"), t("tabs_darcy"), "💧 Water Level",             # Hydraulics (Pump, Darcy, WL)
                t("tabs_stats"), t("tabs_cluster"),                             # Stats (Stats, Cluster)
                "📷 OCR", t("tabs_forecast"), t("tabs_lith"),                   # Utils (OCR, Forecast, Lith)
                "🗺️ Maps", t("tabs_report"), t("tabs_save")                     # Output (Maps, Report, Save)
            ]
            real_tabs = st.tabs(tab_labels)
            
            # Map old indices to new positions in real_tabs
            tabs = [None] * 17
            tabs[0] = real_tabs[0] # Piper
            tabs[1] = real_tabs[1] # Schoeller
            tabs[2] = real_tabs[2] # Gibbs
            tabs[3] = real_tabs[3] # Stiff
            tabs[4] = real_tabs[5] # Agri
            tabs[5] = real_tabs[6] # Pump
            tabs[6] = real_tabs[13] # Lith
            tabs[7] = real_tabs[7] # Darcy
            tabs[8] = real_tabs[11] # OCR
            tabs[9] = real_tabs[12] # Forecast
            tabs[10] = real_tabs[15] # Report
            tabs[11] = real_tabs[14] # Maps
            tabs[12] = real_tabs[16] # Save
            tabs[13] = real_tabs[4] # WQI
            tabs[14] = real_tabs[9] # Stats
            tabs[15] = real_tabs[10] # Cluster
            tabs[16] = real_tabs[8] # Water Level
            # Existing diagrams
            # --- 0. Piper ---
            with tabs[0]: 
                col_p1, col_p2 = st.columns([3, 1])
                with col_p1:
                    fig_p = draw_piper_diagram(df)
                    if fig_p: fig_p.update_layout(title=t('piper_chart_title'))
                    st.plotly_chart(fig_p, use_container_width=True)
                with col_p2:
                    st.markdown(t("piper_info"))
                    st.info(t("piper_desc"))
                    st.markdown(t("piper_zones"))

            # --- 1. Schoeller ---
            with tabs[1]: 
                fig_s = draw_schoeller_diagram(df)
                if fig_s: fig_s.update_layout(title=t('schoeller_chart_title'))
                st.plotly_chart(fig_s, use_container_width=True)
                st.caption(t("schoeller_caption"))

            # --- 2. Gibbs ---
            with tabs[2]: 
                col_g1, col_g2 = st.columns([3, 1])
                with col_g1:
                    fig_g = draw_gibbs_diagram(df)
                    if fig_g: fig_g.update_layout(title=t('gibbs_chart_title'))
                    st.plotly_chart(fig_g, use_container_width=True)
                with col_g2:
                    st.markdown(t("gibbs_title"))
                    st.success(t("gibbs_desc"))
                    st.markdown(t("gibbs_zones"))

                st.divider()
                st.markdown(t("gibbs_engineer_help"))
                
                # Perform analysis
                analysis_res = perform_gibbs_analysis(df)
                
                # 1. Dataset Summary
                st.markdown(t("dominant_process"))
                summary = analysis_res["summary"]
                total = sum(summary.values())
                if total > 0:
                    col_summ1, col_summ2, col_summ3 = st.columns(3)
                    
                    # Manual mapping to maintain order or just iterate
                    # Rock weathering, Evaporation, Precipitation
                    
                    # Rock
                    c = summary.get("Rock weathering", 0)
                    pct = (c / total) * 100 if total else 0
                    col_summ1.metric(t("rock_weathering"), f"{c} ({pct:.0f}%)")
                    
                    # Evaporation
                    c = summary.get("Evaporation", 0)
                    pct = (c / total) * 100 if total else 0
                    col_summ2.metric(t("evaporation"), f"{c} ({pct:.0f}%)")
                    
                    # Precipitation
                    c = summary.get("Precipitation", 0)
                    pct = (c / total) * 100 if total else 0
                    col_summ3.metric(t("precipitation"), f"{c} ({pct:.0f}%)")
                    
                else:
                    st.write(t("no_data"))

                # 2. Per Sample
                st.markdown(t("auto_classification"))
                if analysis_res["samples"]:
                    st.dataframe(pd.DataFrame(analysis_res["samples"]), use_container_width=True)
                
                # 3. QC
                st.markdown(t("qc_warnings"))
                
                qc = analysis_res["qc"]
                # TDS
                if qc["tds_unit_issues"]:
                    st.error(t("qc_tds_issue").format(', '.join(qc['tds_unit_issues'])))
                else:
                    st.success(t("qc_tds_ok"))
                
                # Ratio
                if qc["ratio_issues"]:
                    st.error(t("qc_ratio_issue").format(', '.join(qc['ratio_issues'])))
                else:
                    st.success(t("qc_ratio_ok"))
                
                # Outliers
                if qc["outliers"]:
                    st.warning(t("qc_outlier_issue").format(', '.join(qc['outliers'])))
                else:
                    st.success(t("qc_outlier_ok"))

            # --- 3. Stiff ---
            with tabs[3]: 
                st.subheader(t("stiff_title"))
                st.caption(t("stiff_caption"))
                
                if df.empty:
                    st.warning(t("no_data"))
                else:
                    # -- Controls & Focus --
                    stiff_c1, stiff_c2 = st.columns([1, 2])
                    
                    with stiff_c1:
                        st.markdown(t("controls"))
                        
                        # Focus Selector
                        s_ids = df.get('ID', df.index.astype(str)).tolist()
                        sel_id = st.selectbox(t("select_sample_focus"), s_ids)
                        
                        # Exports
                        st.markdown(t("export"))
                        if st.button(t("download_png_grid")):
                            # On the fly generation
                            n = len(df)
                            cols_n = 4
                            rows_n = (n // cols_n) + (1 if n % cols_n > 0 else 0)
                            fig_grid, axes_grid = plt.subplots(rows_n, cols_n, figsize=(cols_n*3, rows_n*3))
                            axes_grid = axes_grid.flatten()
                            for idx, ax in enumerate(axes_grid):
                                if idx < n:
                                    r_ = df.iloc[idx]
                                    lb_ = r_.get('ID', f"Sample-{idx}")
                                    create_stiff_on_ax(ax, r_, lb_)
                                else:
                                    ax.axis('off')
                            plt.tight_layout()
                            buf = io.BytesIO()
                            fig_grid.savefig(buf, format='png', bbox_inches='tight')
                            plt.close(fig_grid)
                            buf.seek(0)
                            st.download_button(t("download_png_grid"), buf, "stiff_grid.png", "image/png")

                        if st.button(t("download_pdf_report")):
                             pdf_buf = io.BytesIO()
                             with PdfPages(pdf_buf) as pdf:
                                 # Page layout: 6 per page
                                 # Or just simple 1 per page for now or standard grid logic
                                 # Let's do simple loop
                                 for idx, r_ in df.iterrows():
                                     lb_ = r_.get('ID', f"Sample-{idx}")
                                     f_ = create_single_stiff_fig(r_, lb_)
                                     pdf.savefig(f_)
                                     plt.close(f_)
                             pdf_buf.seek(0)
                             st.download_button(t("download_pdf_report"), pdf_buf, "stiff_report.pdf", "application/pdf")
                        
                        # Toggle Map
                        show_stiff_map = st.checkbox(t("map_integration"), value=False)

                    with stiff_c2:
                        if sel_id:
                            # Large View
                            st.markdown(t("focused_view").format(sel_id))
                            selected_row = df[df.get('ID', df.index.astype(str)) == sel_id].iloc[0]
                            f_large = create_single_stiff_fig(selected_row, sel_id, figsize=(6, 6), fontsize=14)
                            if f_large:
                                c_show1, c_show2 = st.columns([1, 1])
                                c_show1.pyplot(f_large)
                                plt.close(f_large)
                                c_show2.dataframe(selected_row.to_frame().T, use_container_width=True)

                    # -- Map View --
                    if show_stiff_map:
                        st.divider()
                        st.markdown(t("stiff_map_preview"))
                        if 'Lat' in df.columns and 'Lon' in df.columns:
                            valid_rows = df.dropna(subset=['Lat', 'Lon'])
                            if not valid_rows.empty:
                                center_lat = valid_rows['Lat'].mean()
                                center_lon = valid_rows['Lon'].mean()
                                m = folium.Map(location=[center_lat, center_lon], zoom_start=10)
                                
                                for idx, r_ in valid_rows.iterrows():
                                    sid = r_.get('ID', f"{idx}")
                                    img_b64 = get_stiff_base64(r_, sid)
                                    if img_b64:
                                        html = f'<div style="text-align:center;"><b>{sid}</b><br><img src="data:image/png;base64,{img_b64}" width="150"></div>'
                                        iframe = folium.IFrame(html, width=180, height=180)
                                        popup = folium.Popup(iframe, max_width=200)
                                        folium.Marker(
                                            [r_['Lat'], r_['Lon']], 
                                            popup=popup, 
                                            tooltip=str(sid),
                                            icon=folium.Icon(color='blue', icon='info-sign')
                                        ).add_to(m)
                                st_folium(m, height=500, width=800, key="stiff_map_view")
                            else:
                                st.warning("No valid Lat/Lon coordinates found.")
                        else:
                            st.warning("Data does not contain 'Lat' and 'Lon' columns.")
                    
                    st.divider()
                    st.markdown(t("all_diagrams"))
                    ch = draw_stiff_diagram(df)
                    if ch:
                        c = st.columns(4) 
                        for i, x in enumerate(ch): c[i % 4].pyplot(x); plt.close(x)

            # --- Aqrar (SAR & Irrigation classification) ---
            with tabs[4]:
                st.subheader(f"🌾 {t('tabs_agri')} {t('agri_analysis')}")
                if all(x in df.columns for x in ['Na', 'Ca', 'Mg']):
                    res_df = df.copy()
                    
                    # 1. Calculations
                    na_meq = res_df['Na']/23.0; ca_meq = res_df['Ca']/20.04; mg_meq = res_df['Mg']/12.15
                    res_df['SAR'] = na_meq / np.sqrt((ca_meq + mg_meq) / 2)
                    res_df['Na %'] = (na_meq / (na_meq + ca_meq + mg_meq)) * 100
                    
                    # Ensure EC and TDS availability
                    if 'EC' not in res_df.columns and 'TDS' in res_df.columns:
                        res_df['EC'] = res_df['TDS'] * 2  # Approx EC from TDS
                    elif 'EC' in res_df.columns and 'TDS' not in res_df.columns:
                         res_df['TDS'] = res_df['EC'] / 2 # Approx TDS from EC
                    elif 'EC' not in res_df.columns and 'TDS' not in res_df.columns:
                         st.error("Dataset needs either 'EC' or 'TDS' column.")
                         st.stop()
                    
                    # Classifications
                    def salinity_class(ec):
                        if ec < 250: return t('c1_class')
                        elif ec < 750: return t('c2_class')
                        elif ec < 2250: return t('c3_class')
                        else: return t('c4_class')
                    
                    def sodium_class(sar):
                        if sar < 10: return t('s1_class')
                        elif sar < 18: return t('s2_class')
                        elif sar < 26: return t('s3_class')
                        else: return t('s4_class')
                        
                    res_df['C_Class'] = res_df['EC'].apply(salinity_class)
                    res_df['S_Class'] = res_df['SAR'].apply(sodium_class)
                    
                    # 2. Logic for new columns: Suitability & Reason
                    def get_suitability_reason(row):
                        c_str = str(row['C_Class'])
                        s_str = str(row['S_Class'])
                        
                        # Extract codes (e.g., C3 from "C3 (High Salinity)")
                        c_code = c_str.split(' ')[0] if len(c_str) > 0 else "C?"
                        s_code = s_str.split(' ')[0] if len(s_str) > 0 else "S?"
                        
                        # Extract numeric level
                        try:
                            c_val = int(c_code[1]) if len(c_code) > 1 and c_code[1].isdigit() else 1
                            s_val = int(s_code[1]) if len(s_code) > 1 and s_code[1].isdigit() else 1
                        except:
                            c_val = 1
                            s_val = 1

                        # Suitability
                        max_val = max(c_val, s_val)
                        if max_val == 1: suitability = "OK"
                        elif max_val == 2: suitability = "OK (Conditional)"
                        elif max_val == 3: suitability = "Caution"
                        else: suitability = "Not Recommended"

                        # Reason construction
                        reasons = []
                        if c_val >= 3:
                            reasons.append(f"{c_code} due to EC={row['EC']:.0f} µS/cm")
                        elif c_val == 2 and suitability == "OK (Conditional)":
                             reasons.append(f"{c_code} (Mod. Salinity)")
                        
                        if s_val >= 3:
                             reasons.append(f"{s_code} due to SAR={row['SAR']:.1f}")
                        elif s_val == 2 and suitability == "OK (Conditional)":
                             reasons.append(f"{s_code} (Mod. Sodium)")
                             
                        if not reasons:
                             reasons.append("Safe limits")

                        return suitability, ", ".join(reasons)

                    res_df[['Overall Suitability', 'Reason']] = res_df.apply(
                        lambda row: pd.Series(get_suitability_reason(row)), axis=1
                    )
                    
                    # Format columns for display
                    res_df['TDS (mg/L)'] = res_df['TDS'] # Renaming for display
                    
                    # --- NEW LAYOUT ---
                    
                    # TOP ROW: Summary + Metrics
                    st.markdown(t("summary_overview"))
                    top_c1, top_c2, top_c3, top_c4 = st.columns([2, 1, 1, 1])
                    
                    params_stats = {
                        "Avg EC": f"{res_df['EC'].mean():.0f} µS/cm",
                        "Max SAR": f"{res_df['SAR'].max():.2f}",
                        "Avg Na%": f"{res_df['Na %'].mean():.1f}%"
                    }
                    dom_c = res_df['C_Class'].mode()[0].split(' ')[0] if not res_df['C_Class'].empty else "N/A"
                    dom_s = res_df['S_Class'].mode()[0].split(' ')[0] if not res_df['S_Class'].empty else "N/A"
                    
                    with top_c1:
                         st.info(t("dominant_classification").format(dom_c, dom_s) + "\n\n" + t("dominant_desc"))
                    with top_c2:
                         st.metric("Avg EC", params_stats["Avg EC"])
                    with top_c3:
                         st.metric("Max SAR", params_stats["Max SAR"])
                    with top_c4:
                         st.metric("Avg Na%", params_stats["Avg Na%"])
                    
                    st.divider()

                    # MIDDLE ROW: Chart (Left) vs Table+Actions (Right)
                    mid_left, mid_right = st.columns([4, 6])
                    
                    with mid_left:
                        st.markdown(f"#### {t('wilcox_title')}")
                        fig = px.scatter(
                            res_df, x="SAR", y="Na %", color="Overall Suitability",
                            hover_data=['ID', 'EC', 'C_Class', 'S_Class'],
                            color_discrete_map={"OK": "green", "OK (Conditional)": "gold", "Caution": "orange", "Not Recommended": "red"}
                        )
                        # Zones / Background
                        fig.add_hrect(y0=0, y1=20, line_width=0, fillcolor="green", opacity=0.05, annotation_text="Excellent")
                        fig.add_hrect(y0=20, y1=40, line_width=0, fillcolor="yellow", opacity=0.05, annotation_text="Good")
                        fig.add_hrect(y0=40, y1=60, line_width=0, fillcolor="orange", opacity=0.05, annotation_text="Permissible")
                        fig.add_hrect(y0=60, y1=80, line_width=0, fillcolor="red", opacity=0.05, annotation_text="Doubtful")
                        fig.add_hrect(y0=80, y1=100, line_width=0, fillcolor="purple", opacity=0.05, annotation_text="Unsuitable")
                        
                        fig.update_layout(yaxis_range=[0, 100], xaxis_title="SAR", yaxis_title="Na %")
                        st.plotly_chart(fig, use_container_width=True)
                        
                    with mid_right:
                        st.markdown(t("classification_table"))
                        # ID, SAR, C class, S class, EC, Na%, TDS (və unit), “Overall suitability”, “Reason”
                        view_cols = ['ID', 'SAR', 'C_Class', 'S_Class', 'EC', 'Na %', 'TDS (mg/L)', 'Overall Suitability', 'Reason']
                        # Reordering columns
                        final_view = res_df[view_cols].copy()
                        # Format floats
                        final_view['SAR'] = final_view['SAR'].map('{:.2f}'.format)
                        final_view['EC'] = final_view['EC'].map('{:.0f}'.format)
                        final_view['Na %'] = final_view['Na %'].map('{:.1f}'.format)
                        if 'TDS (mg/L)' in final_view.columns:
                             final_view['TDS (mg/L)'] = final_view['TDS (mg/L)'].apply(lambda x: f"{x:.0f}" if pd.notnull(x) else "")

                        st.dataframe(final_view, hide_index=True, use_container_width=True)
                        
                        # Actions
                        st.markdown(t("recommended_actions"))
                        # Auto-generate recommendations from dominant classes or worst case in dataset
                        recs = []
                        if 'C3' in dom_c or 'C4' in dom_c:
                            recs.append(t("drainage_improvement"))
                            recs.append(t("salt_tolerant"))
                        if 'S3' in dom_s or 'S4' in dom_s:
                             recs.append(t("gypsum_app"))
                        if not recs:
                             recs.append(t("monitor_irrigation"))
                             
                        for r in recs:
                            st.success(f"✅ {r}")

                    # BOTTOM ROW: Assumptions (Collapsible)
                    with st.expander(t("assumptions_standards")):
                        st.markdown("""
                        **Sources & Standards:**
                        *   **C-Class (Salinity):** USDA Richards (1954). C1 (<250), C2 (250-750), C3 (750-2250), C4 (>2250) µS/cm.
                        *   **S-Class (Sodium):** SAR values. S1 (<10), S2 (10-18), S3 (18-26), S4 (>26).
                        *   **TDS/EC:** Assuming TDS (mg/L) $\\approx$ 0.64 * EC (or EC = TDS * 2 approx depending on salinity). Current calc: EC = TDS * 2.
                        *   **Suitability:**
                            *   **OK:** C1/C2, S1/S2.
                            *   **Caution:** C3 or S3.
                            *   **Not Recommended:** C4 or S4.
                        """)

                else: st.error(t("na_required"))

            # --- Pumping test (Cooper-Jacob) ---
            with tabs[5]:
                st.subheader(t("cooper_title"))
                st.caption(t("cooper_desc"))
                
                c1, c2 = st.columns([1, 2], gap="large")
                with c1:
                    st.markdown(t("input_params"))
                    Q = st.number_input(f"{t('discharge')} (m³/gün)", value=100.0)
                    r_dist = st.number_input("Müşahidə quyusu məsafəsi (r) (m)", value=10.0, min_value=0.1, help="Distance to observation well")
                    
                    st.markdown(t("time_drawdown"))
                    # Enhanced default data
                    pt_data = pd.DataFrame({
                        'Time (min)': [1, 2, 5, 10, 20, 50, 100, 150, 200], 
                        'Drawdown (m)': [0.5, 0.8, 1.5, 2.1, 2.9, 4.5, 6.2, 7.1, 7.8]
                    })
                    pt_df = st.data_editor(pt_data, num_rows="dynamic", key="cooper_editor")
                    
                with c2:
                    if len(pt_df) > 2:
                        try:
                            # 1. Parse Data
                            time_vals = pd.to_numeric(pt_df.iloc[:,0], errors='coerce').dropna().values
                            dd_vals = pd.to_numeric(pt_df.iloc[:,1], errors='coerce').dropna().values
                            
                            if len(time_vals) >= 3:
                                min_t, max_t = float(np.min(time_vals)), float(np.max(time_vals))
                                
                                st.markdown("##### " + t("results_header"))
                                
                                # 2. Interactive Fit Selection
                                st.info(t("select_linear_section"))
                                t_range = st.slider(
                                    t("analysis_range"), 
                                    min_value=min_t, max_value=max_t, value=(min_t, max_t),
                                    key="cj_range_slider"
                                )
                                
                                # 3. Filter Data
                                mask = (time_vals >= t_range[0]) & (time_vals <= t_range[1])
                                t_fit, s_fit = time_vals[mask], dd_vals[mask]
                                
                                if len(t_fit) < 3:
                                    st.warning(t("min_3_points"))
                                else:
                                    # 4. Calculation
                                    log_t = np.log10(t_fit)
                                    slope, intercept = np.polyfit(log_t, s_fit, 1)
                                    
                                    # Transmissivity (T) - Q is m3/day -> T is m2/day
                                    T = (0.183 * Q) / abs(slope)
                                    
                                    # Storativity (S)
                                    # t0 calculation (time where s=0)
                                    t0_min = 10 ** (-intercept / slope)
                                    t0_day = t0_min / 1440.0
                                    S = (2.25 * T * t0_day) / (r_dist ** 2)
                                    
                                    # R2 Fit Quality
                                    s_pred = slope * log_t + intercept
                                    ss_res = np.sum((s_fit - s_pred) ** 2)
                                    ss_tot = np.sum((s_fit - np.mean(s_fit)) ** 2)
                                    r2 = 1 - (ss_res / ss_tot)
                                    
                                    # Display Metrics
                                    m1, m2, m3, m4 = st.columns(4)
                                    m1.metric("Transmissivity (T)", f"{T:.2f} {t('day_unit')}")
                                    m2.metric("Storativity (S)", f"{S:.2e}")
                                    m3.metric("t₀ (min)", f"{t0_min:.2f}")
                                    m4.metric("R² (Fit)", f"{r2:.4f}")
                                    
                                    if r2 < 0.9:
                                        st.warning("⚠️ Diqqət: Uyğunluq (R²) aşağıdır (<0.9). Aralığı dəqiqləşdirin.")
                                    
                                    # 5. Plotting
                                    fig = go.Figure()
                                    
                                    # All Data
                                    fig.add_trace(go.Scatter(
                                        x=time_vals, y=dd_vals, mode='markers', name=t('measured_all'),
                                        marker=dict(color='gray', size=8, opacity=0.5)
                                    ))
                                    
                                    # Fitted Data
                                    fig.add_trace(go.Scatter(
                                        x=t_fit, y=s_fit, mode='markers', name=t('fitted_selected'),
                                        marker=dict(color='blue', size=10, symbol='circle-open', line=dict(width=2))
                                    ))
                                    
                                    # Trend Line (Full Range Projection)
                                    x_proj = np.linspace(min_t, max_t, 100)
                                    y_proj = slope * np.log10(x_proj) + intercept
                                    fig.add_trace(go.Scatter(
                                        x=x_proj, y=y_proj, mode='lines', name=f'Trend (T={T:.1f})',
                                        line=dict(color='red', dash='dash')
                                    ))
                                    
                                    fig.update_xaxes(type="log", title=t('time_log_axis'), dtick="D1")
                                    fig.update_yaxes(title=t('drawdown_axis'))
                                    fig.update_layout(
                                        title=t("time_drawdown_plot"),
                                        margin=dict(l=20, r=20, t=30, b=20),
                                        legend=dict(orientation="h", y=1.1)
                                    )
                                    st.plotly_chart(fig, use_container_width=True)
                                    
                        except Exception as e:
                            st.error(t("calc_error").format(e))
            # --- Lithology section ---
            with tabs[6]:
                # Initialize Session State for Lithology if not present
                if 'lithology_df' not in st.session_state:
                    # Try to migrate from old 'lit' state if exists
                    if 'lit' in st.session_state and not st.session_state['lit'].empty:
                         st.session_state['lithology_df'] = st.session_state['lit'].rename(columns={'Rock': 'Lithology'})
                         # Ensure required columns
                         for col in ['Description', 'Source', 'Confidence', 'ID', 'Color']:
                             if col not in st.session_state['lithology_df'].columns:
                                 st.session_state['lithology_df'][col] = None
                    else:
                        st.session_state['lithology_df'] = pd.DataFrame(columns=[
                            'From', 'To', 'Lithology', 'Description', 'Source', 'Confidence', 'ID', 'Color'
                        ])
                
                if 'lithology_preview' not in st.session_state:
                    st.session_state['lithology_preview'] = None
                
                # --- LAYOUT SETUP ---
                st.subheader(f"{t('tabs_lith')} {t('manager')}")
                
                main_col1, main_col2 = st.columns([35, 65])
                
                # --- LEFT COLUMN: CONTROLS & INPUT ---
                with main_col1:
                    st.markdown(f"### 1. {t('lith_input')}")
                    well_id = st.text_input(t('well_id'), "Q-1", key="lith_well_id")
                    
                    subtabs = st.tabs([t('manual_entry'), t('ai_parse'), t('import')])
                    
                    # TAB 1: MANUAL
                    with subtabs[0]:
                        with st.form("lith_manual_add"):
                            c_f, c_t = st.columns(2)
                            m_from = c_f.number_input(t('from_m'), min_value=0.0, step=0.5)
                            m_to = c_t.number_input(t('to_m'), min_value=0.0, step=0.5)
                            m_lith = st.selectbox(t('lithology'), list(LITHOLOGY_STYLES.keys()))
                            m_desc = st.text_input(t('desc'), "")
                            m_add = st.form_submit_button(t('add_layer'))
                            
                            if m_add:
                                if m_to <= m_from:
                                    st.error("Error: 'To' <= 'From'")
                                else:
                                    pattern_info = LITHOLOGY_STYLES.get(m_lith, LITHOLOGY_STYLES["Unknown"])
                                    new_row = {
                                        'From': m_from, 'To': m_to, 'Lithology': m_lith, 
                                        'Description': m_desc, 'Source': 'Manual', 
                                        'Confidence': 1.0, 'ID': f"L{len(st.session_state['lithology_df'])+1}",
                                        'Color': pattern_info['color']
                                    }
                                    st.session_state['lithology_df'] = pd.concat([
                                        st.session_state['lithology_df'], 
                                        pd.DataFrame([new_row])
                                    ], ignore_index=True)
                                    st.success(t('success'))

                    # TAB 2: AI PARSE
                    with subtabs[1]:
                        st.info(t('paste_log_hint'))
                        ai_txt = st.text_area(t('log_text'), height=150, placeholder="Example: 0-2m clay (brown), 2-10m fine sand, 10-15m limestone...")
                        
                        if st.button(t('parse_ai')):
                            with st.spinner(t('analyzing')):
                                parsed_df = ai_parse_lithology(ai_txt)
                                if parsed_df is not None:
                                    st.session_state['lithology_preview'] = parsed_df
                        
                        if st.session_state['lithology_preview'] is not None:
                            st.write(t('preview'))
                            st.dataframe(st.session_state['lithology_preview'], hide_index=True)
                            
                            c_p1, c_p2 = st.columns(2)
                            if c_p1.button(t('apply_log')):
                                # Set color for parsed items
                                df_to_add = st.session_state['lithology_preview'].copy()
                                df_to_add['Color'] = df_to_add['Lithology'].apply(lambda x: get_lith_style(x)['color'])
                                
                                st.session_state['lithology_df'] = pd.concat([st.session_state['lithology_df'], df_to_add], ignore_index=True)
                                st.session_state['lithology_preview'] = None # Clear preview
                                st.success(t('success'))
                                st.rerun()
                                
                            if c_p2.button(t('discard')):
                                st.session_state['lithology_preview'] = None
                                st.rerun()

                    # TAB 3: IMPORT
                    with subtabs[2]:
                        uploaded_lith = st.file_uploader(t('upload_csv'), type=["csv"])
                        if uploaded_lith:
                            if st.button(t('load_csv')):
                                try:
                                    imp_df = pd.read_csv(uploaded_lith)
                                    # Normalize columns
                                    if 'Rock' in imp_df.columns and 'Lithology' not in imp_df.columns:
                                        imp_df = imp_df.rename(columns={'Rock': 'Lithology'})
                                    
                                    if 'Lithology' not in imp_df.columns:
                                         st.warning("CSV missing 'Lithology' column. Created default.")
                                         imp_df['Lithology'] = "Unknown"
                                         
                                    st.session_state['lithology_df'] = imp_df
                                    st.success(t('success'))
                                except Exception as e:
                                    st.error(f"Error: {e}")

                    st.markdown("---")
                    st.markdown(f"### 2. {t('qc_edit')}")
                    
                    # Validation
                    issues = validate_lithology_log(st.session_state['lithology_df'])
                    if issues:
                        with st.expander(f"⚠️ {len(issues)} {t('issues_found')}", expanded=True):
                            for i in issues: st.warning(i)
                            if st.button(t('auto_fix')):
                                st.session_state['lithology_df'] = st.session_state['lithology_df'].sort_values("From").reset_index(drop=True)
                                st.rerun()
                    else:
                        st.success(t('log_ok'))

                    # Editor
                    edited_df = st.data_editor(
                        st.session_state['lithology_df'],
                        num_rows="dynamic",
                        column_config={
                            "Lithology": st.column_config.SelectboxColumn(
                                "Lithology",
                                options=list(LITHOLOGY_STYLES.keys()),
                                required=True
                            ),
                            "Color": st.column_config.TextColumn("Color", help="Hex color code (e.g. #FF0000)")
                        },
                        hide_index=True,
                        key="lith_editor"
                    )
                    # Sync back edits
                    if not edited_df.equals(st.session_state['lithology_df']):
                        st.session_state['lithology_df'] = edited_df

                # --- RIGHT COLUMN: VISUALIZATION ---
                with main_col2:
                    st.markdown(f"### 3. {t('strip_log_vis')}: **{well_id}**")
                    
                    if not st.session_state['lithology_df'].empty:
                        plot_df = st.session_state['lithology_df'].copy()
                        
                        # Ensure 'Lithology' column exists
                        if 'Lithology' not in plot_df.columns:
                            # Try validation or fallback
                            if 'Rock' in plot_df.columns:
                                plot_df = plot_df.rename(columns={'Rock': 'Lithology'})
                            else:
                                plot_df['Lithology'] = "Unknown"
                        
                        # Ensure numeric
                        plot_df['From'] = pd.to_numeric(plot_df['From'], errors='coerce')
                        plot_df['To'] = pd.to_numeric(plot_df['To'], errors='coerce')
                        plot_df = plot_df.dropna(subset=['From', 'To'])
                        
                        plot_df['Thickness'] = plot_df['To'] - plot_df['From']
                        
                        # Fix colors if missing
                        if 'Color' not in plot_df.columns:
                            plot_df['Color'] = plot_df['Lithology'].apply(lambda x: get_lith_style(x)['color'])
                        else:
                             # fill nan colors
                             plot_df['Color'] = plot_df.apply(lambda r: r['Color'] if pd.notna(r['Color']) else get_lith_style(r['Lithology'])['color'], axis=1)

                        # Sort for drawing
                        plot_df = plot_df.sort_values("From")
                        
                        # Create Strip Log
                        fig = go.Figure()
                        
                        for idx, row in plot_df.iterrows():
                            # Pattern/Hatch
                            style = get_lith_style(row['Lithology'])
                            pattern = style.get("pattern", "")
                            
                            fig.add_trace(go.Bar(
                                x=[1], # Fixed width 
                                y=[row['Thickness']],
                                base=[row['From']],
                                orientation='v',
                                name=row['Lithology'],
                                marker=dict(
                                    color=row['Color'],
                                    line=dict(width=1, color='black'),
                                    pattern=dict(shape=pattern) if pattern else None
                                ),
                                text=f"{row['Lithology']} ({row['From']}-{row['To']}m)",
                                hoverinfo="text+name",
                                hovertext=f"<b>{row['Lithology']}</b><br>Depth: {row['From']} - {row['To']}m<br>Desc: {row.get('Description','')}<br>Conf: {row.get('Confidence',1.0)}"
                            ))
                            
                            # Add text annotation
                            fig.add_annotation(
                                x=0.5, y=row['From'] + (row['Thickness']/2), # Center of bar
                                text=str(row['Lithology']),
                                showarrow=False,
                                font=dict(size=12, color="black"),
                                xref="x", yref="y"
                            )

                        fig.update_layout(
                            yaxis=dict(
                                title="Depth (m)",
                                autorange="reversed", # DEPTH INCREASES DOWN
                                zeroline=False
                            ),
                            xaxis=dict(
                                showticklabels=False,
                                title=f"Well: {well_id}",
                                range=[0, 1]
                            ),
                            barmode='stack', 
                            height=800,
                            showlegend=False,
                            margin=dict(l=50, r=50, t=30, b=30),
                            plot_bgcolor='white',
                            shapes=[
                                dict(
                                    type="rect",
                                    xref="paper", yref="paper",
                                    x0=0, y0=0, x1=1, y1=1,
                                    line=dict(color="black", width=2)
                                )
                            ]
                        )
                        st.plotly_chart(fig, use_container_width=True)
                        
                        # Helper Stats
                        st.markdown("#### Statistics")
                        if not plot_df.empty:
                            total_depth = plot_df['To'].max()
                            st.metric("Total Depth", f"{total_depth} m")
                            
                            stats_df = plot_df.groupby("Lithology")['Thickness'].sum().reset_index()
                            stats_df['%'] = (stats_df['Thickness'] / total_depth * 100).round(1)
                            c_s1, c_s2 = st.columns(2)
                            with c_s1: st.dataframe(stats_df, hide_index=True)
                            with c_s2: 
                                fig_pie = px.pie(stats_df, values='Thickness', names='Lithology', title="Composition")
                                st.plotly_chart(fig_pie, use_container_width=True)
                            
                            st.markdown("---")
                            # --- 4. WELL DESIGN & QC ---
                            st.subheader(t('design_title'))
                            
                            metrics, risks, design = analyze_well_design(plot_df)
                            
                            # Metrics
                            d1, d2, d3 = st.columns(3)
                            d1.metric(t('clay_fraction'), f"{metrics.get('clay_pct',0):.1f}%")
                            d2.metric(t('sand_fraction'), f"{metrics.get('sand_pct',0):.1f}%")
                            d3.metric("Max Depth", f"{metrics.get('total_depth',0)} m")
                            
                            # Risks
                            if risks:
                                with st.expander(f"🚩 {t('risk_flags')} ({len(risks)})", expanded=False):
                                    for r in risks: st.warning(r)
                            
                            # Design Table
                            st.markdown(f"#### {t('design_header')}")
                            if design:
                                d_df = pd.DataFrame(design)
                                st.dataframe(
                                    d_df[['from', 'to', 'lith', 'type', 'slot', 'pack']], 
                                    column_config={
                                        "from": st.column_config.NumberColumn("From (m)", format="%.1f"),
                                        "to": st.column_config.NumberColumn("To (m)", format="%.1f"),
                                        "lith": "Lithology",
                                        "type": "Type",
                                        "slot": t('slot_size'),
                                        "pack": t('filter_pack')
                                    },
                                    use_container_width=True,
                                    hide_index=True
                                )
                                st.caption(t('grain_warn'))
                                
                                # Report Text Generation
                                if st.button(t('design_report_btn')):
                                    report_text = f"ANALYSIS REPORT - {well_id}\n"
                                    report_text += "="*30 + "\n\n"
                                    report_text += f"Total Depth: {metrics['total_depth']} m\n"
                                    report_text += f"{t('clay_fraction')}: {metrics['clay_pct']:.1f}%\n"
                                    report_text += f"{t('sand_fraction')}: {metrics['sand_pct']:.1f}%\n\n"
                                    
                                    report_text += "DETAILED DESIGN:\n"
                                    for item in design:
                                        report_text += f"• {item['from']} - {item['to']}m ({item['lith']}): {item['type']}\n"
                                        report_text += f"  Spec: {item['slot']} | {item['pack']}\n"
                                    
                                    st.text_area("Report Output", report_text, height=300)

                    else:
                        st.info("No lithology data. Use the Input panel to add layers.")
            # --- Darcy section (hydraulic velocity) ---
            with tabs[7]:
                st.subheader(f"💧 {t('darcy_title')}")
                st.caption(t('darcy_caption'))
                
                col1, col2 = st.columns([1, 1], gap="large")
                
                with col1:
                    st.subheader(f"⚙️ {t('input_params')}")
                    k_val = st.number_input(f"{t('darcy_k')} [m/day]", value=5.0, min_value=0.0, format="%.4f", help="Hydraulic Conductivity")
                    if k_val <= 0: st.warning("⚠️ K > 0 recommended.")

                    i_val = st.number_input(f"{t('darcy_i')} [m/m]", value=0.01, min_value=0.0, format="%.5f", step=0.001)
                    if i_val > 0.1: st.warning("⚠️ High gradient (>0.1).")

                    n_pct = st.slider(f"{t('darcy_n')} [%]", min_value=1, max_value=60, value=25)
                    n_val = n_pct / 100.0
                    st.caption(f"ℹ️ {n_pct}% = {n_val:.2f} (fraction)")
                    if n_val > 0.5: st.warning("⚠️ Possibly too high for ne.")

                    # Advanced Params
                    with st.expander("🛠️ Advanced (Pro)"):
                         retardation = st.number_input("Retardation Factor (R)", min_value=1.0, value=1.0, step=0.1, help="Delays solute transport.")
                         st.caption("R=1: Conservative tracer.")
                         alpha_l = st.number_input("Longitudinal Dispersivity (αL) [m]", min_value=0.0, value=10.0, step=1.0, help="Spreading of plume")

                with col2:
                    with st.expander(f"📐 {t('formulas')}", expanded=False):
                        st.latex(r"q = K \cdot i")
                        st.caption(f"q - {t('darcy_v')}")
                        st.latex(r"v = \frac{q}{n_e}")
                        st.caption(f"v - {t('darcy_vs')}")
                    
                    # Calculations
                    q = k_val * i_val
                    v = q / n_val if n_val > 0 else 0
                    v_solute = v / retardation

                    # Interpretation Card
                    st.markdown("### 📊 Interpretation")
                    st.write(f"**{t('darcy_v')}:** `{q:.4f} m/day`")
                    st.write(f"**{t('darcy_vs')}:** `{v:.4f} m/day`")
                    if retardation > 1.0:
                        st.write(f"**Solute v:** `{v_solute:.4f} m/day`")
                    
                    st.info(f"Annual: **{v_solute * 365:.2f} m/yr**")

                st.divider()
                
                # Travel Time
                st.subheader(f"⏱️ {t('travel_time_title')}")
                dist = st.number_input(f"{t('distance')} (m)", value=100.0, step=10.0)
                
                if dist <= 0:
                     st.error("Distance > 0 required.")
                elif v_solute > 0:
                    days = dist / v_solute
                    years = days / 365.25
                    st.success(f"{t('travel_time_res')} **{days:.1f} {t('days')}** (~{years:.2f} {t('years')})")
                    st.progress(min(1.0, 100/days) if days > 100 else 1.0)
                else:
                    st.warning("Velocity is 0.")
            # --- OCR ---
            with tabs[8]:
                st.subheader(f"📷 {t('ocr_title')}")
                
                # Init Session State
                if 'ocr_mode' not in st.session_state: st.session_state['ocr_mode'] = "Hydrochemical (Lab Sheet)"
                if 'ocr_data' not in st.session_state: st.session_state['ocr_data'] = None
                if 'ocr_geotech_json' not in st.session_state: st.session_state['ocr_geotech_json'] = None
                if 'ocr_file_content' not in st.session_state: st.session_state['ocr_file_content'] = None
                if 'ocr_mime_type' not in st.session_state: st.session_state['ocr_mime_type'] = None
                if 'ocr_report' not in st.session_state: st.session_state['ocr_report'] = None
                if 'geo_interpretation' not in st.session_state: st.session_state['geo_interpretation'] = None

                ocr_col1, ocr_col2 = st.columns([1, 1], gap="large")
                
                with ocr_col1:
                    st.info(t('ocr_info'))
                    # 1. Mode Selection
                    opts = [t("ocr_mode_hydro"), t("ocr_mode_geo")]
                    curr = st.session_state.get('ocr_mode', opts[0])
                    idx = opts.index(curr) if curr in opts else 0
                    
                    ocr_mode = st.radio("Document Type", opts, index=idx)
                    st.session_state['ocr_mode'] = ocr_mode
                    
                    u = st.file_uploader(t('upload_file'), ["png", "jpg", "jpeg", "pdf"], key="ocr_unique_uploader")
                    
                    if not st.session_state.get("ai_accept", False):
                        st.warning(t('ai_warning'))
                    
                    analyze_btn = st.button(t("analyze"), type="primary", disabled=not u)

                with ocr_col2:
                    with st.expander(f"❓ {t('ocr_how')}", expanded=True):
                         if ocr_mode == "Borehole Log / Geotechnical (Beta)":
                             st.markdown("""
                             **Geotechnical Parser:**
                             1. Parses Borehole Logs (Scan/PDF).
                             2. Extracts Stratigraphy & SPT.
                             3. Generates Visual Log & Report.
                             """)
                         else:
                             st.markdown(t('ocr_steps'))

                # Process Upload
                if u and analyze_btn:
                    with st.spinner(t('ocr_waiting')):
                        u.seek(0)
                        file_bytes = u.read()
                        mime = u.type
                        st.session_state['ocr_file_content'] = file_bytes
                        st.session_state['ocr_mime_type'] = mime
                        st.session_state['ocr_report'] = None
                        st.session_state['geo_interpretation'] = None
                        
                        if ocr_mode == "Borehole Log / Geotechnical (Beta)":
                            res = ai_parse_geotech_log(file_bytes, mime)
                            st.session_state['ocr_geotech_json'] = res
                            st.session_state['ocr_data'] = None # Reset other mode
                        else:
                            d = ai_read_image(file_bytes, mime)
                            st.session_state['ocr_data'] = d
                            st.session_state['ocr_geotech_json'] = None

                # --- MODE 1: HYDROCHEMICAL ---
                if st.session_state['ocr_mode'] == t("ocr_mode_hydro"):
                    if st.session_state['ocr_data'] is not None and not st.session_state['ocr_data'].empty: 
                        st.success(f"✅ {t('success')} ({len(st.session_state['ocr_data'])} rows)")
                        
                        # Editable result
                        st.markdown(f"### 📝 {t('ocr_result')}")
                        edited_ocr = st.data_editor(st.session_state['ocr_data'], num_rows="dynamic", use_container_width=True, key="ocr_results_editor")
                        
                        # Generate Automatic Report
                        st.markdown("---")
                        st.subheader(f"🧠 {t('ocr_ai_report')}")
                        
                        if st.button("📝 Generate Detailed Report (Multi-Modal)", type="primary"):
                            with st.spinner(t('report_writing')):
                                data_summary = edited_ocr.to_string()
                                f_bytes = st.session_state.get('ocr_file_content')
                                f_mime = st.session_state.get('ocr_mime_type')
                                c_lang = st.session_state.get("lang", "en")
                                
                                report = ai_generate_ocr_detailed_report(f_bytes, f_mime, data_summary, c_lang)
                                st.session_state['ocr_report'] = report
                                
                        if st.session_state['ocr_report']:
                            st.markdown(st.session_state['ocr_report'])
                            st.download_button(
                                label=t('download_word'), 
                                data=create_word_report(st.session_state['ocr_report'], {}), 
                                file_name="HydroMind_OCR_Report.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                    elif st.session_state['ocr_data'] is not None and st.session_state['ocr_data'].empty:
                        st.error(t('no_data'))
                        st.warning(t('ocr_error_hint'))

                # --- MODE 2: GEOTECH ---
                elif st.session_state['ocr_mode'] == t("ocr_mode_geo"):
                    res = st.session_state.get('ocr_geotech_json')
                    if res:
                        if "error" in res:
                             st.error(res["error"])
                        else:
                             st.success("✅ Borehole Log Parsed Successfully")
                             
                             # 1. Metadata Panel
                             with st.expander("📌 1. Header & Quality Control", expanded=True):
                                 meta = res.get("metadata", {})
                                 qc = res.get("qc", {})
                                 
                                 qc1, qc2 = st.columns([1, 3])
                                 qc1.metric("OCR Quality", f"{meta.get('confidence', 0)}%")
                                 if qc.get('issues'): qc2.warning(f"Issues: {', '.join(qc['issues'])}")
                                 
                                 # Editable Meta
                                 new_meta_df = st.data_editor(pd.DataFrame([meta]), key="geo_meta_edit", use_container_width=True)
                                 meta_dict = new_meta_df.iloc[0].to_dict()

                             # 2. Stratigraphy Panel
                             st.markdown("### 🪨 2. Stratigraphy (Layers)")
                             layers = res.get("stratigraphy", [])
                             if layers:
                                 df_layers = pd.DataFrame(layers)
                                 df_layers_edited = st.data_editor(df_layers, num_rows="dynamic", key="geo_layers_edit", use_container_width=True)
                             else:
                                 st.warning("No layers detected.")
                                 df_layers_edited = pd.DataFrame()

                             # 3. SPT Panel
                             st.markdown("### 🔨 3. In-situ Tests (SPT)")
                             spts = res.get("tests", [])
                             if spts:
                                 df_spt = pd.DataFrame(spts)
                                 df_spt_edited = st.data_editor(df_spt, num_rows="dynamic", key="geo_spt_edit", use_container_width=True)
                             else:
                                 st.info("No SPT data found.")
                                 df_spt_edited = pd.DataFrame()

                             # 4. Actions
                             st.divider()
                             st.subheader("📊 4. Analysis & Report")
                             
                             col_a1, col_a2 = st.columns([2, 1])
                             
                             with col_a1:
                                 st.markdown("**Visual Strip Log**")
                                 if not df_layers_edited.empty:
                                     fig_log = render_geotech_visual_log(df_layers_edited, df_spt_edited, meta_dict)
                                     st.plotly_chart(fig_log, use_container_width=True)
                             
                             with col_a2:
                                 st.markdown("**AI Engineering Interpretation**")
                                 if st.button("🧠 Generate Interpretation"):
                                     with st.spinner("Analyzing soil mechanics..."):
                                         # Quick Interpretation prompt
                                         prompt_interp = f"Based on these SPT values: {df_spt_edited.to_dict('records') if not df_spt_edited.empty else 'None'} and Stratigraphy: {df_layers_edited.to_dict('records')}, write a short engineering summary (bearing capacity indications, consistencies)."
                                         client = get_genai_client()
                                         try:
                                             resp = call_gemini_with_retry(client, model="gemini-3-flash-preview", contents=prompt_interp)
                                             st.session_state['geo_interpretation'] = resp.text
                                         except: st.session_state['geo_interpretation'] = "Error generating."

                                 if st.session_state.get('geo_interpretation'):
                                     st.info(st.session_state['geo_interpretation'])
                                     
                                     word_file = create_geotech_report_word(meta_dict, df_layers_edited, df_spt_edited, st.session_state['geo_interpretation'])
                                     st.download_button(
                                         "📄 Download Report (.docx)", 
                                         word_file, 
                                         "Geotech_Log_Report.docx",
                                         "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                     )

                                 # --- Exports (CSV) ---
                                 exports = res.get("exports", {})
                                 if exports:
                                     st.divider()
                                     st.markdown("###### 📥 CSV Exports")
                                     if "lithology_csv_rows" in exports:
                                         ld = exports["lithology_csv_rows"]
                                         if ld:
                                             csv_l = pd.DataFrame(ld).to_csv(index=False).encode('utf-8')
                                             st.download_button("Lithology CSV", csv_l, "lithology.csv", "text/csv", key="dl_lith")
                                     
                                     if "spt_csv_rows" in exports:
                                         sd = exports["spt_csv_rows"]
                                         if sd:
                                             csv_s = pd.DataFrame(sd).to_csv(index=False).encode('utf-8')
                                             st.download_button("SPT CSV", csv_s, "spt.csv", "text/csv", key="dl_spt")

            # --- Proqnoz ---
            with tabs[9]:
                st.subheader(f"🔮 {t('forecast_title')} {t('forecast_suffix')}")
                st.caption(t("forecast_extended_caption"))
                
                fc_col1, fc_col2 = st.columns([1, 2], gap="large")
                
                with fc_col1:
                    st.markdown(f"##### 1. {t('history_data')}")
                    # Default template with 'Year'
                    fc_data = pd.DataFrame({
                        'Year': [2015, 2016, 2017, 2018, 2019, 2020, 2021, 2022, 2023, 2024], 
                        'Value': [10.8, 10.6, 10.5, 10.2, 10.0, 9.8, 9.5, 9.1, 9.0, 8.8]
                    })
                    hdf = st.data_editor(fc_data, num_rows="dynamic", key="forecast_editor")
                    
                    st.markdown("##### Parameters")
                    years_forward = st.slider(t('years_forward'), 1, 10, 5)
                    model_type = st.selectbox("Model Type", ["Linear (Recommended)", "Polynomial (Degree 2)"])
                    
                    # Data Suitability Check
                    n_samples = len(hdf)
                    if n_samples < 5:
                        confidence_level = "Very Low"
                        conf_color = "red"
                    elif n_samples < 10:
                        confidence_level = "Low"
                        conf_color = "orange"
                    elif n_samples < 30:
                        confidence_level = "Medium"
                        conf_color = "blue"
                    else:
                        confidence_level = "High"
                        conf_color = "green"
                        
                    st.markdown(f"**Data Confidence:** :{conf_color}[{confidence_level}] (n={n_samples})")
                    if n_samples < 10:
                        st.caption("⚠️ Less than 10 data points may lead to overfitting.")

                with fc_col2:
                    if len(hdf) >= 3:
                        m_sel = "Linear" if "Linear" in model_type else "Polynomial"
                        res = predict_future(hdf, years_to_predict=years_forward, model_type=m_sel)
                        
                        if res is not None:
                            # Unpack results
                            fy = res["years"]
                            fv = res["values"]
                            metrics = res["metrics"]
                            
                            # Big Metric for next year
                            next_val = fv[0]
                            # Find last actual value
                            last_actual = res["history_y"][-1]
                            delta = next_val - last_actual
                            
                            st.markdown(f"##### 2. {t('results')}")
                            
                            # Metrics Row
                            m1, m2, m3, m4 = st.columns(4)
                            m1.metric(f"{fy[0]} Forecast", f"{next_val:.2f}", delta=f"{delta:.2f}")
                            m2.metric("R² Score", f"{metrics['R2']:.3f}")
                            m3.metric("MAE", f"{metrics['MAE']:.3f}")
                            if metrics['Slope']:
                                m4.metric("Slope (Change/Yr)", f"{metrics['Slope']:.4f}")
                            else:
                                m4.metric("Model", "Polynomial")
                            
                            # Chart
                            fig = go.Figure()
                            
                            # Confidence Intervals (Area)
                            if res["ci_lower"] is not None:
                                fig.add_trace(go.Scatter(
                                    x=np.concatenate([fy, fy[::-1]]),
                                    y=np.concatenate([res["ci_upper"], res["ci_lower"][::-1]]),
                                    fill='toself',
                                    fillcolor='rgba(255, 0, 0, 0.2)',
                                    line=dict(color='rgba(255,255,255,0)'),
                                    hoverinfo="skip",
                                    name='95% Confidence Interval'
                                ))

                            # Observed Data
                            fig.add_trace(go.Scatter(
                                x=res["history_x"], y=res["history_y"], 
                                mode='markers', name='Actual Data',
                                marker=dict(color='blue', size=8)
                            ))
                            
                            # Historical Fit Line
                            fig.add_trace(go.Scatter(
                                x=res["history_x"], y=res["history_pred"],
                                mode='lines', name='Model Fit',
                                line=dict(color='gray', dash='dash')
                            ))

                            # Future Forecast
                            fig.add_trace(go.Scatter(
                                x=fy, y=fv, 
                                mode='lines+markers', name='Forecast',
                                line=dict(color='red', width=3)
                            ))
                            
                            title_text = f"Trend Forecast ({m_sel} Regression)"
                            fig.update_layout(
                                title=title_text, 
                                template="plotly_white",
                                xaxis_title="Year",
                                yaxis_title="Value",
                                hovermode="x"
                            )
                            st.plotly_chart(fig, use_container_width=True)
                            
                        else:
                            st.warning("Sütun adlarını yoxlayın: 'Year'/'Value' və ya 'İl'/'Dəyər' olmalıdır.")
                    else:
                        st.info("Proqnoz üçün ən azı 3 illik məlumat daxil edin.")
            # --- AI Report ---
            with tabs[10]:
                st.subheader(f"📝 {t('tabs_report')}")
                st.caption(t('ai_report_caption'))
                
                c_rep1, c_rep2 = st.columns([1, 2], gap="large")
                with c_rep1:
                    st.markdown("##### ⚙️ Report Settings")
                    rep_proj = st.text_input("Project Name / Location", value="Hydrogeological Assessment")
                    rep_client = st.text_input("Client Name", value="Confidential Client")
                    
                    st.markdown("**Report Type:**")
                    rep_type_choice = st.radio(
                        "Template", 
                        ["A", "B", "C", "Auto"],
                        format_func=lambda x: {
                            "A": "A — Water Quality (Hydrochem)",
                            "B": "B — Irrigation (Agri)",
                            "C": "C — Pumping Test",
                            "Auto": "Auto — Detect from Data"
                        }[x]
                    )

                    st.divider()
                    st.info(t('ai_report_info'))
                    
                    if st.button(t('generate_report'), type="primary", key="btn_gen_rep_new"):
                        with st.spinner(t('report_writing')):
                            try:
                                # Meta Info
                                meta = f"Project: {rep_proj}, Client: {rep_client}, Date: {datetime.now().strftime('%Y-%m-%d')}"
                                
                                # Gather Available Charts/Outputs roughly
                                outputs = []
                                if not df.empty: outputs.append("Dataset Table")
                                if 'Lat' in df.columns: outputs.append("Well Map")
                                # We can't easily valid what charts were actually rendered in other tabs without session state checks
                                # But we can infer possible charts:
                                if 'Ca' in df.columns and 'Cl' in df.columns: outputs.append("Piper Diagram, Schoeller Diagram, Stiff Map")
                                if 'Water Level' in df.columns: outputs.append("Water Level Chart")
                                if 'Lithology' in st.session_state.get('lithology_df', pd.DataFrame()).columns: outputs.append("Lithology Log")
                                
                                # Call V2
                                msg = ai_generate_report_v2(df, rep_type_choice, meta, ", ".join(outputs))
                                st.session_state['lr'] = msg
                                
                            except Exception as e:
                                st.error(f"Error: {e}")
                
                with c_rep2:
                    if st.session_state.get('lr'):
                        st.markdown(f"### 📄 {t('result')}")
                        st.markdown(st.session_state['lr'])
                        
                        try:
                            # Pass empty dict for stats as it's implied in text now
                            docx_data = create_word_report(st.session_state['lr'], {})
                            st.download_button(
                                t('download_word'), 
                                docx_data, 
                                f"Report_{rep_proj.replace(' ', '_')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        except Exception as e:
                            st.error(f"Error creating Word file: {e}")
                    else:
                        st.markdown(t('report_structure'))
                        st.info("👈 Fill details on the left to start.")

            # --- Maps (Advanced GIS) ---
            with tabs[11]:
                st.subheader(f"🗺️ {t('maps_title')}")
                
                if 'Lat' in df.columns and 'Lon' in df.columns:
                    # Layout: Sidebar (Settings) + Main (Map)
                    
                    # We'll use tabs for different GIS modes
                    gis_tabs = st.tabs([t("gis_tab_interpolation"), t("gis_tab_cv"), t("gis_tab_flow")])
                    
                    # SHARED SETTINGS (in a container or columns above?)
                    # Let's put common settings in the first tab's sidebar-like col
                    
                    # ---- TAB 1: INTERPOLATION ----
                    with gis_tabs[0]:
                        row_cfg, row_map = st.columns([1, 3])
                        
                        with row_cfg:
                            st.markdown(t("gis_data_layer"))
                            # Filter numeric columns
                            num_cols = [c for c in df.select_dtypes(include=np.number).columns if c not in ['Lat', 'Lon']]
                            param = st.selectbox(t("parameter"), num_cols, index=0 if num_cols else None)
                            
                            st.markdown(t("gis_method"))
                            method_options = ["IDW", "Kriging (Ordinary)", "Spline / Cubic", "Natural Neighbor / Linear"]
                            method_labels = {
                                "IDW": t("method_idw"),
                                "Kriging (Ordinary)": t("method_kriging"),
                                "Spline / Cubic": t("method_spline"),
                                "Natural Neighbor / Linear": t("method_natural")
                            }
                            method = st.selectbox(
                                t("gis_interpolation_method"),
                                method_options,
                                format_func=lambda x: method_labels.get(x, x)
                            )
                            
                            st.markdown(t("gis_formatting"))
                            cmap = st.selectbox(t("gis_colormap"), ["jet", "viridis", "plasma", "RdBu", "YlGnBu"], index=0)
                            n_levels = st.slider(t("gis_contour_levels"), 5, 30, 15)
                            opacity = st.slider(t("opacity"), 0.0, 1.0, 0.6)
                            
                            st.markdown(t("gis_constraints"))
                            clip_hull = st.checkbox(t("gis_clip_hull"), value=True, help=t("gis_clip_hull_help"))
                            res_factor = st.slider(t("gis_resolution"), 50, 200, 100, help=t("gis_resolution_help"))

                            # Kriging Specifics
                            variogram_model = 'linear'
                            sill, range_v, nugget = None, None, None
                            
                            if "Kriging" in method:
                                with st.expander(t("gis_variogram_settings")):
                                    variogram_model = st.selectbox(t("gis_variogram_model"), ["linear", "power", "gaussian", "spherical"], index=0)
                                    # Optional manual parameters (advanced)
                                    if st.checkbox(t("gis_manual_params")):
                                        sill = st.number_input(t("gis_sill"), value=1.0)
                                        range_v = st.number_input(t("gis_range"), value=10.0)
                                        nugget = st.number_input(t("gis_nugget"), value=0.0)
                                    
                        with row_map:
                            if param:
                                map_df = df.dropna(subset=['Lat', 'Lon', param])
                                if len(map_df) < 4:
                                    st.warning(t("gis_not_enough_points").format(n=len(map_df)))
                                else:
                                    # Create base map
                                    m_center = [map_df['Lat'].mean(), map_df['Lon'].mean()]
                                    
                                    # Perform Interpolation
                                    with st.spinner(t("gis_computing").format(method=method_labels.get(method, method))):
                                        err, fe_json, cm_obj, field_grid, uncertainty_grid, grid_coords = generate_contours_advanced(
                                            map_df, param, n_levels, cmap, method, 
                                            resolution_factor=res_factor,
                                            variogram_model=variogram_model,
                                            sill_val=sill, range_val=range_v, nugget_val=nugget,
                                            clip_hull=clip_hull
                                        )
                                    
                                    # TABS for Map vs Uncertainty
                                    if "Kriging" in method:
                                         view_mode_labels = {
                                             "Interpolated Value": t("gis_layer_interpolated"),
                                             "Uncertainty (Variance)": t("gis_layer_uncertainty")
                                         }
                                         view_mode = st.radio(
                                             t("gis_layer_view"),
                                             ["Interpolated Value", "Uncertainty (Variance)"],
                                             horizontal=True,
                                             format_func=lambda x: view_mode_labels.get(x, x)
                                         )
                                    else:
                                         view_mode = "Interpolated Value"
                                         
                                    m = folium.Map(location=m_center, zoom_start=10)
                                    
                                    # Add Contours
                                    if fe_json:
                                        display_json = fe_json
                                        display_cm = cm_obj
                                        
                                        if view_mode == "Uncertainty (Variance)":
                                            if uncertainty_grid is not None and grid_coords:
                                                try:
                                                    xi = grid_coords['xi']
                                                    yi = grid_coords['yi']
                                                    
                                                    fig_u = plt.figure()
                                                    ax_u = fig_u.add_subplot(111)
                                                    # Plot uncertainty
                                                    cntr_u = ax_u.contourf(xi, yi, uncertainty_grid, levels=10, cmap='Reds', alpha=0.7)
                                                    
                                                    display_json = geojsoncontour.contourf_to_geojson(
                                                        contourf=cntr_u,
                                                        min_angle_deg=3.0,
                                                        ndigits=3,
                                                        stroke_width=2,
                                                        fill_opacity=0.5
                                                    )
                                                    display_cm = cm.LinearColormap(colors=['white', 'red'], vmin=np.nanmin(uncertainty_grid), vmax=np.nanmax(uncertainty_grid))
                                                    plt.close(fig_u)
                                                except Exception as e:
                                                    st.error(t("uncertainty_plot_error").format(error=e))

                                        folium.GeoJson(
                                            display_json,
                                            style_function=lambda x: {
                                                'fillColor': x['properties']['fill'],
                                                'color': 'black',
                                                'weight': 0.5,
                                                'fillOpacity': opacity
                                            },
                                            name=t("gis_layer_name")
                                        ).add_to(m)
                                        
                                        if display_cm: m.add_child(display_cm)
                                    
                                    # Add Points
                                    for _, r in map_df.iterrows():
                                        folium.CircleMarker(
                                            [r['Lat'], r['Lon']],
                                            radius=5, color='black', fill=True, fill_color='white',
                                            tooltip=f"<b>{r.get('ID','?')}</b><br>{param}: {r[param]}"
                                        ).add_to(m)
                                        
                                    st_folium(m, height=600, use_container_width=True)
                                    
                    # ---- TAB 2: CROSS VALIDATION ----
                    with gis_tabs[1]:
                        st.markdown(t("gis_cv_title"))
                        st.caption(t("gis_cv_caption"))
                        
                        col_cv_sets, col_cv_res = st.columns([1, 2])
                        
                        with col_cv_sets:
                             cv_param = st.selectbox(t("gis_cv_param"), num_cols, index=0 if num_cols else None, key="cv_p")
                             cv_method = st.selectbox(
                                 t("gis_cv_method"),
                                 ["IDW", "Kriging (Ordinary)"],
                                 key="cv_m",
                                 format_func=lambda x: method_labels.get(x, x)
                             )
                             if st.button(t("gis_run_validation"), type="primary"):
                                 if cv_param:
                                     with st.spinner(t("gis_cv_running")):
                                         metrics, res_df = perform_cv_analysis(
                                             df.dropna(subset=['Lat','Lon',cv_param]), 
                                             cv_param, cv_method
                                         )
                                         st.session_state['cv_metrics'] = metrics
                                         st.session_state['cv_res'] = res_df
                        
                        with col_cv_res:
                             if 'cv_metrics' in st.session_state and st.session_state['cv_metrics']:
                                 met = st.session_state['cv_metrics']
                                 m1, m2, m3 = st.columns(3)
                                 m1.metric("RMSE", f"{met['RMSE']:.3f}", help=t("gis_rmse_help"))
                                 m2.metric("MAE", f"{met['MAE']:.3f}")
                                 m3.metric("R² Score", f"{met['R2']:.3f}", help=t("gis_r2_help"))
                                 
                                 res_df = st.session_state['cv_res']
                                 
                                 # Scatter Plot Actual vs Predicted
                                 fig_cv = px.scatter(
                                     res_df, x="Actual", y="Predicted", hover_data=["ID", "Residual"],
                                     title=t("gis_cv_plot_title").format(param=cv_param),
                                     trendline="ols"
                                 )
                                 
                                 # 1:1 Ideal Line
                                 min_val = min(res_df["Actual"].min(), res_df["Predicted"].min())
                                 max_val = max(res_df["Actual"].max(), res_df["Predicted"].max())
                                 
                                 fig_cv.add_shape(
                                     type="line", 
                                     x0=min_val, y0=min_val,
                                     x1=max_val, y1=max_val,
                                     line=dict(color="green", dash="dot"),
                                     name=t("gis_ideal_line")
                                 )

                                 try:
                                     results = px.get_trendline_results(fig_cv)
                                     model = results.px_fit_results.iloc[0]
                                     alpha = model.params[0] # Intercept
                                     beta = model.params[1]  # Slope
                                     r_sq = model.rsquared
                                     
                                     fig_cv.add_annotation(
                                         x=min_val, y=max_val,
                                         text=f'y = {beta:.2f}x + {alpha:.2f}<br>R² = {r_sq:.2f}',
                                         showarrow=False,
                                         align="left",
                                         bgcolor="rgba(255, 255, 255, 0.8)",
                                         bordercolor="black"
                                     )

                                     # Calculate Intersection (x = ax + b => x(1-a) = b => x = b / (1-a))
                                     # Avoid division by zero if lines are parallel (beta=1)
                                     if abs(1 - beta) > 1e-5:
                                         intersect_val = alpha / (1 - beta)
                                         
                                         # Add marker for intersection
                                         fig_cv.add_trace(go.Scatter(
                                             x=[intersect_val], 
                                             y=[intersect_val],
                                             mode='markers',
                                             marker=dict(size=12, color='orange', symbol='x'),
                                             name=t("gis_intersection_name")
                                         ))
                                         
                                         # Add annotation pointing to it
                                         fig_cv.add_annotation(
                                             x=intersect_val, y=intersect_val,
                                             text=t("gis_intersection"),
                                             showarrow=True,
                                             arrowhead=1,
                                             ax=20, ay=-20
                                         )

                                 except:
                                     pass

                                 st.plotly_chart(fig_cv, use_container_width=True)
                                 
                                 with st.expander(t("gis_detailed_residuals")):
                                     st.dataframe(res_df, use_container_width=True)
                             else:
                                 st.info(t("gis_select_params_prompt"))

                    # ---- TAB 3: FLOW ----
                    with gis_tabs[2]:
                        st.markdown(t("gis_flow_title"))
                        st.info(t("gis_flow_info"))
                        
                        if 'Water Level' not in df.columns:
                            st.error(t("gis_flow_missing"))
                        else:
                            col_f1, col_f2 = st.columns([1, 3])
                            with col_f1:
                                f_res = st.slider(t("gis_flow_grid_resolution"), 20, 100, 40)
                            
                            with col_f2:
                                wl_df = df.dropna(subset=['Lat', 'Lon', 'Water Level'])
                                if len(wl_df) > 3:
                                    # Calc Grid
                                    x = wl_df['Lon'].values
                                    y = wl_df['Lat'].values
                                    z = wl_df['Water Level'].values
                                    
                                    xi = np.linspace(min(x), max(x), f_res)
                                    yi = np.linspace(min(y), max(y), f_res)
                                    xi, yi = np.meshgrid(xi, yi)
                                    
                                    # Smooth surface first (Cubic splice)
                                    zi = griddata((x, y), z, (xi, yi), method='cubic')
                                    
                                    # Gradient
                                    dx, dy = calculate_flow_direction(zi, xi, yi)
                                    
                                    # Plot
                                    fig_flow = plt.figure(figsize=(10, 8))
                                    ax = fig_flow.add_subplot(111)
                                    
                                    # Contour Water Level
                                    cs = ax.contour(xi, yi, zi, levels=10, colors='blue', alpha=0.5)
                                    ax.clabel(cs, inline=1, fontsize=8)
                                    
                                    # Quiver (Flow is negative gradient usually, high to low)
                                    # dx, dy is gradient (uphill). Flow is downhill (-dx, -dy).
                                    # Skip every nth point for clarity
                                    skip = 2
                                    ax.quiver(xi[::skip, ::skip], yi[::skip, ::skip], -dx[::skip, ::skip], -dy[::skip, ::skip], color='red')
                                    
                                    ax.scatter(x, y, c='black', s=20, label='Wells')
                                    ax.set_title("Groundwater Flow Field (Approximated)")
                                    ax.set_xlabel("Longitude")
                                    ax.set_ylabel("Latitude")
                                    
                                    st.pyplot(fig_flow)
                                else:
                                    st.warning(t("gis_not_enough_data"))
                else:
                    st.warning(t('lat_lon_required'))

            # --- Save ---
            with tabs[12]:
                st.subheader(f"💾 {t('tabs_save')}")
                c_save1, c_save2 = st.columns([2, 1])
                with c_save1:
                    n = st.text_input(t("project_name_lbl"), f"Project_{datetime.now().strftime('%Y%m%d_%H%M')}")
                    if st.button(t("save_project_btn"), type="primary"): 
                        save_project(st.session_state['username'], n, "Auto", df)
                        st.success(f"{t('project_saved')} '{n}'")
                with c_save2:
                    st.info(t('save_info'))

            # --- WQI ---
            with tabs[13]:
                st.subheader(f"💧 {t('wqi_title')}")
                st.caption("Enhanced Water Quality Index with Diagnostics")
                
                # 1. Calculation
                wqi_df, weights, used_standards = calculate_wqi(df)
                
                if wqi_df.empty:
                    st.warning("Could not calculate WQI. Please check your data columns matches standard parameters (pH, TDS, Ca, Mg, etc.).")
                else:
                    combined = pd.concat([df.reset_index(drop=True), wqi_df], axis=1)
                    
                    # 2. Methodology & Standards (Glass Box)
                    with st.expander("📝 Methodology, Formula & Standards", expanded=False):
                        st.markdown(r"""
                        **Weighted Arithmetic Water Quality Index (WQI)**
                        
                        Formula:
                        $$
                        WQI = \frac{\sum (W_i \times q_i)}{\sum W_i}
                        $$
                        *   $W_i$ is the unit weight for parameter $i$.
                        *   $q_i$ is the quality rating scale for parameter $i$.
                        
                        **Classification Scale:**
                        *   **0 - 25**: Excellent
                        *   **26 - 50**: Good
                        *   **51 - 70**: Poor
                        *   **71 - 90**: Very Poor
                        *   **> 90**: Unsuitable
                        """)
                        
                        st.write("**Standards ($S_i$) & Relative Unit Weights ($W_i$) Used:**")
                        std_data = []
                        for p, s in used_standards.items():
                            std_data.append({
                                "Parameter": p, 
                                "Standard (Si)": s, 
                                "Weight (Wi)": f"{weights.get(p,0):.4f}"
                            })
                        st.dataframe(pd.DataFrame(std_data), use_container_width=True)

                    # 3. Overview Metrics
                    avg_wqi = combined['WQI'].mean()
                    
                    # Determine main driver for the dataset
                    contrib_cols = [c for c in combined.columns if c.endswith('_contribution_pct')]
                    if contrib_cols:
                        avg_contrib = combined[contrib_cols].mean().sort_values(ascending=False)
                        top_driver = avg_contrib.index[0].replace('_contribution_pct', '')
                        top_driver_pct = avg_contrib.iloc[0]
                        driver_text = f"Primarily driven by **{top_driver}** ({top_driver_pct:.1f}%)"
                    else:
                        driver_text = "Drivers not analyzed"

                    m1, m2 = st.columns(2)
                    m1.metric(t('avg_wqi'), f"{avg_wqi:.1f}")
                    m2.info(f"💡 Automatic Summary: Overall quality average is **{avg_wqi:.1f}**. {driver_text}.")

                    # 4. Main Analysis (Split View)
                    col_main, col_drill = st.columns([2, 1])
                    
                    # Prepare sorted data for chart
                    combined_sorted = combined.sort_values('WQI', ascending=True)

                    with col_main:
                        st.markdown("#### 📊 WQI Ranking")
                        fig = px.bar(
                            combined_sorted, x='ID', y='WQI', 
                            color='Class', 
                            title="Water Quality Index by Sample (Sorted)",
                            color_discrete_map={
                                "Excellent": "#00CC96", "Good": "#636EFA", 
                                "Poor": "#EF553B", "Very Poor": "#AB63FA", "Unsuitable": "#19D3F3"
                            },
                            hover_data=['WQI', 'Class']
                        )
                        fig.add_hline(y=50, line_dash="dash", line_color="green", annotation_text="Good Limit")
                        fig.add_hline(y=100, line_dash="dash", line_color="red", annotation_text="Limit")
                        st.plotly_chart(fig, use_container_width=True)
                        
                        # Spatial Map
                        st.markdown("#### 🗺️ Spatial WQI Map")
                        if 'Lat' in combined_sorted.columns and 'Lon' in combined_sorted.columns:
                            # Drop NaNs for map
                            map_df = combined_sorted.dropna(subset=['Lat', 'Lon'])
                            if not map_df.empty:
                                m = folium.Map(location=[map_df['Lat'].mean(), map_df['Lon'].mean()], zoom_start=10)
                                color_map = {
                                    "Excellent": "green", "Good": "blue", 
                                    "Poor": "orange", "Very Poor": "purple", "Unsuitable": "red"
                                }
                                for _, row in map_df.iterrows():
                                    c = color_map.get(row['Class'], 'black')
                                    folium.CircleMarker(
                                        location=[row['Lat'], row['Lon']],
                                        radius=8,
                                        popup=folium.Popup(f"<b>ID: {row['ID']}</b><br>WQI: {row['WQI']:.1f}<br>Class: {row['Class']}", max_width=300),
                                        color='white',
                                        weight=1,
                                        fill=True,
                                        fill_color=c,
                                        fill_opacity=0.8
                                    ).add_to(m)
                                st_folium(m, height=400, use_container_width=True)
                            else:
                                st.warning("Lat/Lon data available but contains NaNs.")
                        else:
                            st.info("Lat/Lon columns not found. Map disabled.")

                    with col_drill:
                        st.markdown("#### 🔍 Failure Diagnostics")
                        st.caption("Select a sample to see why it has that WQI.")
                        
                        selected_id_wqi = st.selectbox("Select Sample ID", combined_sorted['ID'].unique())
                        
                        if selected_id_wqi:
                            sel_row = combined[combined['ID'] == selected_id_wqi].iloc[0]
                            st.metric("Sample WQI", f"{sel_row['WQI']:.1f}", sel_row['Class'])
                            
                            # Prepare contributors
                            c_data = {}
                            for col in combined.columns:
                                if col.endswith('_contribution_pct'):
                                    p_name = col.replace('_contribution_pct', '')
                                    c_data[p_name] = sel_row[col]
                            
                            c_df = pd.DataFrame(list(c_data.items()), columns=['Parameter', 'Contribution %'])
                            c_df = c_df.sort_values('Contribution %', ascending=True) 
                            
                            fig_d = px.bar(
                                c_df, x='Contribution %', y='Parameter', orientation='h',
                                title=f"Drivers for {selected_id_wqi}",
                                text_auto='.1f'
                            )
                            st.plotly_chart(fig_d, use_container_width=True)

                    with st.expander(t('detailed_table')):
                         st.dataframe(combined, use_container_width=True)
            # --- Statistics ---
            with tabs[14]:
                st.subheader(f"📊 {t('tabs_stats')}")
                st.caption(t('stats_caption'))
                
                with st.expander(f"ℹ️ {t('what_is_this')}", expanded=False):
                    st.markdown(t('stats_explainer'))

                selected_cols = [c for c in active_cols if c in df.columns]
                
                if len(selected_cols) == 0:
                    st.warning(t('select_param_warning'))
                else:
                    # --- 1. Top Summary Cards ---
                    stats, corr = compute_statistics(df, selected_cols)
                    
                    st.markdown("#### 1. Top Summary Cards")
                    
                    # Calculate Hardness & CBE
                    calc_df = df.copy() 
                    calc_df['Calc_Hardness'] = calc_df.apply(calculate_hardness, axis=1)
                    calc_df['Calc_CBE'] = calc_df.apply(calculate_cbe, axis=1).abs() # Absolute CBE
                    
                    col_met1, col_met2, col_met3, col_met4 = st.columns(4)
                    
                    with col_met1:
                        st.metric("Samples (N)", len(calc_df))
                        
                    with col_met2:
                        if 'TDS' in calc_df.columns and pd.api.types.is_numeric_dtype(calc_df['TDS']):
                            min_t = calc_df['TDS'].min()
                            max_t = calc_df['TDS'].max()
                            mean_t = calc_df['TDS'].mean()
                            st.metric("TDS Range", f"{min_t:.0f} - {max_t:.0f}", f"Avg: {mean_t:.0f}")
                        else:
                            st.metric("TDS", "N/A")

                    with col_met3:
                        if 'pH' in calc_df.columns and pd.api.types.is_numeric_dtype(calc_df['pH']):
                             st.metric("pH Range", f"{calc_df['pH'].min():.1f} - {calc_df['pH'].max():.1f}", f"Avg: {calc_df['pH'].mean():.1f}", delta_color="off")
                        else:
                             st.metric("pH", "N/A")

                    with col_met4:
                        if calc_df['Calc_Hardness'].notna().any():
                             st.metric("Hardness (Avg)", f"{calc_df['Calc_Hardness'].mean():.1f} mg/L")
                        else:
                             st.metric("Hardness", "N/A")

                    # --- 2. Auto-Insights ---
                    st.markdown("#### 2. Auto-Insights")
                    insights_list = generate_insights(calc_df, selected_cols)
                    for txt in insights_list:
                        st.info(txt, icon="💡")
                    
                    # --- 3. Data Quality & Correlations ---
                    st.markdown(f"#### 3. {t('correlation')} & QC")
                    
                    c_qc1, c_qc2 = st.columns([1, 2])
                    
                    with c_qc1:
                        st.markdown("**Data Quality (QC)**")
                        mean_cbe = calc_df['Calc_CBE'].mean()
                        max_cbe = calc_df['Calc_CBE'].max()
                        
                        st.metric("Mean CBE %", f"{mean_cbe:.2f}%")
                        if max_cbe > 10:
                            st.warning(f"Max CBE is **{max_cbe:.1f}%** (>10%)\nCheck ion balance.")
                        else:
                            st.success(f"Max CBE is {max_cbe:.1f}% (OK)")
                            
                        if len(calc_df) < 15:
                            st.warning("⚠️ Low sample size (N<15). Correlations may be unreliable.")

                    with c_qc2:
                        st.markdown("**Correlation Heatmap**")
                        if len(selected_cols) >= 2:
                             corr_method = st.radio("Correlation Method", ["pearson", "spearman"], horizontal=True, key="stat_corr_method")
                             
                             numeric_df = calc_df[selected_cols].select_dtypes(include=np.number)
                             if not numeric_df.empty:
                                corr_val = numeric_df.corr(method=corr_method)
                                
                                heat = go.Figure(data=go.Heatmap(
                                    z=corr_val.values, x=corr_val.columns, y=corr_val.index, 
                                    colorscale='RdBu', zmin=-1, zmax=1,
                                    text=corr_val.values.round(2), texttemplate="%{text}"
                                ))
                                heat.update_layout(height=450)
                                st.plotly_chart(heat, use_container_width=True)
                        else:
                             st.info(t('corr_req'))
                    
                    # --- 4. Top Relationships ---
                    st.markdown(f"#### 4. Top Relationships (Scatter)")
                    numeric_df = calc_df[selected_cols].select_dtypes(include=np.number)
                    if not numeric_df.empty and len(selected_cols) >= 2:
                        corr_matrix = numeric_df.corr().abs()
                        upper = corr_matrix.where(np.triu(np.ones(corr_matrix.shape), k=1).astype(bool))
                        top_pairs = upper.stack().sort_values(ascending=False).head(3)
                        
                        if not top_pairs.empty:
                            cols_scatter = st.columns(3)
                            for i, ((c1, c2), val) in enumerate(top_pairs.items()):
                                if i < 3:
                                    with cols_scatter[i]:
                                        fig_sc = px.scatter(calc_df, x=c1, y=c2, title=f"{c1} vs {c2} (r={val:.2f})")
                                        st.plotly_chart(fig_sc, use_container_width=True)
                        else:
                            st.info("No strong correlations found.")
                    else:
                        st.info("Not enough numeric columns for scatter plots.")

                    # --- 5. Full Stats & Matrix ---
                    st.markdown("#### 5. Detailed Data")
                    with st.expander(f"Show {t('general_stats')} Table"):
                        st.dataframe(stats.style.background_gradient(cmap="Blues"), use_container_width=True)
                         
                    c_full1, c_full2 = st.columns([1,3])
                    with c_full1:
                        show_matrix = st.checkbox("Show Scatter Matrix (Slow)?")
                    
                    if show_matrix:
                         st.markdown(f"##### {t('scatter_matrix')}")
                         if len(selected_cols) > 5:
                             st.warning(t('scatter_warning'))
                         
                         sm_fig = px.scatter_matrix(
                             calc_df[selected_cols],
                             dimensions=selected_cols,
                             color=selected_cols[0] if len(selected_cols)>0 else None
                         )
                         sm_fig.update_layout(height=600)
                         st.plotly_chart(sm_fig, use_container_width=True)
            # --- Clustering ---
            with tabs[15]:
                st.subheader(f"🧩 {t('tabs_cluster')} (K-Means Analysis)")
                st.caption(t('cluster_caption'))
                
                # --- A. Settings Panel ---
                col_sets, col_vis = st.columns([1, 2])
                
                with col_sets:
                    st.markdown("##### ⚙️ Settings")
                    
                    # 1. Feature Selection
                    # Determine numeric columns
                    num_cols = df.select_dtypes(include=[np.number]).columns.tolist()
                    # Filter out non-chemical/irrelevant (e.g., Lat, Lon) if needed or let user choose
                    rec_cols = ['Ca','Mg','Na','K','Cl','SO4','HCO3','TDS','pH','EC','SAR','Water Level']
                    default_cluster_cols = [c for c in rec_cols if c in num_cols]
                    if not default_cluster_cols and len(num_cols)>0: default_cluster_cols = num_cols[:3]
                    
                    cluster_features = st.multiselect("Features to Cluster", num_cols, default=default_cluster_cols)
                    
                    # 2. Standardization
                    use_scaling = st.checkbox("Standardize Features (Z-Score)", value=True, help="Crucial for Hydrochemistry (e.g. pH=7 vs TDS=500)")
                    
                    # 3. K Selection
                    n_samples = len(df)
                    max_k = min(6, n_samples - 1) if n_samples > 2 else 2
                    if max_k < 2: max_k = 2
                    
                    k = st.slider(t('cluster_count'), min_value=2, max_value=max_k, value=min(3, max_k))
                    
                    if n_samples < 15:
                        st.warning(f"⚠️ Low sample size (N={n_samples}). Clustering is illustrative only.")
                    
                    if len(cluster_features) < 2:
                        st.error("Select at least 2 features.")
                        # Stop execution locally within this tab content if possible, or just don't run
                        run_clustering = False
                    else:
                        run_clustering = True

                # --- B. Architecture & Processing ---
                if run_clustering:
                    try:
                        X = df[cluster_features].copy().fillna(0)
                        
                        # Scaling
                        if use_scaling:
                            scaler = StandardScaler()
                            X_scaled = scaler.fit_transform(X)
                        else:
                            X_scaled = X.values
                            
                        # K-Means
                        kmeans = KMeans(n_clusters=k, random_state=42, n_init=10)
                        labels = kmeans.fit_predict(X_scaled)
                        
                        # Store in a temporary View DF so we don't mess up main DF globally yet (or maybe we do want to?)
                        # Let's keep it local to viz to avoid state issues unless saved. 
                        # Actually good to have in DF for export.
                        df['Cluster'] = labels.astype(str) 
                        
                        # PCA for Viz
                        pca = PCA(n_components=2)
                        components = pca.fit_transform(X_scaled)
                        df['PC1'] = components[:, 0]
                        df['PC2'] = components[:, 1]
                        var_explained = pca.explained_variance_ratio_
                        
                    except Exception as e:
                        st.error(f"Clustering Error: {e}")
                        run_clustering = False

                # --- C. Visualization (PCA) ---
                with col_vis:
                    if run_clustering:
                        tab_viz1, tab_viz2 = st.tabs(["🔍 PCA View (Recommended)", "📈 Raw Feature View"])
                        
                        with tab_viz1:
                            fig_pca = px.scatter(
                                df, x='PC1', y='PC2', color='Cluster',
                                hover_data=['ID'] + cluster_features,
                                title=f"PCA Cluster View ({var_explained[0]*100:.1f}% + {var_explained[1]*100:.1f}% explained)",
                                template="plotly_white"
                            )
                            st.plotly_chart(fig_pca, use_container_width=True)
                            st.caption("PCA reduces dimensions to 2D while preserving max variance.")
                            
                        with tab_viz2:
                            cx = st.selectbox("X Axis", cluster_features, index=0)
                            cy = st.selectbox("Y Axis", cluster_features, index=1 if len(cluster_features)>1 else 0)
                            fig_raw = px.scatter(
                                df, x=cx, y=cy, color='Cluster', hover_data=['ID'],
                                title=f"Raw Data: {cx} vs {cy}",
                                template="plotly_white"
                            )
                            st.plotly_chart(fig_raw, use_container_width=True)

                # --- D. Cluster Interpretation (Cards) ---
                if run_clustering and run_clustering: # redundant check but ok
                    st.divider()
                    st.markdown("#### 🏷️ Cluster Interpretation")
                    
                    c_cols = st.columns(k)
                    for i in range(k):
                        c_data = df[df['Cluster'] == str(i)]
                        count = len(c_data)
                        
                        # Centroid calculation (Mean of raw data)
                        c_mean = c_data[cluster_features].mean()
                        
                        # Heuristic Labeling
                        dom_ion = ""
                        if 'Na' in c_mean and 'Ca' in c_mean:
                            dom_ion = "Na-Type" if c_mean['Na'] > c_mean['Ca'] else "Ca-Type"
                        
                        salinity = ""
                        if 'TDS' in c_mean:
                            tds = c_mean['TDS']
                            if tds < 500: salinity = "Fresh"
                            elif tds < 1500: salinity = "Brackish"
                            else: salinity = "Saline"
                        
                        with c_cols[i]:
                            st.info(f"**Cluster {i}** (n={count})")
                            st.markdown(f"**{salinity} {dom_ion}**")
                            # Show top key params
                            disp_cols = [c for c in ['TDS', 'Na', 'Cl', 'SO4', 'HCO3'] if c in c_mean]
                            txt = ""
                            for dc in disp_cols:
                                txt += f"- **{dc}:** {c_mean[dc]:.0f}\n"
                            st.markdown(txt)

                    # Export
                    csv_cluster = df[['ID', 'Cluster'] + cluster_features].to_csv(index=False).encode('utf-8')
                    st.download_button("📥 Download Result CSV", csv_cluster, "clustered_data.csv")

            # --- Groundwater Level ---
            with tabs[16]:
                st.subheader(f"💧 {t('tabs_level')}")
                
                if 'Water Level' not in df.columns:
                    st.error(t('water_level_req'))
                else:
                    wl_df = df[['ID', 'Water Level'] + ([c for c in ['Lat', 'Lon'] if c in df.columns])].copy()
                    wl_df = wl_df.dropna(subset=['Water Level'])
                    
                    if wl_df.empty:
                        st.warning(t('no_data'))
                    else:
                        with st.expander(f"ℹ️ {t('classification_method')}", expanded=False):
                            st.markdown(t('level_class_explainer'))

                        # Basic stats in big cards
                        s1, s2, s3 = st.columns(3)
                        s1.metric(t('min_level'), f"{wl_df['Water Level'].min():.2f} m")
                        s2.metric(t('avg_level'), f"{wl_df['Water Level'].mean():.2f} m")
                        s3.metric(t('max_level'), f"{wl_df['Water Level'].max():.2f} m")

                        # Simple classification by quartiles
                        q25 = float(wl_df['Water Level'].quantile(0.25))
                        q75 = float(wl_df['Water Level'].quantile(0.75))
                        def wl_class(v):
                            if v <= q25: return 'Low (<25%)'
                            if v >= q75: return 'High (>75%)'
                            return 'Normal'
                        wl_df['Status'] = wl_df['Water Level'].apply(wl_class)

                        c1, c2 = st.columns(2)
                        with c1:
                            st.markdown(f"#### {t('level_chart')}")
                            fig_wl = px.bar(
                                wl_df, x='ID', y='Water Level', color='Status', 
                                title=t('static_levels'),
                                color_discrete_map={'Low (<25%)': 'blue', 'Normal': 'green', 'High (>75%)': 'red'}
                            )
                            st.plotly_chart(fig_wl, use_container_width=True)
                        with c2:
                            st.markdown(f"#### {t('histogram')}")
                            fig_hist = px.histogram(wl_df, x='Water Level', nbins=10, title=t('freq_chart'))
                            fig_hist.update_layout(bargap=0.1)
                            st.plotly_chart(fig_hist, use_container_width=True)

                        # Map/contours if coordinates exist
                        if all(col in wl_df.columns for col in ['Lat', 'Lon']):
                            st.markdown(f"### 🗺️ {t('level_map')}")
                            mm1, mm2 = st.columns([1, 3])
                            with mm1:
                                method = st.selectbox(t('interpolation'), ["Cubic", "Kriging (RBF)"], key='wl_interp')
                                levels = st.slider(t('isoline_count'), 5, 25, 12, key='wl_levels')
                            with mm2:
                                map_df = wl_df.dropna(subset=['Lat', 'Lon', 'Water Level'])
                                mp = folium.Map([map_df['Lat'].mean(), map_df['Lon'].mean()], zoom_start=12)
                                
                                # Use existing contour function
                                l, f, cm = generate_contours(map_df, 'Water Level', levels, 'Blues', method)
                                
                                if f: folium.GeoJson(f, style_function=lambda x: {'fillColor':x['properties']['fill'], 'color':'black', 'weight':0.5, 'fillOpacity':0.6}).add_to(mp)
                                if cm: mp.add_child(cm)
                                
                                for _,r in map_df.iterrows():
                                    folium.CircleMarker(
                                        [r['Lat'],r['Lon']], 
                                        radius=4, 
                                        color='black', fill=True, fillColor='white',
                                        tooltip=f"{r['ID']}: {r['Water Level']}m"
                                    ).add_to(mp)
                                    
                                st_folium(mp, height=500, use_container_width=True)
                        else:
                            st.info(t('lat_lon_required'))

                        with st.expander(t('assumptions')):
                            st.markdown(t('level_assumptions_text'))
    # --- ARXIV MODULU ---
    elif st.session_state['page_key'] == "archive":
        d = get_user_projects(st.session_state.get('username','Qonaq'))
        for r in d:
            with st.expander(f"{r[1]}"):
                st.download_button("Excel", r[3], f"{r[1]}.xlsx")


# --- APP ENTRYPOINT (Login removed) ---
# Always start the app directly. Keep a guest username unless you later re-introduce auth.
st.session_state['logged_in'] = True
if 'username' not in st.session_state or not st.session_state['username']:
    st.session_state['username'] = "Guest"

main_app()